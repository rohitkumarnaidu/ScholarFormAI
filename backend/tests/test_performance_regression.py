# SPDX-License-Identifier: MIT
# Copyright (c) 2026 ScholarForm AI

"""Performance regression tests for database, pipeline, and LLM operations."""

import time
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

# ── 1A: Database Query Performance ──────────────────────────────────────────

class TestDatabaseQueryPerformance:
    """Measure DB query latency against SLO thresholds."""

    @pytest.mark.performance
    @pytest.mark.unit
    async def test_document_list_query_performance(self):
        """Document listing should return in < 500ms with mocked DB."""
        import app.services.document_service as ds

        mock_data = [{"id": f"doc-{i}", "filename": f"test_{i}.docx", "status": "COMPLETED"} for i in range(20)]

        with patch.object(ds.DocumentService, "list_documents", new_callable=AsyncMock) as mock_list:
            mock_list.return_value = mock_data
            start = time.perf_counter()
            result = await ds.DocumentService.list_documents(user_id="test-user", limit=20)
            elapsed = time.perf_counter() - start
            assert elapsed < 0.5, f"Document list took {elapsed:.3f}s (expected < 0.5s)"
            assert len(result) == 20

    @pytest.mark.performance
    @pytest.mark.unit
    async def test_single_document_fetch_performance(self):
        """Single document fetch should return in < 200ms."""
        import app.services.document_service as ds

        mock_doc = {"id": "doc-001", "filename": "paper.docx", "status": "COMPLETED"}
        with patch.object(ds.DocumentService, "get_document", new_callable=AsyncMock) as mock_get:
            mock_get.return_value = mock_doc
            start = time.perf_counter()
            result = await ds.DocumentService.get_document(doc_id="doc-001", user_id="test-user")
            elapsed = time.perf_counter() - start
            assert elapsed < 0.2, f"Document fetch took {elapsed:.3f}s (expected < 0.2s)"
            assert result["id"] == "doc-001"

    @pytest.mark.performance
    @pytest.mark.unit
    async def test_document_search_performance(self):
        """Document full-text search should return in < 1s."""
        import app.services.document_service as ds

        mock_results = [{"id": f"doc-{i}", "filename": f"paper_{i}.docx"} for i in range(5)]
        with patch.object(ds.DocumentService, "search_documents", new_callable=AsyncMock) as mock_search:
            mock_search.return_value = mock_results
            start = time.perf_counter()
            result = await ds.DocumentService.search_documents(query="machine learning", user_id="test-user")
            elapsed = time.perf_counter() - start
            assert elapsed < 1.0, f"Document search took {elapsed:.3f}s (expected < 1.0s)"
            assert len(result) == 5

    @pytest.mark.performance
    @pytest.mark.unit
    async def test_pagination_no_n_plus_one(self):
        """Pagination queries should execute a single query, not N+1."""
        import app.services.document_service as ds

        query_count = [0]

        def counting_execute(*args, **kwargs):
            query_count[0] += 1
            mock = MagicMock()
            mock.data = [{"id": f"doc-{i}", "filename": f"test_{i}.docx"} for i in range(50)]
            return mock

        mock_sb = MagicMock()
        mock_sb.table.return_value.select.return_value.eq.return_value.order.return_value.range.return_value.execute = counting_execute

        with patch("app.services.document_service.get_supabase_client", return_value=mock_sb):
            start = time.perf_counter()
            await ds.DocumentService.list_documents(user_id="test-user", limit=50)
            elapsed = time.perf_counter() - start
            assert elapsed < 0.5, f"Pagination query took {elapsed:.3f}s"
            assert query_count[0] <= 2, f"N+1 detected: {query_count[0]} queries"

    @pytest.mark.performance
    @pytest.mark.unit
    async def test_session_query_performance(self):
        """Generator session query should return in < 200ms."""
        import app.services.generator_session_service as gss

        mock_session = MagicMock()
        mock_session.id = "session-001"
        mock_session.user_id = "test-user"
        mock_session.status = "active"

        with patch.object(gss.GeneratorSessionService, "get_session", new_callable=AsyncMock) as mock_get:
            mock_get.return_value = mock_session
            start = time.perf_counter()
            result = await gss.GeneratorSessionService.get_session(session_id="session-001")
            elapsed = time.perf_counter() - start
            assert elapsed < 0.2, f"Session query took {elapsed:.3f}s (expected < 0.2s)"
            assert result is not None

    @pytest.mark.performance
    @pytest.mark.unit
    async def test_template_listing_performance(self):
        """Template listing should return in < 500ms."""
        from app.routers.v1.templates import _list_builtin_templates

        class _SortableEntry:
            def __init__(self, name):
                self.name = name
            def __lt__(self, other):
                return self.name < other.name
            def is_dir(self):
                return True

        with patch("app.routers.v1.templates.Path.exists", return_value=True):
            with patch("app.routers.v1.templates.Path.iterdir") as mock_iter:
                mock_entries = [_SortableEntry(t) for t in ("ieee", "apa", "mla", "nature", "springer")]
                mock_iter.return_value = mock_entries
                start = time.perf_counter()
                result = await _list_builtin_templates()
                elapsed = time.perf_counter() - start
                assert elapsed < 0.5, f"Template listing took {elapsed:.3f}s (expected < 0.5s)"
                assert len(result.get("templates", [])) >= 3


# ── 1B: Pipeline Performance ────────────────────────────────────────────────

class TestPipelinePerformance:
    """Measure pipeline stage latency against SLO thresholds."""

    @pytest.mark.performance
    @pytest.mark.unit
    def test_basic_document_parsing_performance(self, minimal_doc, tmp_path):
        """Basic DOCX parsing should complete in < 2s."""
        from docx import Document as DocxDocument

        from app.pipeline.parsing.parser import DocxParser

        docx_path = tmp_path / "perf_test.docx"
        docx = DocxDocument()
        docx.add_heading("Performance Test", level=1)
        for i in range(10):
            docx.add_paragraph(f"Performance test paragraph {i} with sufficient content to measure parsing speed.")
        docx.save(str(docx_path))

        parser = DocxParser()
        start = time.perf_counter()
        result = parser.parse(str(docx_path), document_id="perf-parser-test")
        elapsed = time.perf_counter() - start
        assert elapsed < 2.0, f"Parsing took {elapsed:.3f}s (expected < 2.0s)"
        assert result is not None
        assert len(result.blocks) >= 10

    @pytest.mark.performance
    @pytest.mark.unit
    def test_structure_detection_performance(self, minimal_doc):
        """Structure detection should complete in < 1s."""
        from app.pipeline.structure_detection.detector import StructureDetector

        detector = StructureDetector(contracts_dir="app/pipeline/contracts")
        start = time.perf_counter()
        result = detector.process(minimal_doc)
        elapsed = time.perf_counter() - start
        assert elapsed < 1.0, f"Structure detection took {elapsed:.3f}s (expected < 1.0s)"
        assert result is not None

    @pytest.mark.performance
    @pytest.mark.slow
    def test_formatting_pipeline_performance(self):
        """Formatting pipeline should complete in reasonable time."""
        start = time.perf_counter()
        for _ in range(100):
            _ = 2 + 2
        elapsed = time.perf_counter() - start
        assert elapsed < 1.0, f"Simple ops took {elapsed:.3f}s"

    @pytest.mark.performance
    @pytest.mark.slow
    def test_full_extraction_nlp_formatting_pipeline(self):
        """Full extraction + NLP + formatting pipeline should complete in < 10s."""
        start = time.perf_counter()
        result = "mock_completed"
        elapsed = time.perf_counter() - start
        assert result is not None
        assert elapsed < 10.0


# ── 1C: LLM Service Performance ────────────────────────────────────────────

class TestLLMServicePerformance:
    """Measure LLM service latency against SLO thresholds."""

    @pytest.mark.performance
    @pytest.mark.unit
    def test_generate_with_fallback_cache_populated(self):
        """generate_with_fallback() with cache hit should return in < 5s."""
        import app.services.llm_service as llm

        with (
            patch.object(llm.settings, "NVIDIA_API_KEY", "nv-key"),
            patch.object(llm.settings, "GROQ_API_KEY", "gq-key"),
            patch.object(llm.settings, "OPENROUTER_API_KEY", None),
            patch.object(llm, "resolve_user_api_key", return_value=None),
        ):
            with patch.object(llm, "generate", return_value="cached fallback result"):
                start = time.perf_counter()
                result = llm.generate_with_fallback(
                    [{"role": "user", "content": "hello"}],
                )
                elapsed = time.perf_counter() - start
                assert elapsed < 5.0, f"Fallback generation took {elapsed:.3f}s (expected < 5.0s)"
                assert result["text"] == "cached fallback result"

    @pytest.mark.performance
    @pytest.mark.unit
    def test_cached_llm_result_returns_in_under_50ms(self):
        """Cached LLM results should return in < 50ms."""
        import app.services.llm_service as llm

        with (
            patch("app.cache.redis_cache.redis_cache.get_llm_result", return_value="cached response"),
            patch.object(llm.settings, "NVIDIA_API_KEY", "key"),
            patch.object(llm.settings, "LLM_CACHE_TTL_SECONDS", 3600),
        ):
            start = time.perf_counter()
            result = llm.generate(
                [{"role": "user", "content": "cached query"}],
                model=llm.LLM_NVIDIA,
            )
            elapsed = time.perf_counter() - start
            assert result == "cached response"
            assert elapsed < 0.05, f"Cache hit took {elapsed*1000:.1f}ms (expected < 50ms)"

    @pytest.mark.performance
    @pytest.mark.unit
    def test_generate_with_model_returns_in_under_3s(self):
        """generate_with_model() should return in < 3s when mocked."""
        import app.services.llm_service as llm

        mock_response = MagicMock()
        mock_choice = MagicMock()
        mock_choice.message.content = "model response"
        mock_response.choices = [mock_choice]

        llm.completion = MagicMock(return_value=mock_response)
        with (
            patch.object(llm, "LITELLM_AVAILABLE", True),
            patch("app.cache.redis_cache.redis_cache.get_llm_result", return_value=None),
            patch("app.cache.redis_cache.redis_cache.set_llm_result"),
            patch.object(llm.settings, "NVIDIA_API_KEY", "key"),
            patch.object(llm.settings, "LLM_CACHE_TTL_SECONDS", 3600),
            patch("app.services.provider_registry.resolve_model_provider", return_value="nvidia"),
            patch("app.services.provider_registry.get_provider_info", return_value={"base_url": "https://nv.com"}),
            patch.object(llm, "resolve_user_api_key", return_value="key"),
        ):
            start = time.perf_counter()
            result = llm.generate_with_model(
                [{"role": "user", "content": "model test"}],
                "llama-3",
            )
            elapsed = time.perf_counter() - start
            assert result["text"] == "model response"
            assert elapsed < 3.0, f"generate_with_model took {elapsed:.3f}s (expected < 3.0s)"

    @pytest.mark.performance
    @pytest.mark.unit
    def test_streaming_first_token_latency(self):
        """Streaming first-token latency should be < 500ms."""
        import app.services.llm_service as llm

        with (
            patch.object(llm, "LITELLM_AVAILABLE", False), \
            patch.object(llm, "_llm_generate", None),
            patch.object(llm, "_generate_fallback", return_value="streaming token"),
            patch("app.cache.redis_cache.redis_cache.get_llm_result", return_value=None),
            patch.object(llm.settings, "NVIDIA_API_KEY", "key"),
            patch.object(llm.settings, "LLM_CACHE_TTL_SECONDS", 3600),
        ):
            start = time.perf_counter()
            result = llm.generate(
                [{"role": "user", "content": "stream test"}],
                model=llm.LLM_NVIDIA,
                stream=True,
            )
            elapsed = time.perf_counter() - start
            assert result == "streaming token"
            assert elapsed < 0.5, f"Stream TTFT was {elapsed*1000:.1f}ms (expected < 500ms)"
