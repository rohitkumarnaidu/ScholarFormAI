from __future__ import annotations

import html
import logging
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.domain.models import (
    DomainManuscript,
    DomainParagraph,
    DomainReference,
    DomainSection,
)
from app.schemas.models import FormattingOptions, Manuscript
from app.services.style_registry import FormattingStyle

logger = logging.getLogger(__name__)


class ReferenceRenderer:
    """CSL citation list string formatting and bibliography rendering."""

    def format_reference_string(self, ref: DomainReference, style: FormattingStyle, index: int = 1) -> tuple[str, str]:
        """Returns (prefix, body_text) tuple for a reference entry."""
        numbered = "numbered" in style.reference_format

        if numbered:
            prefix = f"[{index}] "
        else:
            authors_text = ""
            if ref.authors:
                formatted_authors = []
                for a in ref.authors:
                    last = a.last_name or (a.name.split()[-1] if a.name else "")
                    first_initial = (a.first_name[0] + ".") if a.first_name else ""
                    if last:
                        formatted_authors.append(f"{last}, {first_initial}" if first_initial else last)
                    elif a.name:
                        formatted_authors.append(a.name)
                authors_text = ", ".join(formatted_authors)
            prefix = (
                f"{authors_text} ({ref.year}). "
                if authors_text and ref.year
                else (f"{authors_text} " if authors_text else "")
            )

        title_text = ref.title or ""
        if ref.journal:
            title_text = f"{title_text}. {ref.journal}"
        if ref.volume:
            title_text = f"{title_text}, {ref.volume}"
            if ref.issue:
                title_text = f"{title_text}({ref.issue})"
        if ref.pages:
            title_text = f"{title_text}, {ref.pages}"
        if ref.doi:
            title_text = f"{title_text}. https://doi.org/{ref.doi}"

        return prefix, title_text

    def render_references(self, doc: Document, references: list[DomainReference], style: FormattingStyle):
        if not references:
            return

        doc.add_page_break()

        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run("References")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = style.font_family

        for i, ref in enumerate(references, 1):
            prefix, body_text = self.format_reference_string(ref, style, index=i)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)

            run_pre = p.add_run(prefix)
            run_pre.font.size = Pt(style.font_size)
            run_pre.font.name = style.font_family

            run_body = p.add_run(body_text)
            run_body.font.size = Pt(style.font_size)
            run_body.font.name = style.font_family


class DocumentLayoutEngine:
    """OpenXML document margins, typography, page setup, running headers, title page, and section styling."""

    def __init__(self, reference_renderer: ReferenceRenderer | None = None):
        self.reference_renderer = reference_renderer or ReferenceRenderer()

    def set_margins(self, doc: Document, style: FormattingStyle):
        for section in doc.sections:
            section.top_margin = Inches(style.margin_inches)
            section.bottom_margin = Inches(style.margin_inches)
            section.left_margin = Inches(style.margin_inches)
            section.right_margin = Inches(style.margin_inches)

    def set_default_font(self, doc: Document, style: FormattingStyle):
        style_obj = doc.styles["Normal"]
        font = style_obj.font
        font.name = style.font_family
        font.size = Pt(style.font_size)
        pf = style_obj.paragraph_format
        pf.line_spacing = style.line_spacing
        pf.space_after = Pt(style.paragraph_spacing)

    def apply_line_spacing(self, doc: Document, style: FormattingStyle):
        style_obj = doc.styles["Normal"]
        style_obj.paragraph_format.line_spacing = style.line_spacing

    def create_title_page(self, doc: Document, manuscript: DomainManuscript, style: FormattingStyle):
        for _ in range(3):
            doc.add_paragraph()

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(manuscript.title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = style.font_family

        authors_text = (
            ", ".join(
                f"{a.first_name} {a.last_name}".strip() if (a.first_name or a.last_name) else a.name
                for a in manuscript.authors
            )
            if manuscript.authors
            else ""
        )
        if authors_text:
            auth_para = doc.add_paragraph()
            auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = auth_para.add_run(authors_text)
            run.font.size = Pt(style.font_size)
            run.font.name = style.font_family

        if manuscript.corresponding_author:
            corr = manuscript.corresponding_author
            corr_name = (
                f"{corr.first_name} {corr.last_name}".strip() if (corr.first_name or corr.last_name) else corr.name
            )
            corr_para = doc.add_paragraph()
            corr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = corr_para.add_run(f"Corresponding Author: {corr_name}")
            run.font.size = Pt(10)
            run.font.name = style.font_family
            if corr.email:
                email_para = doc.add_paragraph()
                email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = email_para.add_run(corr.email)
                run.font.size = Pt(10)

        doc.add_page_break()

    def add_running_header(self, doc: Document, manuscript: DomainManuscript, style: FormattingStyle):
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(self.generate_running_head(manuscript.title))
            run.font.size = Pt(10)
            run.font.name = style.font_family

    def generate_running_head(self, title: str) -> str:
        words = re.sub(r"[^a-zA-Z0-9\s]", "", title).split()
        short = " ".join(words[:5])
        return short.upper() if len(short) <= 50 else short[:50].upper()

    def add_abstract(self, doc: Document, manuscript: DomainManuscript, style: FormattingStyle):
        if not manuscript.abstract:
            return

        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run("Abstract")
        run.bold = True
        run.font.size = Pt(style.font_size)
        run.font.name = style.font_family

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(style.first_line_indent)
        run = p.add_run(manuscript.abstract)
        run.font.size = Pt(style.font_size)
        run.font.name = style.font_family

    def add_keywords(self, doc: Document, manuscript: DomainManuscript, style: FormattingStyle):
        if not manuscript.keywords:
            return

        p = doc.add_paragraph()
        run = p.add_run("Keywords: ")
        run.bold = True
        run.font.size = Pt(style.font_size)
        run.font.name = style.font_family
        run = p.add_run(", ".join(manuscript.keywords))
        run.font.size = Pt(style.font_size)
        run.font.name = style.font_family

    def create_body(self, doc: Document, manuscript: DomainManuscript, style: FormattingStyle):
        for section in manuscript.sections:
            self.render_section(doc, section, style, 1)

    def render_section(self, doc: Document, section: DomainSection, style: FormattingStyle, level: int):
        heading_style = style.heading_styles.get(level, {})
        hs = doc.add_paragraph()

        align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        hs.alignment = align_map.get(heading_style.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)

        heading_text = section.heading or section.title
        run = hs.add_run(heading_text)
        run.bold = heading_style.get("bold", level == 1)
        run.italic = heading_style.get("italic", False)
        run.font.size = Pt(heading_style.get("font_size", 14 - level))
        run.font.name = style.font_family

        for para in section.content:
            self.render_paragraph(doc, para, style)

        for subsection in section.subsections:
            self.render_section(doc, subsection, style, level + 1)

    def render_paragraph(self, doc: Document, para: DomainParagraph | Any, style: FormattingStyle):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(style.first_line_indent)

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        text = getattr(para, "text", str(para))
        alignment = getattr(para, "alignment", None)
        para_style = getattr(para, "style", None)
        bullet = getattr(para, "bullet", False)

        if alignment:
            p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        if bullet:
            p.style = doc.styles["List Bullet"]

        run = p.add_run(text)
        run.font.name = style.font_family
        run.font.size = Pt(style.font_size)

        if para_style == "italic":
            run.italic = True
        elif para_style == "bold":
            run.bold = True

    def add_page_numbers(self, doc: Document):
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = p.add_run()
            fld_char1 = OxmlElement("w:fldChar")
            fld_char1.set(qn("w:fldCharType"), "begin")
            run._element.append(fld_char1)

            run2 = p.add_run()
            instr_text = OxmlElement("w:instrText")
            instr_text.set(qn("xml:space"), "preserve")
            instr_text.text = " PAGE "
            run2._element.append(instr_text)

            run3 = p.add_run()
            fld_char2 = OxmlElement("w:fldChar")
            fld_char2.set(qn("w:fldCharType"), "end")
            run3._element.append(fld_char2)

    def render_document(
        self,
        manuscript: DomainManuscript,
        style: FormattingStyle,
        output_path: str,
        options: FormattingOptions | None = None,
        reference_renderer: ReferenceRenderer | None = None,
    ) -> str:
        doc = Document()
        ref_renderer = reference_renderer or self.reference_renderer

        self.set_margins(doc, style)
        self.set_default_font(doc, style)
        self.apply_line_spacing(doc, style)

        if style.title_page:
            self.create_title_page(doc, manuscript, style)

        if style.running_header:
            self.add_running_header(doc, manuscript, style)

        self.add_abstract(doc, manuscript, style)
        self.add_keywords(doc, manuscript, style)
        self.create_body(doc, manuscript, style)
        ref_renderer.render_references(doc, manuscript.references, style)

        if style.page_numbers:
            self.add_page_numbers(doc)

        doc.save(output_path)
        logger.info("Saved formatted document to %s", output_path)
        return output_path


class PageEstimator:
    """In-memory paragraph and page calculation (without re-parsing saved docx from disk)."""

    def estimate_pages(self, target: DomainManuscript | Manuscript | str | Any) -> int:
        if isinstance(target, str):
            try:
                from docx import Document as DocxDoc

                doc = DocxDoc(target)
                para_count = len(doc.paragraphs)
                return max(1, para_count // 25)
            except Exception:
                return 1

        if not isinstance(target, DomainManuscript) and hasattr(target, "sections"):
            target = DomainManuscript.from_pydantic(target)

        if isinstance(target, DomainManuscript):
            return self.estimate_pages_from_manuscript(target)

        return 1

    def estimate_pages_from_manuscript(self, manuscript: DomainManuscript) -> int:
        para_count = 0
        if manuscript.title:
            para_count += 1
        if manuscript.authors:
            para_count += 1
        if manuscript.abstract:
            para_count += 2
        if manuscript.keywords:
            para_count += 1

        def count_section_paragraphs(section: DomainSection) -> int:
            cnt = 1
            cnt += len(section.content)
            for sub in section.subsections:
                cnt += count_section_paragraphs(sub)
            return cnt

        for section in manuscript.sections:
            para_count += count_section_paragraphs(section)

        if manuscript.references:
            para_count += 1 + len(manuscript.references)

        return max(1, para_count // 25)


class HTMLPreviewRenderer:
    """Plaintext/HTML preview generation."""

    def generate_html_preview(self, manuscript: DomainManuscript | Manuscript, style: FormattingStyle) -> str:
        if not isinstance(manuscript, DomainManuscript):
            manuscript = DomainManuscript.from_pydantic(manuscript)

        escaped_title = html.escape(manuscript.title)
        parts = [
            f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>{escaped_title}</title>
<style>
  body {{ """
            f"""font-family: {style.font_family}, serif; font-size: {style.font_size}pt; """
            f"""line-height: {style.line_spacing}; margin: {style.margin_inches}in; }}
  h1 {{ font-size: 16pt; font-weight: bold; text-align: center; }}
  h2 {{ font-size: 14pt; font-weight: bold; }}
  h3 {{ font-size: 12pt; font-weight: bold; font-style: italic; }}
  .abstract {{ margin: 1em 0; }}
  .keywords {{ margin: 0.5em 0; }}
  .references {{ margin-top: 2em; }}
  .ref-entry {{ margin-left: 0.5in; text-indent: -0.5in; }}
  @media print {{ .page-break {{ page-break-after: always; }} }}
</style></head><body>"""
        ]

        parts.append(f"<h1>{escaped_title}</h1>")
        if manuscript.authors:
            formatted_authors = [
                html.escape(f"{a.first_name} {a.last_name}".strip() if (a.first_name or a.last_name) else a.name)
                for a in manuscript.authors
            ]
            parts.append(f"<p style='text-align:center;'>{', '.join(formatted_authors)}</p>")

        if manuscript.abstract:
            parts.append(f"<div class='abstract'><h2>Abstract</h2><p>{html.escape(manuscript.abstract)}</p></div>")

        if manuscript.keywords:
            parts.append(
                f"<p class='keywords'><strong>Keywords:</strong> "
                f"{', '.join(html.escape(kw) for kw in manuscript.keywords)}</p>"
            )

        for section in manuscript.sections:
            parts.append(self._render_section_html(section, 2))

        if manuscript.references:
            parts.append("<div class='references'><h2>References</h2>")
            for ref in manuscript.references:
                formatted_authors = []
                for a in ref.authors:
                    last = a.last_name or (a.name.split()[-1] if a.name else "")
                    first_initial = (a.first_name[0] + ".") if a.first_name else ""
                    if last:
                        formatted_authors.append(
                            f"{html.escape(last)}, {html.escape(first_initial)}" if first_initial else html.escape(last)
                        )
                    elif a.name:
                        formatted_authors.append(html.escape(a.name))
                authors_str = ", ".join(formatted_authors) if formatted_authors else ""
                year_str = html.escape(ref.year or "")
                title_str = html.escape(ref.title or "")
                journal_str = html.escape(ref.journal or "")
                parts.append(
                    f"<p class='ref-entry'>{authors_str} ({year_str}). {title_str}. <em>{journal_str}</em>.</p>"
                )
            parts.append("</div>")

        parts.append("</body></html>")
        return "\n".join(parts)

    def _render_section_html(self, section: DomainSection, level: int) -> str:
        tag = f"h{min(level, 6)}"
        heading_text = section.heading or section.title
        parts = [f"<{tag}>{html.escape(heading_text)}</{tag}>"]
        for para in section.content:
            text = getattr(para, "text", str(para))
            parts.append(f"<p>{html.escape(text)}</p>")
        for sub in section.subsections:
            parts.append(self._render_section_html(sub, level + 1))
        return "\n".join(parts)


class ManuscriptFormatter:
    """Delegator class providing backwards compatibility for legacy manuscript formatting calls."""

    def __init__(
        self,
        layout_engine: DocumentLayoutEngine | None = None,
        reference_renderer: ReferenceRenderer | None = None,
        page_estimator: PageEstimator | None = None,
        html_renderer: HTMLPreviewRenderer | None = None,
    ):
        self.reference_renderer = reference_renderer or ReferenceRenderer()
        self.layout_engine = layout_engine or DocumentLayoutEngine(reference_renderer=self.reference_renderer)
        self.page_estimator = page_estimator or PageEstimator()
        self.html_renderer = html_renderer or HTMLPreviewRenderer()

    def format(
        self,
        manuscript: DomainManuscript | Manuscript,
        style: FormattingStyle,
        output_path: str,
        options: FormattingOptions | None = None,
    ) -> str:
        if not isinstance(manuscript, DomainManuscript):
            manuscript = DomainManuscript.from_pydantic(manuscript)

        return self.layout_engine.render_document(
            manuscript=manuscript,
            style=style,
            output_path=output_path,
            options=options,
            reference_renderer=self.reference_renderer,
        )

    def estimate_pages(self, target: DomainManuscript | Manuscript | str) -> int:
        return self.page_estimator.estimate_pages(target)

    def generate_html_preview(self, manuscript: DomainManuscript | Manuscript, style: FormattingStyle) -> str:
        if not isinstance(manuscript, DomainManuscript):
            manuscript = DomainManuscript.from_pydantic(manuscript)

        return self.html_renderer.generate_html_preview(manuscript, style)
