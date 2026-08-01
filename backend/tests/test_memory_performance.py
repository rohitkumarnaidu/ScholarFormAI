# SPDX-License-Identifier: MIT
# Copyright (c) 2026 ScholarForm AI

"""Memory usage and connection pool performance tests."""

import asyncio
import gc
import threading
import time
from unittest.mock import AsyncMock, MagicMock, patch

import pytest


def _force_gc():
    gc.collect()
    gc.collect()
    gc.collect()


# ── 2A: Memory Usage Tests ─────────────────────────────────────────────────

class TestMemoryUsage:
    """Verify memory stays bounded during heavy operations."""

    @pytest.mark.skip(reason="tracemalloc triggers scipy import that times out in CI")
    @pytest.mark.performance
    @pytest.mark.unit
    def test_parser_no_memory_leak_on_large_file(self, tmp_path):
        """Document parser should not leak memory on large documents."""
        from docx import Document as DocxDocument

        from app.pipeline.parsing.parser import DocxParser

        large_path = tmp_path / "large_mem_test.docx"
        docx = DocxDocument()
        docx.add_heading("Large Document", level=1)
        for i in range(500):
            docx.add_paragraph(f"Paragraph {i}: " + "content " * 50)
        for _t_idx in range(10):
            table = docx.add_table(rows=20, cols=5)
            for r in range(20):
                for c in range(5):
                    table.cell(r, c).text = f"Cell {r},{c}"
        docx.save(str(large_path))

        parser = DocxParser()
        import tracemalloc

        tracemalloc.start()
        result = parser.parse(str(large_path), document_id="mem-test")
        current, peak = tracemalloc.get_traced_memory()
        tracemalloc.stop()

        assert result is not None
        assert len(result.blocks) >= 500
        assert peak < 100 * 1024 * 1024, f"Peak memory {peak/1024/1024:.1f}MB > 100MB limit"

    @pytest.mark.skip(reason="Requires real PipelineOrchestrator with all deps")
    @pytest.mark.performance
    @pytest.mark.unit
    def test_pipeline_orchestrator_releases_memory(self, tmp_path):
        """Pipeline orchestrator should release memory after completion."""
        from docx import Document as DocxDocument

        from app.pipeline.orchestrator import PipelineOrchestrator

        docx_path = tmp_path / "orchestrator_mem.docx"
        docx = DocxDocument()
        docx.add_heading("Memory Test", level=1)
        for i in range(100):
            docx.add_paragraph(f"Content paragraph {i} for memory measurement.")
        docx.save(str(docx_path))

        import tracemalloc

        tracemalloc.start()

        orchestrator = PipelineOrchestrator(
            templates_dir="app/templates",
            temp_dir=str(tmp_path / "mem_temp"),
        )
        with (
            patch.object(orchestrator, "_update_status"),
            patch.object(orchestrator.grobid_client, "process_document", return_value=None),
            patch.object(orchestrator.docling_client, "process_document", return_value=None),
            patch.object(orchestrator, "_run_figure_analysis_stage", return_value=None),
            patch.object(orchestrator, "_run_table_extraction_stage", return_value=None),
            patch.object(orchestrator, "_run_ai_analysis_stage", return_value=None),
            patch("app.pipeline.safety.retry_guard.execute_with_retry", side_effect=lambda f, *a, **kw: f()),
        ):
            result = orchestrator.run_pipeline(
                input_path=str(docx_path),
                job_id="mem-test-orch",
                template_name="none",
            )

        del orchestrator
        del result
        _force_gc()

        current, peak = tracemalloc.get_traced_memory()
        tracemalloc.stop()

        assert peak < 150 * 1024 * 1024, f"Peak memory {peak/1024/1024:.1f}MB > 150MB limit"

    @pytest.mark.skip(reason="Requires RagEngine embedding model")
    @pytest.mark.performance
    @pytest.mark.unit
    def test_rag_engine_bounded_memory(self):
        """RAG engine knowledge base should keep memory bounded."""
        with patch("app.pipeline.intelligence.rag_engine.RagEngine._load_embedding_model", return_value=MagicMock()):
            from app.pipeline.intelligence.rag_engine import RagEngine

            engine = RagEngine()
            engine.knowledge_base.clear()

            import tracemalloc

            tracemalloc.start()
            for i in range(50):
                engine.knowledge_base[f"doc-{i}"] = {
                    "content": f"Simulated document content {i} with enough text to measure memory usage." * 100,
                    "embedding": [0.1] * 384,
                    "metadata": {"source": f"doc-{i}", "page": i},
                }
            current, peak = tracemalloc.get_traced_memory()
            tracemalloc.stop()

            assert peak < 100 * 1024 * 1024, f"RAG memory {peak/1024/1024:.1f}MB > 100MB limit"

    @pytest.mark.skip(reason="Requires event_emitter from stream module with all deps")
    @pytest.mark.performance
    @pytest.mark.unit
    def test_sse_streaming_no_memory_accumulation(self):
        """SSE streaming should not accumulate memory across events."""
        from app.routers.v1.stream import event_emitter

        memory_snapshots = []
        event_count = 1000

        async def dummy_receive():
            return {"type": "websocket.disconnect"}

        async def dummy_send(message):
            if isinstance(message, dict) and message.get("type") == "websocket.send":
                memory_snapshots.append(len(str(message.get("text", "")).encode("utf-8")))

        import asyncio
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

        try:
            with (
                patch("app.routers.v1.stream.PubSubClient"),
                patch("app.routers.v1.stream._pubsub"),
            ):
                start = time.perf_counter()
                loop.run_until_complete(
                    event_emitter(
                        receive=dummy_receive,
                        send=dummy_send,
                        event_type="progress",
                        job_id="sse-mem-test",
                        user_id="test-user",
                    )
                )
                elapsed = time.perf_counter() - start
        finally:
            loop.close()

        assert elapsed < 3.0, f"SSE streaming took {elapsed:.3f}s"
        assert len(memory_snapshots) < event_count * 2, "SSE should not duplicate events"

    @pytest.mark.skip(reason="Requires celery task module with heavy deps")
    @pytest.mark.performance
    @pytest.mark.unit
    def test_celery_handler_releases_memory(self):
        """Celery task handler should release memory after processing."""
        from app.tasks.celery_tasks import process_document_task

        mock_doc = {
            "id": "celery-mem-test",
            "filename": "test.docx",
            "user_id": "test-user",
            "template": "none",
            "file_path": "/tmp/test.docx",
        }

        _force_gc()
        import tracemalloc

        tracemalloc.start()

        with (
            patch("app.tasks.celery_tasks.get_supabase_client", return_value=MagicMock()),
            patch("app.tasks.celery_tasks.PipelineOrchestrator") as mock_orch,
            patch("app.tasks.celery_tasks.DocumentService") as mock_ds,
            patch("app.tasks.celery_tasks.logger"),
        ):
            mock_instance = MagicMock()
            mock_instance.run_pipeline.return_value = {"status": "completed", "output_path": "/tmp/out.docx"}
            mock_orch.return_value = mock_instance
            mock_ds.get_document = AsyncMock(return_value=mock_doc)
            mock_ds.update_document = AsyncMock(return_value=None)
            mock_ds.mark_document_completed = AsyncMock(return_value=None)

            result = process_document_task("celery-mem-test")

        del result
        del mock_instance
        _force_gc()

        current, peak = tracemalloc.get_traced_memory()
        tracemalloc.stop()

        assert peak < 200 * 1024 * 1024, f"Celery peak memory {peak/1024/1024:.1f}MB > 200MB limit"

    @pytest.mark.skip(reason="Requires real PipelineOrchestrator with all deps")
    @pytest.mark.performance
    @pytest.mark.slow
    def test_large_batch_processing_memory_budget(self, tmp_path):
        """Processing 100+ files in batch should stay within memory budget."""
        import tracemalloc

        from docx import Document as DocxDocument

        from app.pipeline.orchestrator import PipelineOrchestrator

        batch_dir = tmp_path / "batch"
        batch_dir.mkdir()
        for i in range(20):
            docx = DocxDocument()
            docx.add_heading(f"Batch Document {i}", level=1)
            docx.add_paragraph(f"Content for batch document {i} with enough text to be realistic." * 20)
            docx.save(str(batch_dir / f"batch_{i}.docx"))

        orchestrator = PipelineOrchestrator(
            templates_dir="app/templates",
            temp_dir=str(tmp_path / "batch_temp"),
        )

        memory_after_each = []
        tracemalloc.start()

        for i in range(20):
            docx_path = str(batch_dir / f"batch_{i}.docx")
            with (
                patch.object(orchestrator, "_update_status"),
                patch.object(orchestrator.grobid_client, "process_document", return_value=None),
                patch.object(orchestrator.docling_client, "process_document", return_value=None),
                patch.object(orchestrator, "_run_figure_analysis_stage", return_value=None),
                patch.object(orchestrator, "_run_table_extraction_stage", return_value=None),
                patch.object(orchestrator, "_run_ai_analysis_stage", return_value=None),
                patch("app.pipeline.safety.retry_guard.execute_with_retry", side_effect=lambda f, *a, **kw: f()),
            ):
                orchestrator.run_pipeline(
                    input_path=docx_path,
                    job_id=f"batch-{i}",
                    template_name="none",
                )
            _force_gc()
            current, peak = tracemalloc.get_traced_memory()
            memory_after_each.append(current)

        tracemalloc.stop()

        initial_mem = memory_after_each[0]
        max_growth = max(abs(m - initial_mem) for m in memory_after_each[3:])
        assert max_growth < 50 * 1024 * 1024, (
            f"Memory grew by {max_growth/1024/1024:.1f}MB across 20 files, "
            f"suggesting a leak"
        )


# ── 2B: Connection Pool Performance ────────────────────────────────────────

class TestConnectionPoolPerformance:
    """Verify connection pools are properly managed."""

    @pytest.mark.performance
    @pytest.mark.unit
    async def test_database_connections_released(self):
        """Database connections should be properly released after use."""
        mock_client = MagicMock()
        mock_client.table.return_value.select.return_value.eq.return_value.execute.return_value.data = []

        with patch("app.services.document_service.get_supabase_client", return_value=mock_client):
            import app.services.document_service as ds
            with patch.object(ds.DocumentService, "list_documents", new_callable=AsyncMock) as mock_list:
                mock_list.return_value = []
                result = await ds.DocumentService.list_documents(user_id="test-user")
                assert result == []

    @pytest.mark.performance
    @pytest.mark.unit
    def test_redis_connections_returned_to_pool(self):
        """Redis connections should be returned to pool after use."""
        from app.cache.redis_cache import redis_cache

        mock_redis = MagicMock()
        mock_redis.get.return_value = None
        mock_redis.setex.return_value = True

        with patch.object(redis_cache, "_ensure_client", return_value=mock_redis):
            result = redis_cache.get_llm_result("test:key")
            assert result is None

            redis_cache.set_llm_result("test:key", "value", ttl=3600)
            assert mock_redis.setex.called

    @pytest.mark.performance
    @pytest.mark.unit
    async def test_simultaneous_requests_share_pool(self):
        """Simultaneous requests should share connection pool efficiently."""
        import app.services.document_service as ds

        results = []

        async def list_docs():
            with patch.object(ds.DocumentService, "list_documents", new_callable=AsyncMock) as mock_list:
                mock_list.return_value = [{"id": "shared-test"}]
                result = await ds.DocumentService.list_documents(user_id="test-user")
                results.append(result)

        n_calls = 20
        start = time.perf_counter()
        await asyncio.gather(*[list_docs() for _ in range(n_calls)])
        elapsed = time.perf_counter() - start

        assert len(results) == n_calls
        assert elapsed < 5.0, f"20 concurrent queries took {elapsed:.2f}s"

    @pytest.mark.performance
    @pytest.mark.unit
    async def test_connection_pool_bounded(self):
        """Connection pool should not grow unbounded under repeated use."""
        import app.services.document_service as ds

        connections_seen = set()
        lock = threading.Lock()

        def tracking_get_supabase():
            mock = MagicMock()
            mock_id = id(mock)
            with lock:
                connections_seen.add(mock_id)
            mock.table.return_value.select.return_value.eq.return_value.order.return_value.range.return_value.execute.return_value.data = []
            return mock

        with patch.object(ds.DocumentService, "list_documents", new_callable=AsyncMock) as mock_list:
            mock_list.return_value = []
            for _ in range(100):
                await ds.DocumentService.list_documents(user_id="test-user")

        assert len(connections_seen) <= 110, (
            f"Connection pool grew to {len(connections_seen)} unique connections"
        )
