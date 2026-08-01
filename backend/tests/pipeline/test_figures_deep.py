# SPDX-License-Identifier: MIT
# Copyright (c) 2026 ScholarForm AI

from __future__ import annotations
from unittest.mock import MagicMock, patch
import pytest
pytestmark = [pytest.mark.pipeline]


class TestFigureAnalyzer:
    def test_init_defaults(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        assert a.min_width == 300
        assert a.min_height == 300
        assert a.min_dpi == 150

    def test_init_custom(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer(min_width=400, min_height=200, min_dpi=72)
        assert a.min_width == 400
        assert a.min_height == 200
        assert a.min_dpi == 72

    def test_downsample_if_needed_file_not_found(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        with patch("os.path.exists", return_value=False):
            result = a.downsample_if_needed("/nonexistent.png")
            assert result is None

    def test_downsample_if_needed_under_threshold(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        with patch("os.path.exists", return_value=True):
            with patch("os.path.getsize", return_value=1000):
                result = a.downsample_if_needed("/small.png")
                assert result == "/small.png"

    def test_downsample_if_needed_over_threshold(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        mock_img = MagicMock()
        with patch("os.path.exists", return_value=True):
            with patch("os.path.getsize", return_value=3_000_000):
                with patch("PIL.Image.open") as mock_open:
                    mock_open.return_value.__enter__.return_value = mock_img
                    result = a.downsample_if_needed("/large.png")
                    assert "_downsampled" in result
                    mock_img.thumbnail.assert_called_once()

    def test_downsample_if_needed_exception(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        with patch("os.path.exists", return_value=True):
            with patch("os.path.getsize", return_value=3_000_000):
                with patch("PIL.Image.open") as mock_open:
                    mock_open.return_value.__enter__.side_effect = Exception("Open failed")
                    result = a.downsample_if_needed("/bad.png")
                    assert result == "/bad.png"

    def test_analyze_image_not_found(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        with patch("os.path.exists", return_value=False):
            result = a.analyze_image("/missing.png")
            assert "error" in result
            assert result["error"] == "File not found"

    def test_analyze_image_valid(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        mock_img = MagicMock()
        mock_img.size = (800, 600)
        mock_img.format = "PNG"
        mock_img.mode = "RGB"
        mock_img.info = {"dpi": (300, 300)}
        with patch("os.path.exists", return_value=True):
            with patch("PIL.Image.open") as mock_open:
                mock_open.return_value.__enter__.return_value = mock_img
                result = a.analyze_image("/good.png")
                assert result["valid"] is True
                assert result["width"] == 800
                assert result["height"] == 600
                assert result["format"] == "PNG"

    def test_analyze_image_low_resolution(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        mock_img = MagicMock()
        mock_img.size = (100, 100)
        mock_img.format = "JPEG"
        mock_img.mode = "L"
        mock_img.info = {"dpi": (72, 72)}
        with patch("os.path.exists", return_value=True):
            with patch("PIL.Image.open") as mock_open:
                mock_open.return_value.__enter__.return_value = mock_img
                result = a.analyze_image("/small.png")
                assert result["valid"] is False
                assert len(result["issues"]) > 0

    def test_analyze_image_low_dpi(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        mock_img = MagicMock()
        mock_img.size = (800, 600)
        mock_img.format = "PNG"
        mock_img.mode = "RGB"
        mock_img.info = {"dpi": (72, 72)}
        with patch("os.path.exists", return_value=True):
            with patch("PIL.Image.open") as mock_open:
                mock_open.return_value.__enter__.return_value = mock_img
                result = a.analyze_image("/lowdpi.png")
                assert "Low DPI" in str(result["issues"])

    def test_analyze_image_dpi_as_int(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        mock_img = MagicMock()
        mock_img.size = (800, 600)
        mock_img.format = "PNG"
        mock_img.mode = "RGB"
        mock_img.info = {"dpi": 200}
        with patch("os.path.exists", return_value=True):
            with patch("PIL.Image.open") as mock_open:
                mock_open.return_value.__enter__.return_value = mock_img
                result = a.analyze_image("/dpi_int.png")
                assert "200" in str(result["dpi"])

    def test_analyze_image_exception(self):
        from app.pipeline.figures.analyzer import FigureAnalyzer
        a = FigureAnalyzer()
        with patch("os.path.exists", return_value=True):
            with patch("PIL.Image.open") as mock_open:
                mock_open.return_value.__enter__.side_effect = Exception("Corrupt image")
                result = a.analyze_image("/corrupt.png")
                assert "error" in result

    def test_global_instance(self):
        from app.pipeline.figures.analyzer import figure_analyzer
        assert figure_analyzer is not None


class TestFigureCaptionMatcher:
    def test_init_defaults(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        m = CaptionMatcher()
        assert m.max_distance == 2
        assert m.enable_vision is False
        assert m.vision_client is None

    def test_init_enable_vision_unavailable(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        with (
            patch("app.pipeline.figures.caption_matcher.logger") as mock_log,
            patch("app.services.nvidia_client.get_nvidia_client", return_value=None),
        ):
            m = CaptionMatcher(enable_vision=True)
            assert m.enable_vision is True
            assert m.vision_client is None

    def test_caption_pattern_matches(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        m = CaptionMatcher()
        assert m.caption_pattern.match("Figure 1: Architecture")
        assert m.caption_pattern.match("Fig. 2: Results")
        assert m.caption_pattern.match("Figure 3-a")

    def test_caption_pattern_no_match(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        m = CaptionMatcher()
        assert not m.caption_pattern.match("Table 1: Data")
        assert not m.caption_pattern.match("Just text")

    def test_process_no_figures(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.pipeline_document import PipelineDocument
        doc = PipelineDocument(document_id="test")
        m = CaptionMatcher()
        result = m.process(doc)
        assert result is doc

    def test_process_with_matches(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.pipeline_document import PipelineDocument
        from app.models.figure import Figure
        from app.models.block import Block, BlockType
        blocks = [
            Block(block_id="b1", text="Figure 1: System Architecture", index=0, block_type=BlockType.BODY),
            Block(block_id="b2", text="Some text", index=1, block_type=BlockType.BODY),
        ]
        figures = [Figure(figure_id="fig_1", index=0, metadata={"block_index": 1})]
        doc = PipelineDocument(document_id="test", blocks=blocks, figures=figures)
        m = CaptionMatcher()
        result = m.process(doc)
        assert result.figures[0].caption_text == "Figure 1: System Architecture"
        assert result.figures[0].caption_block_id == "b1"

    def test_process_skip_headings(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.pipeline_document import PipelineDocument
        from app.models.figure import Figure
        from app.models.block import Block, BlockType
        blocks = [
            Block(block_id="b1", text="Figure 1 Analysis", index=0, block_type=BlockType.HEADING_1),
            Block(block_id="b2", text="Figure 1: Real Caption", index=1, block_type=BlockType.BODY),
        ]
        figures = [Figure(figure_id="fig_1", index=0, metadata={"block_index": 1})]
        doc = PipelineDocument(document_id="test", blocks=blocks, figures=figures)
        m = CaptionMatcher()
        result = m.process(doc)
        assert result.figures[0].caption_text == "Figure 1: Real Caption"

    def test_process_exception_handling(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        doc = MagicMock()
        doc.figures = [MagicMock()]
        doc.blocks.side_effect = Exception("Boom")
        m = CaptionMatcher()
        result = m.process(doc)
        assert result is doc

    def test_find_caption_candidates(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.block import Block, BlockType
        blocks = [
            Block(block_id="b1", text="Figure 1: Caption", index=0, block_type=BlockType.BODY),
            Block(block_id="b2", text="Introduction", index=1, block_type=BlockType.HEADING_1),
            Block(block_id="b3", text="Some paragraph", index=2, block_type=BlockType.BODY),
        ]
        m = CaptionMatcher()
        candidates = m._find_caption_candidates(blocks)
        assert 0 in candidates
        assert 1 not in candidates

    def test_match_candidates_figure_found(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.figure import Figure
        from app.models.block import Block, BlockType
        blocks = [
            Block(block_id="b1", text="Some text", index=0, block_type=BlockType.BODY),
            Block(block_id="b2", text="Figure 1: Caption", index=1, block_type=BlockType.BODY),
        ]
        figures = [Figure(figure_id="fig_1", index=0, metadata={"block_index": 0})]
        m = CaptionMatcher()
        matches = m._match_candidates(blocks, figures, [1])
        assert len(matches) == 1
        assert matches[0][0].figure_id == "fig_1"

    def test_match_candidates_skips_assigned(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.figure import Figure
        from app.models.block import Block, BlockType
        blocks = [
            Block(block_id="b1", text="Some text", index=0, block_type=BlockType.BODY),
            Block(block_id="b2", text="Figure 1: Cap", index=1, block_type=BlockType.BODY),
            Block(block_id="b3", text="Figure 2: Cap2", index=2, block_type=BlockType.BODY),
        ]
        figures = [
            Figure(figure_id="fig_1", index=0, metadata={"block_index": 0}),
            Figure(figure_id="fig_2", index=1, metadata={"block_index": 0}),
        ]
        m = CaptionMatcher()
        matches = m._match_candidates(blocks, figures, [1, 2])
        assert len(matches) == 2

    def test_match_candidates_tie_break_above(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.figure import Figure
        from app.models.block import Block, BlockType
        blocks = [
            Block(block_id="b1", text="Fig above", index=0, block_type=BlockType.BODY),
            Block(block_id="b2", text="Fig below", index=2, block_type=BlockType.BODY),
            Block(block_id="b3", text="Figure 1: Cap", index=1, block_type=BlockType.BODY),
        ]
        figures = [
            Figure(figure_id="fig_1", index=0, metadata={"block_index": 0}),
            Figure(figure_id="fig_2", index=1, metadata={"block_index": 2}),
        ]
        m = CaptionMatcher(max_distance=3)
        matches = m._match_candidates(blocks, figures, [1])
        assert len(matches) == 1

    def test_match_candidates_no_valid_block_map(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.figure import Figure
        from app.models.block import Block, BlockType
        blocks = [Block(block_id="b1", text="Figure 1: Cap", index=0, block_type=BlockType.BODY)]
        figures = [Figure(figure_id="fig_1", index=0, metadata={"block_index": 99})]
        m = CaptionMatcher()
        matches = m._match_candidates(blocks, figures, [0])
        assert len(matches) == 0

    def test_enhance_captions_with_vision_no_path(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.figure import Figure
        figs = [Figure(figure_id="f1", index=0)]
        m = CaptionMatcher()
        result = m._enhance_captions_with_vision(figs)
        assert result == 0

    def test_enhance_captions_with_vision_no_client(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.figure import Figure
        figs = [Figure(figure_id="f1", index=0, export_path="/tmp/test.png")]
        with patch("app.services.nvidia_client.get_nvidia_client", return_value=None):
            m = CaptionMatcher(enable_vision=True)
            result = m._enhance_captions_with_vision(figs)
            assert result == 0

    def test_enhance_captions_vision_exception(self):
        from app.pipeline.figures.caption_matcher import CaptionMatcher
        from app.models.figure import Figure
        figs = [Figure(figure_id="f1", index=0, export_path="/tmp/test.png")]
        m = CaptionMatcher(enable_vision=True)
        m.vision_client = MagicMock()
        m.vision_client.analyze_figure.side_effect = Exception("Vision failed")
        with patch("os.path.exists", return_value=True):
            result = m._enhance_captions_with_vision(figs)
            assert result == 0

    def test_convenience_function(self):
        from app.pipeline.figures.caption_matcher import link_figures
        from app.models.pipeline_document import PipelineDocument
        doc = PipelineDocument(document_id="test")
        result = link_figures(doc, enable_vision=False)
        assert result is doc


class TestFigureRenderer:
    def test_calculate_image_size_with_dimensions(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        fig = Figure(figure_id="f1", index=0, width=1920, height=1080)
        renderer = FigureRenderer()
        width, height = renderer.calculate_image_size(fig)
        assert width is not None

    def test_calculate_image_size_no_dimensions(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        fig = Figure(figure_id="f1", index=0)
        renderer = FigureRenderer()
        width, height = renderer.calculate_image_size(fig)
        assert width is not None
        assert height is None

    def test_calculate_image_size_too_wide(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        fig = Figure(figure_id="f1", index=0, width=10000, height=1000)
        renderer = FigureRenderer()
        width, height = renderer.calculate_image_size(fig)
        from docx.shared import Inches
        assert width <= Inches(6.5)

    def test_calculate_image_size_too_narrow(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        fig = Figure(figure_id="f1", index=0, width=50, height=50)
        renderer = FigureRenderer()
        width, height = renderer.calculate_image_size(fig)
        from docx.shared import Inches
        assert width >= Inches(2.0)

    def test_render_from_export_path(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        doc = MagicMock()
        fig = Figure(figure_id="f1", index=0, export_path="/tmp/fig.png")
        renderer = FigureRenderer()
        with patch("os.path.exists", return_value=True):
            with patch.object(renderer, "calculate_image_size", return_value=(MagicMock(), MagicMock())):
                with patch.object(renderer, "_add_caption"):
                    renderer.render(doc, fig, 1)
                    doc.add_paragraph.assert_called()

    def test_render_from_export_path_failure(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        doc = MagicMock()
        paragraph = MagicMock()
        run = MagicMock()
        paragraph.add_run.return_value = run
        doc.add_paragraph.return_value = paragraph
        run.add_picture.side_effect = Exception("Picture failed")
        fig = Figure(figure_id="f1", index=0, export_path="/tmp/bad.png")
        renderer = FigureRenderer()
        with patch("os.path.exists", return_value=True):
            with patch.object(renderer, "calculate_image_size", return_value=(MagicMock(), MagicMock())):
                with patch.object(renderer, "_add_caption"):
                    renderer.render(doc, fig, 1)
                    paragraph.add_run.assert_called()

    def test_render_from_image_data(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        doc = MagicMock()
        paragraph = MagicMock()
        run = MagicMock()
        paragraph.add_run.return_value = run
        doc.add_paragraph.return_value = paragraph
        fig = Figure(figure_id="f1", index=0, image_data=b"fake_image_bytes")
        renderer = FigureRenderer()
        with patch.object(renderer, "calculate_image_size", return_value=(MagicMock(), MagicMock())):
            with patch.object(renderer, "_add_caption"):
                renderer.render(doc, fig, 1)
                run.add_picture.assert_called()

    def test_render_from_image_data_no_height(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        doc = MagicMock()
        paragraph = MagicMock()
        run = MagicMock()
        paragraph.add_run.return_value = run
        doc.add_paragraph.return_value = paragraph
        fig = Figure(figure_id="f1", index=0, image_data=b"fake")
        renderer = FigureRenderer()
        with patch.object(renderer, "calculate_image_size", return_value=(MagicMock(), None)):
            with patch.object(renderer, "_add_caption"):
                renderer.render(doc, fig, 1)
                run.add_picture.assert_called()

    def test_render_from_image_data_failure(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        doc = MagicMock()
        paragraph = MagicMock()
        run = MagicMock()
        paragraph.add_run.return_value = run
        doc.add_paragraph.return_value = paragraph
        run.add_picture.side_effect = Exception("Bad image")
        fig = Figure(figure_id="f1", index=0, image_data=b"bad")
        renderer = FigureRenderer()
        with patch.object(renderer, "calculate_image_size", return_value=(MagicMock(), MagicMock())):
            with patch.object(renderer, "_add_caption"):
                renderer.render(doc, fig, 1)
                doc.add_paragraph.assert_called()

    def test_render_no_image(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        doc = MagicMock()
        fig = Figure(figure_id="f1", index=0)
        renderer = FigureRenderer()
        with patch.object(renderer, "calculate_image_size", return_value=(MagicMock(), MagicMock())):
            with patch.object(renderer, "_add_caption"):
                renderer.render(doc, fig, 1)
                doc.add_paragraph.assert_called()

    def test_add_caption_no_text(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        doc = MagicMock()
        fig = Figure(figure_id="f1", index=0)
        renderer = FigureRenderer()
        renderer._add_caption(doc, fig, 1)
        doc.add_paragraph.assert_not_called()

    def test_add_caption_with_text_matching_number(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        doc = MagicMock()
        fig = Figure(figure_id="f1", index=0, caption_text="Figure 1: Test caption")
        renderer = FigureRenderer()
        renderer._add_caption(doc, fig, 1)
        doc.add_paragraph.assert_called_with(style="Caption")

    def test_add_caption_with_text_not_matching(self):
        from app.pipeline.figures.renderer import FigureRenderer
        from app.models.figure import Figure
        doc = MagicMock()
        fig = Figure(figure_id="f1", index=0, caption_text="Custom description")
        renderer = FigureRenderer()
        renderer._add_caption(doc, fig, 1)
        doc.add_paragraph.assert_called_with(style="Caption")
