# SPDX-License-Identifier: MIT
# Copyright (c) 2026 ScholarForm AI

"""
Figure Renderer - Renders Figure models into python-docx documents.
"""

import logging
import os
from io import BytesIO
from typing import Optional, Tuple

from docx.shared import Inches

from app.models import Figure
from app.pipeline.safety.safe_execution import safe_function

logger = logging.getLogger(__name__)

# Page constraints (standard letter: 8.5" wide, 1" margins each side = 6.5" usable)
MAX_WIDTH = Inches(6.5)
MAX_HEIGHT = Inches(9.0)
MIN_WIDTH = Inches(2.0)
DEFAULT_WIDTH = Inches(5.0)


class FigureRenderer:
    """
    Renders Figure objects into a Word document with dynamic sizing.
    """

    @safe_function(fallback_value=None, error_message="Image sizing failed")
    def calculate_image_size(self, figure: Figure) -> Tuple[Inches, Optional[Inches]]:
        """Calculate optimal image size based on actual dimensions and page constraints."""
        if figure.width and figure.height:
            img_width_inches = Inches(figure.width / 96.0)
            img_height_inches = Inches(figure.height / 96.0)
            aspect_ratio = figure.width / figure.height

            if img_width_inches > MAX_WIDTH:
                final_width = MAX_WIDTH
                final_height = Inches(MAX_WIDTH.inches / aspect_ratio)
            elif img_width_inches < MIN_WIDTH:
                final_width = MIN_WIDTH
                final_height = Inches(MIN_WIDTH.inches / aspect_ratio)
            else:
                final_width = img_width_inches
                final_height = img_height_inches

            if final_height > MAX_HEIGHT:
                final_height = MAX_HEIGHT
                final_width = Inches(MAX_HEIGHT.inches * aspect_ratio)

            return final_width, final_height
        else:
            return DEFAULT_WIDTH, None

    @safe_function(fallback_value=None, error_message="Figure rendering failed")
    def render(self, doc, figure: Figure, number: int):
        """Render a figure with dynamic sizing based on image dimensions."""
        width, height = self.calculate_image_size(figure)

        if figure.export_path and os.path.exists(figure.export_path):
            try:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(figure.export_path, width=width, height=height)
                paragraph.alignment = 1
            except Exception as e:
                logger.warning("Failed to render figure from export_path: %s", e)
                p = doc.add_paragraph(f"[Image: {figure.export_path}]")
                p.alignment = 1
        elif figure.image_data:
            try:
                image_stream = BytesIO(figure.image_data)
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                if height:
                    run.add_picture(image_stream, width=width, height=height)
                else:
                    run.add_picture(image_stream, width=width)
                paragraph.alignment = 1
                logger.info("Rendered figure %d from image_data (%d bytes)", number, len(figure.image_data))
            except Exception as e:
                logger.warning("Failed to render figure from image_data: %s", e)
                p = doc.add_paragraph(f"[Figure {number} - Image rendering failed: {str(e)[:50]}]")
                p.alignment = 1
        else:
            p = doc.add_paragraph(f"[Figure {number} Placeholder - No image data]")
            p.alignment = 1

        self._add_caption(doc, figure, number)

    def _add_caption(self, doc, figure: Figure, number: int):
        """Add caption with bold prefix below the figure."""
        if not figure.caption_text:
            return
        caption_p = doc.add_paragraph(style="Caption")
        caption_lower = figure.caption_text.lower().strip()
        if caption_lower.startswith(f"figure {number}:"):
            run = caption_p.add_run(f"Figure {number}: ")
            run.bold = True
            rest_text = figure.caption_text[len(f"Figure {number}:"):].strip()
            caption_p.add_run(rest_text)
        else:
            run = caption_p.add_run(f"Figure {number}: ")
            run.bold = True
            caption_p.add_run(figure.caption_text)
        caption_p.alignment = 1
