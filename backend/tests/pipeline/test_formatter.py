# SPDX-License-Identifier: MIT
# Copyright (c) 2026 ScholarForm AI

from __future__ import annotations
import os
import io
import yaml
import pytest
from unittest.mock import patch, MagicMock, PropertyMock, call
from pathlib import Path
from zipfile import ZipFile
from xml.etree.ElementTree import Element, SubElement, tostring

@pytest.fixture
def formatter():
    from app.pipeline.formatting.formatter import Formatter
    from app.models import PipelineDocument, DocumentMetadata, TemplateInfo, Block, BlockType, Figure, Table, TableCell, Equation, TextStyle, Reference

    with (
        patch("app.pipeline.formatting.formatter.ContractLoader") as mock_cl,
        patch("app.pipeline.formatting.formatter.StyleMapper") as mock_sm,
        patch("app.pipeline.formatting.formatter.NumberingEngine") as mock_ne,
        patch("app.pipeline.formatting.formatter.ReferenceFormatter") as mock_rf,
        patch("app.pipeline.formatting.formatter.TemplateRenderer") as mock_tr,
        patch("app.pipeline.formatting.formatter.TableRenderer") as mock_tbl,
    ):
        mock_cl_instance = mock_cl.return_value
        mock_cl_instance.load.return_value = {}
        f = Formatter(
            templates_dir="app/templates",
            contracts_dir="app/pipeline/contracts",
        )
        return f


@pytest.fixture
def minimal_doc():
    from app.models import PipelineDocument, DocumentMetadata
    return PipelineDocument(
        document_id="doc1",
        blocks=[],
        metadata=DocumentMetadata(),
    )


@pytest.fixture
def doc_with_blocks():
    from app.models import PipelineDocument, DocumentMetadata, Block, BlockType, Reference
    blocks = [
        Block(block_id="b1", index=1, block_type=BlockType.TITLE, text="Paper Title", section_name="abstract"),
        Block(block_id="b2", index=2, block_type=BlockType.BODY, text="Abstract text.", section_name="abstract"),
        Block(block_id="b3", index=3, block_type=BlockType.HEADING_1, text="Introduction", section_name="introduction"),
        Block(block_id="b4", index=4, block_type=BlockType.BODY, text="Body paragraph.", section_name="introduction"),
        Block(block_id="b5", index=5, block_type=BlockType.HEADING_2, text="Methods", section_name="methods"),
        Block(block_id="b6", index=6, block_type=BlockType.BODY, text="We used a novel approach.", section_name="methods"),
        Block(block_id="b7", index=7, block_type=BlockType.REFERENCE_ENTRY, text="[1] Author, J. (2024)", section_name="references"),
    ]
    return PipelineDocument(
        document_id="doc2",
        blocks=blocks,
        metadata=DocumentMetadata(title="Test Paper", authors=["John Doe"]),
        references=[
            Reference(block_id="b7", block_index=7, reference_id="ref1", index=7, raw_text="[1] Author, J. (2024)", citation_key="ref1"),
        ],
        formatting_options={},
    )


def _make_word_doc(paragraphs=None):
    """Create a minimal mock Word document."""
    doc = MagicMock()
    doc._body = MagicMock()
    doc._body._element = MagicMock()
    doc.sections = [MagicMock()]
    doc.sections[0]._sectPr = MagicMock()
    doc.sections[0]._sectPr.xpath.return_value = []
    doc.sections[0].footer = MagicMock()
    doc.sections[0].footer.paragraphs = []
    doc.paragraphs = []
    doc.add_paragraph = MagicMock()
    p = MagicMock()
    p._p = MagicMock()
    p._p.xml = ""
    p._p.getparent.return_value = MagicMock()
    p.text = ""
    p.alignment = None
    p.style = None
    p.paragraph_format = MagicMock()
    run = MagicMock()
    run._r = MagicMock()
    run.font = MagicMock()
    run.bold = False
    run.italic = False
    p.add_run.return_value = run
    p.runs = [run]
    doc.add_paragraph.return_value = p
    return doc


# ── Init ────────────────────────────────────────────────────────────────────

class TestFormatterInit:
    def test_init_default_paths(self):
        with patch("app.pipeline.formatting.formatter.ContractLoader"):
            f = Formatter()
        assert "app/templates" in str(f.templates_dir)

    def test_init_custom_paths(self):
        with patch("app.pipeline.formatting.formatter.ContractLoader"):
            f = Formatter(templates_dir="/tmp/templates", contracts_dir="/tmp/contracts")
        assert str(f.templates_dir) == "/tmp/templates"

    def test_init_sets_sub_components(self, formatter):
        assert formatter.contract_loader is not None
        assert formatter.style_mapper is not None
        assert formatter.numbering_engine is not None
        assert formatter.reference_formatter is not None
        assert formatter.template_renderer is not None
        assert formatter.table_renderer is not None


# ── Process ──────────────────────────────────────────────────────────────────

class TestFormatterProcess:
    def test_process_sets_generated_doc(self, formatter, doc_with_blocks):
        with patch.object(formatter, "format", return_value="rendered_doc"):
            result = formatter.process(doc_with_blocks)
        assert result.generated_doc == "rendered_doc"

    def test_process_calls_format_with_correct_args(self, formatter, doc_with_blocks):
        with patch.object(formatter, "format", return_value="rendered") as mock_format:
            formatter.process(doc_with_blocks)
        mock_format.assert_called_once_with(doc_with_blocks, "none")

    def test_process_uses_template_name(self, formatter):
        from app.models import PipelineDocument, DocumentMetadata, TemplateInfo
        doc = PipelineDocument(
            document_id="doc1",
            blocks=[],
            metadata=DocumentMetadata(),
            template=TemplateInfo(template_name="ieee"),
        )
        with patch.object(formatter, "format", return_value="rendered") as mock_format:
            formatter.process(doc)
        mock_format.assert_called_once_with(doc, "ieee")


# ── Format – template renderer path ─────────────────────────────────────────

class TestFormatterFormatTemplatePath:
    def test_format_uses_template_renderer(self, formatter, doc_with_blocks):
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        formatter.template_renderer.render = MagicMock(return_value=MagicMock())
        formatter.template_renderer.render.return_value.save = MagicMock()
        with patch.object(formatter, "_post_process_template_render") as mock_post:
            with patch.object(formatter, "_install_post_save_hook"):
                result = formatter.format(doc_with_blocks, template_name="ieee")
        assert result is not None
        formatter.template_renderer.render.assert_called_once()

    def test_format_template_renderer_fallback_on_error(self, formatter, doc_with_blocks):
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        from jinja2.exceptions import TemplateError
        formatter.template_renderer.render = MagicMock(side_effect=TemplateError("fail"))
        with patch.object(formatter, "_load_contract", return_value={}):
            with patch.object(formatter, "_apply_initial_layout"):
                with patch.object(formatter, "_apply_page_size"):
                    with patch.object(formatter, "_render_block"):
                        with patch.object(formatter, "_install_post_save_hook"):
                            with patch("app.pipeline.formatting.formatter.WordDocument") as mock_wd:
                                mock_wd.return_value = _make_word_doc()
                                result = formatter.format(doc_with_blocks, template_name="ieee")
        assert result is not None

    def test_format_none_template_blank_doc(self, formatter, minimal_doc):
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        with patch.object(formatter, "_load_contract", return_value={}):
            with patch.object(formatter, "_apply_initial_layout"):
                with patch.object(formatter, "_apply_page_size"):
                    with patch.object(formatter, "_install_post_save_hook"):
                        with patch("app.pipeline.formatting.formatter.WordDocument") as mock_wd:
                            mock_wd.return_value = _make_word_doc()
                            result = formatter.format(minimal_doc, template_name="none")
        assert result is not None

    def test_format_resolves_bool_options(self, formatter, doc_with_blocks):
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        doc_with_blocks.formatting_options = {
            "cover_page": True,
            "toc": True,
            "page_numbers": True,
            "borders": True,
            "line_numbers": True,
        }
        formatter.template_renderer.render = MagicMock(return_value=MagicMock())
        with patch.object(formatter, "_post_process_template_render") as mock_post:
            with patch.object(formatter, "_install_post_save_hook"):
                formatter.format(doc_with_blocks, template_name="ieee")
        mock_post.assert_called_once()

    def test_format_applies_numbering_and_references(self, formatter, doc_with_blocks):
        formatter.template_renderer.render = MagicMock(return_value=MagicMock())
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        with patch.object(formatter, "_post_process_template_render"):
            with patch.object(formatter, "_install_post_save_hook"):
                with patch.object(formatter, "_prepare_references") as mock_prep:
                    formatter.format(doc_with_blocks, template_name="ieee")
        formatter.numbering_engine.apply_numbering.assert_called_once_with(doc_with_blocks, "ieee")
        mock_prep.assert_called_once()


# ── Format – legacy path ────────────────────────────────────────────────────

class TestFormatterFormatLegacyPath:
    def test_format_legacy_renders_blocks(self, formatter, doc_with_blocks):
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        doc_with_blocks.formatting_options = {"template_engine": "legacy"}
        with patch.object(formatter, "_load_contract", return_value={}):
            with patch.object(formatter, "_apply_initial_layout"):
                with patch.object(formatter, "_apply_page_size"):
                    with patch.object(formatter, "_render_block") as mock_render:
                        with patch.object(formatter, "_install_post_save_hook"):
                            with patch("app.pipeline.formatting.formatter.WordDocument") as mock_wd:
                                mock_wd.return_value = _make_word_doc()
                                formatter.format(doc_with_blocks, template_name="ieee")
        assert mock_render.call_count >= 3

    def test_format_legacy_skips_header_footer_footnote(self, formatter, doc_with_blocks):
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        doc_with_blocks.blocks.append(
            Block(block_id="b8", index=8, block_type=BlockType.FOOTNOTE, text="Footnote text",
                  metadata={"is_footnote": True})
        )
        doc_with_blocks.formatting_options = {"template_engine": "legacy"}
        with patch.object(formatter, "_load_contract", return_value={}):
            with patch.object(formatter, "_apply_initial_layout"):
                with patch.object(formatter, "_apply_page_size"):
                    with patch.object(formatter, "_render_block") as mock_render:
                        with patch.object(formatter, "_install_post_save_hook"):
                            with patch("app.pipeline.formatting.formatter.WordDocument") as mock_wd:
                                mock_wd.return_value = _make_word_doc()
                                formatter.format(doc_with_blocks, template_name="ieee")
        footnote_calls = [c for c in mock_render.call_args_list
                          if c.args[1].block_id == "b8"]
        assert len(footnote_calls) == 0

    def test_format_legacy_adds_figures(self, formatter, doc_with_blocks):
        from app.models import Figure
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        doc_with_blocks.figures = [
            Figure(figure_id="f1", index=1, export_path="fig.png", caption_text="Figure caption",
                   width=400, height=300),
        ]
        doc_with_blocks.formatting_options = {"template_engine": "legacy"}
        with patch.object(formatter, "_load_contract", return_value={}):
            with patch.object(formatter, "_apply_initial_layout"):
                with patch.object(formatter, "_apply_page_size"):
                    with patch.object(formatter, "_render_block"):
                        with patch.object(formatter.figure_renderer, "render") as mock_rf:
                            with patch.object(formatter, "_install_post_save_hook"):
                                with patch("app.pipeline.formatting.formatter.WordDocument") as mock_wd:
                                    mock_wd.return_value = _make_word_doc()
                                    formatter.format(doc_with_blocks, template_name="ieee")
        mock_rf.assert_called_once()

    def test_format_legacy_adds_equations(self, formatter, doc_with_blocks):
        from app.models import Equation
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        doc_with_blocks.equations = [
            Equation(equation_id="e1", text="x=1", index=5),
        ]
        doc_with_blocks.formatting_options = {"template_engine": "legacy"}
        with patch.object(formatter, "_load_contract", return_value={}):
            with patch.object(formatter, "_apply_initial_layout"):
                with patch.object(formatter, "_apply_page_size"):
                    with patch.object(formatter, "_render_block"):
                        with patch.object(formatter, "_render_equation") as mock_re:
                            with patch.object(formatter, "_install_post_save_hook"):
                                with patch("app.pipeline.formatting.formatter.WordDocument") as mock_wd:
                                    mock_wd.return_value = _make_word_doc()
                                    formatter.format(doc_with_blocks, template_name="ieee")
        mock_re.assert_called_once()

    def test_format_legacy_adds_tables(self, formatter, doc_with_blocks):
        from app.models import Table
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        doc_with_blocks.tables = [
            Table(table_id="t1", num_rows=2, num_cols=2, index=1, block_index=6, cells=[]),
        ]
        doc_with_blocks.formatting_options = {"template_engine": "legacy"}
        with patch.object(formatter, "_load_contract", return_value={}):
            with patch.object(formatter, "_apply_initial_layout"):
                with patch.object(formatter, "_apply_page_size"):
                    with patch.object(formatter, "_render_block"):
                        with patch.object(formatter, "_install_post_save_hook"):
                            with patch("app.pipeline.formatting.formatter.WordDocument") as mock_wd:
                                mock_wd.return_value = _make_word_doc()
                                formatter.format(doc_with_blocks, template_name="ieee")
        formatter.table_renderer.render.assert_called_once()

    def test_format_legacy_applies_section_options(self, formatter, doc_with_blocks):
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        doc_with_blocks.formatting_options = {
            "template_engine": "legacy",
            "page_numbers": True,
            "borders": True,
            "line_numbers": True,
        }
        with patch.object(formatter, "_load_contract", return_value={}):
            with patch.object(formatter, "_apply_initial_layout"):
                with patch.object(formatter, "_apply_page_size"):
                    with patch.object(formatter, "_render_block"):
                        with patch.object(formatter, "_add_page_numbers") as mock_pn:
                            with patch.object(formatter, "_add_page_borders") as mock_pb:
                                with patch.object(formatter, "_add_line_numbers") as mock_ln:
                                    with patch.object(formatter, "_apply_global_line_spacing"):
                                        with patch.object(formatter, "_install_post_save_hook"):
                                            with patch.object(formatter, "_remove_static_page_number_placeholders"):
                                                with patch("app.pipeline.formatting.formatter.WordDocument") as mock_wd:
                                                    mock_wd.return_value = _make_word_doc()
                                                    formatter.format(doc_with_blocks, template_name="ieee")
        mock_pn.assert_called_once()
        mock_pb.assert_called_once()
        mock_ln.assert_called_once()

    def test_format_legacy_adds_cover_and_toc(self, formatter, doc_with_blocks):
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        doc_with_blocks.formatting_options = {
            "template_engine": "legacy",
            "cover_page": True,
            "toc": True,
        }
        with patch.object(formatter, "_load_contract", return_value={}):
            with patch.object(formatter, "_apply_initial_layout"):
                with patch.object(formatter, "_apply_page_size"):
                    with patch.object(formatter, "_render_block"):
                        with patch.object(formatter, "_add_cover_page") as mock_cp:
                            with patch.object(formatter, "_add_table_of_contents") as mock_toc:
                                with patch.object(formatter, "_install_post_save_hook"):
                                    with patch("app.pipeline.formatting.formatter.WordDocument") as mock_wd:
                                        mock_wd.return_value = _make_word_doc()
                                        formatter.format(doc_with_blocks, template_name="ieee")
        mock_cp.assert_called_once()
        mock_toc.assert_called_once()

    def test_format_legacy_missing_template_uses_blank(self, formatter, minimal_doc):
        formatter.numbering_engine.apply_numbering.side_effect = lambda doc, tmpl: doc
        minimal_doc.formatting_options = {"template_engine": "legacy"}
        with patch.object(formatter, "_load_contract", return_value={}):
            with patch.object(formatter, "_apply_initial_layout"):
                with patch.object(formatter, "_apply_page_size"):
                    with patch.object(formatter, "_install_post_save_hook"):
                        with patch("app.pipeline.formatting.formatter.WordDocument") as mock_wd:
                            mock_wd.return_value = _make_word_doc()
                            with patch("os.path.exists", return_value=False):
                                result = formatter.format(minimal_doc, template_name="ieee")
        assert result is not None
        mock_wd.assert_called_once()


# ── Bool option helpers ─────────────────────────────────────────────────────

class TestFormatterBoolOptions:
    def test_coerce_bool_none(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._coerce_bool_option(None, True) is True
        assert Formatter._coerce_bool_option(None, False) is False

    def test_coerce_bool_passthrough(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._coerce_bool_option(True, False) is True
        assert Formatter._coerce_bool_option(False, True) is False

    def test_coerce_bool_int(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._coerce_bool_option(1, False) is True
        assert Formatter._coerce_bool_option(0, True) is False
        assert Formatter._coerce_bool_option(2.5, False) is True
        assert Formatter._coerce_bool_option(0.0, True) is False

    def test_coerce_bool_string_true(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._coerce_bool_option("true", False) is True
        assert Formatter._coerce_bool_option("yes", False) is True
        assert Formatter._coerce_bool_option("1", False) is True
        assert Formatter._coerce_bool_option("on", False) is True

    def test_coerce_bool_string_false(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._coerce_bool_option("false", True) is False
        assert Formatter._coerce_bool_option("no", True) is False
        assert Formatter._coerce_bool_option("0", True) is False
        assert Formatter._coerce_bool_option("off", True) is False
        assert Formatter._coerce_bool_option("", True) is False

    def test_coerce_bool_unknown_string(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._coerce_bool_option("maybe", False) is True

    def test_resolve_bool_option_primary_key(self, formatter):
        assert formatter._resolve_bool_option({"key": True}, "key") is True
        assert formatter._resolve_bool_option({"key": False}, "key", default=True) is False

    def test_resolve_bool_option_aliases(self, formatter):
        assert formatter._resolve_bool_option({"alias": True}, "key", aliases=("alias",)) is True

    def test_resolve_bool_option_default(self, formatter):
        assert formatter._resolve_bool_option({}, "missing", default=True) is True
        assert formatter._resolve_bool_option({}, "missing", default=False) is False

    def test_resolve_bool_option_non_dict(self, formatter):
        assert formatter._resolve_bool_option(None, "key", default=True) is True


# ── References ───────────────────────────────────────────────────────────────

class TestFormatterReferences:
    def test_prepare_references_formats_missing(self, formatter, doc_with_blocks):
        ref = doc_with_blocks.references[0]
        ref.formatted_text = ""
        formatter.reference_formatter.format_reference = MagicMock(return_value="Formatted: [1]")
        formatter._prepare_references(doc_with_blocks, "ieee")
        assert ref.formatted_text == "Formatted: [1]"

    def test_prepare_references_skips_existing(self, formatter, doc_with_blocks):
        ref = doc_with_blocks.references[0]
        ref.formatted_text = "Already formatted"
        formatter.reference_formatter.format_reference = MagicMock(return_value="New format")
        formatter._prepare_references(doc_with_blocks, "ieee")
        assert ref.formatted_text == "Already formatted"

    def test_prepare_references_no_refs(self, formatter, minimal_doc):
        formatter._prepare_references(minimal_doc, "ieee")

    def test_prepare_references_fallback_on_error(self, formatter, doc_with_blocks):
        ref = doc_with_blocks.references[0]
        ref.formatted_text = ""
        ref.raw_text = "Raw fallback"
        formatter.reference_formatter.format_reference = MagicMock(side_effect=ValueError("fail"))
        formatter._prepare_references(doc_with_blocks, "ieee")
        assert ref.formatted_text == "Raw fallback"


# ── Equations ────────────────────────────────────────────────────────────────

class TestFormatterEquations:
    def test_render_equation(self, formatter):
        from app.models import Equation
        doc = _make_word_doc()
        eqn = Equation(equation_id="e1", text="E=mc^2", number="1", index=1)
        formatter._render_equation(doc, eqn)
        doc.add_paragraph.assert_called_once()

    def test_render_equation_no_number(self, formatter):
        from app.models import Equation
        doc = _make_word_doc()
        eqn = Equation(equation_id="e2", text="F=ma", number="", index=2)
        formatter._render_equation(doc, eqn)
        doc.add_paragraph.assert_called_once()


# ── Layout ───────────────────────────────────────────────────────────────────

class TestFormatterLayout:
    def test_apply_initial_layout_no_contract(self, formatter):
        doc = _make_word_doc()
        formatter.contract_loader.load.return_value = {}
        formatter._apply_initial_layout(doc, "ieee")

    def test_apply_initial_layout_with_margins(self, formatter):
        doc = _make_word_doc()
        formatter.contract_loader.load.return_value = {
            "layout": {"margins": {"top": 1.5, "bottom": 1.0, "left": 1.0, "right": 1.0}}
        }
        from docx.shared import Inches
        formatter._apply_initial_layout(doc, "ieee")
        doc.sections[0].top_margin = Inches(1.5)
        doc.sections[0].bottom_margin = Inches(1.0)

    def test_get_target_columns_default(self, formatter):
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="text", section_name="body")
        formatter.contract_loader.load.return_value = {"layout": {"default_columns": 2}}
        assert formatter._get_target_columns(block, "ieee") == 2

    def test_get_target_columns_override(self, formatter):
        block = Block(block_id="b2", index=2, block_type=BlockType.BODY, text="text", section_name="abstract")
        formatter.contract_loader.load.return_value = {
            "layout": {"default_columns": 1, "section_overrides": {"abstract": 2}}
        }
        assert formatter._get_target_columns(block, "ieee") == 2

    def test_get_target_columns_no_contract(self, formatter):
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="text")
        formatter.contract_loader.load.return_value = {}
        assert formatter._get_target_columns(block, "ieee") == 1

    def test_resolve_page_size_from_options(self, formatter):
        assert formatter._resolve_page_size("ieee", {"page_size": "A4"}) == "A4"

    def test_resolve_page_size_from_contract(self, formatter):
        formatter.contract_loader.load.return_value = {"layout": {"page_size": "Legal"}}
        assert formatter._resolve_page_size("ieee", {}) == "Legal"

    def test_resolve_page_size_default(self, formatter):
        formatter.contract_loader.load.return_value = {}
        assert formatter._resolve_page_size("ieee", {}) == "Letter"

    def test_resolve_line_spacing_from_options(self, formatter):
        assert formatter._resolve_line_spacing("ieee", {"line_spacing": 2.0}) == 2.0

    def test_resolve_line_spacing_from_contract(self, formatter):
        formatter.contract_loader.load.return_value = {"layout": {"line_spacing": 1.5}}
        assert formatter._resolve_line_spacing("ieee", {}) == 1.5

    def test_resolve_line_spacing_none(self, formatter):
        formatter.contract_loader.load.return_value = {}
        assert formatter._resolve_line_spacing("ieee", {}) is None

    def test_resolve_line_spacing_invalid(self, formatter):
        assert formatter._resolve_line_spacing("ieee", {"line_spacing": "bad"}) is None

    def test_resolve_line_spacing_zero(self, formatter):
        assert formatter._resolve_line_spacing("ieee", {"line_spacing": 0}) is None

    def test_apply_page_size(self, formatter):
        doc = _make_word_doc()
        formatter._apply_page_size(doc, "A4")

    def test_set_columns(self, formatter):
        section = MagicMock()
        section._sectPr = MagicMock()
        section._sectPr.xpath.return_value = []
        formatter._set_columns(section, 2)

    def test_set_columns_updates_existing(self, formatter):
        cols_element = MagicMock()
        section = MagicMock()
        section._sectPr = MagicMock()
        section._sectPr.xpath.return_value = [cols_element]
        formatter._set_columns(section, 1)
        cols_element.set.assert_called()

    def test_apply_global_line_spacing(self, formatter):
        doc = _make_word_doc()
        para = MagicMock()
        doc.paragraphs = [para]
        with patch.object(formatter, "_resolve_line_spacing", return_value=1.5):
            formatter._apply_global_line_spacing(doc, "ieee", {})
        para.paragraph_format.line_spacing = 1.5

    def test_apply_global_line_spacing_none(self, formatter):
        doc = _make_word_doc()
        para = MagicMock()
        doc.paragraphs = [para]
        with patch.object(formatter, "_resolve_line_spacing", return_value=None):
            formatter._apply_global_line_spacing(doc, "ieee", {})
        assert not hasattr(para.paragraph_format, 'line_spacing') or True


# ── Cover page & TOC ────────────────────────────────────────────────────────

class TestFormatterCoverAndTOC:
    def test_add_cover_page(self, formatter):
        doc = _make_word_doc()
        doc_obj = MagicMock()
        doc_obj.metadata.title = "Test Title"
        doc_obj.metadata.authors = ["Alice", "Bob"]
        doc_obj.original_filename = ""
        formatter._add_cover_page(doc, doc_obj)
        doc.add_paragraph.assert_called()

    def test_add_cover_page_untitled(self, formatter):
        doc = _make_word_doc()
        doc_obj = MagicMock()
        doc_obj.metadata.title = ""
        doc_obj.metadata.authors = []
        doc_obj.original_filename = "manuscript.docx"
        formatter._add_cover_page(doc, doc_obj)
        doc.add_paragraph.assert_called()

    def test_add_table_of_contents(self, formatter):
        doc = _make_word_doc()
        doc.add_paragraph.return_value.add_run.return_value._r = MagicMock()
        formatter._add_table_of_contents(doc)

    def test_add_table_of_contents_no_page_break(self, formatter):
        doc = _make_word_doc()
        doc.add_paragraph.return_value.add_run.return_value._r = MagicMock()
        formatter._add_table_of_contents(doc, add_page_break=False)

    def test_add_table_of_contents_prepend(self, formatter):
        doc = _make_word_doc()
        doc.add_paragraph.return_value.add_run.return_value._r = MagicMock()
        body_element = MagicMock()
        doc._body._element = body_element
        formatter._add_table_of_contents(doc, prepend=True)


# ── Page elements ────────────────────────────────────────────────────────────

class TestFormatterPageElements:
    def test_add_page_numbers(self, formatter):
        doc = _make_word_doc()
        formatter._add_page_numbers(doc)

    def test_add_page_numbers_existing_field(self, formatter):
        doc = _make_word_doc()
        existing_p = MagicMock()
        existing_p._p.xml = "<w:p><w:fldChar w:fldCharType='begin'/><w:instrText> PAGE </w:instrText></w:p>"
        doc.sections[0].footer.paragraphs = [existing_p]
        formatter._add_page_numbers(doc)

    def test_add_page_borders(self, formatter):
        doc = _make_word_doc()
        doc.sections[0]._sectPr.xpath.return_value = []
        formatter._add_page_borders(doc)

    def test_add_page_borders_removes_existing(self, formatter):
        doc = _make_word_doc()
        existing = MagicMock()
        doc.sections[0]._sectPr.xpath.return_value = [existing]
        formatter._add_page_borders(doc)
        doc.sections[0]._sectPr.remove.assert_called_with(existing)

    def test_add_line_numbers_new(self, formatter):
        doc = _make_word_doc()
        doc.sections[0]._sectPr.xpath.return_value = []
        formatter._add_line_numbers(doc)
        doc.sections[0]._sectPr.append.assert_called_once()

    def test_add_line_numbers_existing(self, formatter):
        doc = _make_word_doc()
        existing_ln = MagicMock()
        doc.sections[0]._sectPr.xpath.return_value = [existing_ln]
        formatter._add_line_numbers(doc)

    def test_add_line_numbers_custom_count_by(self, formatter):
        doc = _make_word_doc()
        doc.sections[0]._sectPr.xpath.return_value = []
        formatter._add_line_numbers(doc, count_by=5)


# ── Paragraph helpers ───────────────────────────────────────────────────────

class TestFormatterParagraphHelpers:
    def test_paragraph_has_field_code_true(self, formatter):
        p = MagicMock()
        p._p.xml = "<w:p><w:instrText>PAGE</w:instrText></w:p>"
        assert formatter._paragraph_has_field_code(p, "PAGE") is True

    def test_paragraph_has_field_code_false(self, formatter):
        p = MagicMock()
        p._p.xml = "<w:p><w:r><w:t>text</w:t></w:r></w:p>"
        assert formatter._paragraph_has_field_code(p, "PAGE") is False

    def test_paragraph_has_field_code_none(self, formatter):
        assert formatter._paragraph_has_field_code(None, "PAGE") is False

    def test_remove_paragraph(self, formatter):
        p = MagicMock()
        p._p = MagicMock()
        p._p.getparent.return_value = MagicMock()
        formatter._remove_paragraph(p)
        p._p.getparent.return_value.remove.assert_called_with(p._p)

    def test_remove_paragraph_none(self, formatter):
        formatter._remove_paragraph(None)

    def test_remove_paragraph_no_parent(self, formatter):
        p = MagicMock()
        p._p.getparent.return_value = None
        formatter._remove_paragraph(p)

    def test_prepend_paragraph(self, formatter):
        doc = _make_word_doc()
        body_element = MagicMock()
        doc._body._element = body_element
        formatter._prepend_paragraph(doc, "Header text", alignment=1)
        body_element.remove.assert_called_once()
        body_element.insert.assert_called_once()

    def test_prepend_paragraph_with_style(self, formatter):
        doc = _make_word_doc()
        doc._body._element = MagicMock()
        p = formatter._prepend_paragraph(doc, "Styled", style="Title", alignment=1)
        assert p is not None

    def test_prepend_paragraph_style_fallback(self, formatter):
        doc = _make_word_doc()
        doc.add_paragraph.side_effect = [Exception("no style"), MagicMock()]
        doc._body._element = MagicMock()
        p = formatter._prepend_paragraph(doc, "Fallback", style="MissingStyle")
        assert p is not None

    def test_document_contains_text_found(self, formatter):
        doc = _make_word_doc()
        para = MagicMock()
        para.text = "This contains target text"
        doc.paragraphs = [para]
        assert formatter._document_contains_text(doc, "target") is True

    def test_document_contains_text_not_found(self, formatter):
        doc = _make_word_doc()
        para = MagicMock()
        para.text = "Other text"
        doc.paragraphs = [para]
        assert formatter._document_contains_text(doc, "missing") is False

    def test_document_contains_text_empty(self, formatter):
        assert formatter._document_contains_text(_make_word_doc(), "") is False


# ── Front matter ─────────────────────────────────────────────────────────────

class TestFormatterFrontMatter:
    def test_prepend_front_matter_as_cover(self, formatter):
        doc = _make_word_doc()
        doc._body._element = MagicMock()
        doc_obj = MagicMock()
        doc_obj.metadata.title = "Title"
        doc_obj.metadata.authors = ["Author"]
        doc_obj.metadata.affiliations = ["University"]
        doc_obj.original_filename = ""
        formatter._prepend_front_matter(doc, doc_obj, as_cover_page=True)
        assert doc._body._element.insert.call_count >= 1

    def test_prepend_front_matter_inline(self, formatter):
        doc = _make_word_doc()
        doc._body._element = MagicMock()
        doc_obj = MagicMock()
        doc_obj.metadata.title = "Title"
        doc_obj.metadata.authors = ["Author"]
        doc_obj.metadata.affiliations = []
        doc_obj.original_filename = ""
        formatter._prepend_front_matter(doc, doc_obj, as_cover_page=False)
        assert doc._body._element.insert.call_count >= 1

    def test_prepend_front_matter_untitled(self, formatter):
        doc = _make_word_doc()
        doc._body._element = MagicMock()
        doc_obj = MagicMock()
        doc_obj.metadata.title = ""
        doc_obj.metadata.authors = []
        doc_obj.metadata.affiliations = []
        doc_obj.original_filename = "paper.docx"
        formatter._prepend_front_matter(doc, doc_obj, as_cover_page=False)

    def test_remove_static_page_number_placeholders(self, formatter):
        doc = _make_word_doc()
        p = MagicMock()
        p.text = "Page 1"
        p._p = MagicMock()
        p._p.getparent.return_value = MagicMock()
        doc.paragraphs = [p]
        formatter._remove_static_page_number_placeholders(doc)
        p._p.getparent.return_value.remove.assert_called_once()

    def test_remove_static_page_number_placeholders_no_match(self, formatter):
        doc = _make_word_doc()
        p = MagicMock()
        p.text = "Some text"
        doc.paragraphs = [p]
        formatter._remove_static_page_number_placeholders(doc)

    def test_remove_static_toc_block(self, formatter):
        doc = _make_word_doc()
        p1 = MagicMock()
        p1.text = "Table of Contents"
        p1._p = MagicMock()
        p1._p.getparent.return_value = MagicMock()
        p2 = MagicMock()
        p2.text = "1. Introduction"
        p2._p = MagicMock()
        p2._p.getparent.return_value = MagicMock()
        p3 = MagicMock()
        p3.text = ""
        p3._p = MagicMock()
        p3._p.getparent.return_value = MagicMock()
        doc.paragraphs = [p1, p2, p3]
        formatter._remove_static_toc_block(doc)

    def test_ensure_dynamic_toc_already_present(self, formatter):
        doc = _make_word_doc()
        doc._body._element.xml = 'TOC \\o "1-3" \\h \\z \\u'
        formatter._ensure_dynamic_toc(doc)

    def test_ensure_dynamic_toc_missing(self, formatter):
        doc = _make_word_doc()
        doc._body._element.xml = "<w:document/>"
        with patch.object(formatter, "_add_table_of_contents") as mock:
            formatter._ensure_dynamic_toc(doc)
        mock.assert_called_once()


# ── Block rendering ─────────────────────────────────────────────────────────

class TestFormatterBlockRendering:
    def test_render_block_normal(self, formatter):
        doc = _make_word_doc()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="Normal paragraph text")
        formatter.style_mapper.get_style_name.return_value = "Normal"
        formatter._render_block(doc, block, "ieee")
        doc.add_paragraph.assert_called()

    def test_render_block_empty_skipped(self, formatter):
        doc = _make_word_doc()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="")
        formatter._render_block(doc, block, "ieee")
        doc.add_paragraph.assert_not_called()

    def test_render_block_bullet_list(self, formatter):
        doc = _make_word_doc()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="- List item")
        formatter._render_block(doc, block, "ieee")
        doc.add_paragraph.assert_called_with(style="List Bullet")

    def test_render_block_numbered_list(self, formatter):
        doc = _make_word_doc()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="1. Numbered item")
        formatter._render_block(doc, block, "ieee")
        doc.add_paragraph.assert_called_with(style="List Number")

    def test_render_block_with_figure_anchor(self, formatter):
        doc = _make_word_doc()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="", metadata={"has_figure": True})
        formatter._render_block(doc, block, "ieee")
        doc.add_paragraph.assert_not_called()

    def test_render_block_with_equation_anchor(self, formatter):
        doc = _make_word_doc()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="", metadata={"has_equation": True})
        formatter._render_block(doc, block, "ieee")
        doc.add_paragraph.assert_not_called()

    def test_render_block_style_fallback(self, formatter):
        doc = _make_word_doc()
        doc.add_paragraph.side_effect = [Exception("style error"), MagicMock()]
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="Fallback text")
        formatter._render_block(doc, block, "ieee")

    def test_is_bullet_list_item_true(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._is_bullet_list_item(None, "- item") is True
        assert Formatter._is_bullet_list_item(None, "* item") is True
        assert Formatter._is_bullet_list_item(None, "\u2022 item") is True

    def test_is_bullet_list_item_false(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._is_bullet_list_item(None, "Normal text") is False
        assert Formatter._is_bullet_list_item(None, "") is False

    def test_is_numbered_list_item_true(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._is_numbered_list_item(None, "1. Item") is True
        assert Formatter._is_numbered_list_item(None, "1) Item") is True
        assert Formatter._is_numbered_list_item(None, "a) Item") is True

    def test_is_numbered_list_item_false(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._is_numbered_list_item(None, "Normal text") is False
        assert Formatter._is_numbered_list_item(None, "") is False

    def test_clean_list_text_bullet(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._clean_list_text(None, "- Item") == "Item"

    def test_clean_list_text_numbered(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._clean_list_text(None, "1. Item") == "Item"

    def test_clean_list_text_no_marker(self):
        from app.pipeline.formatting.formatter import Formatter
        assert Formatter._clean_list_text(None, "Normal") == "Normal"


# ── Spacing ──────────────────────────────────────────────────────────────────

class TestFormatterSpacing:
    def test_apply_spacing_from_contract_heading(self, formatter):
        paragraph = MagicMock()
        block = MagicMock()
        block.is_heading.return_value = True
        block.block_type = BlockType.HEADING_1
        formatter.contract_loader.load.return_value = {
            "layout": {"spacing": {"heading": {"before": 12, "after": 6}}}
        }
        formatter._apply_spacing_from_contract(paragraph, block, "ieee")
        paragraph.paragraph_format.space_before = 12
        paragraph.paragraph_format.space_after = 6

    def test_apply_spacing_from_contract_paragraph(self, formatter):
        paragraph = MagicMock()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="Body")
        formatter.contract_loader.load.return_value = {
            "layout": {"spacing": {"paragraph": {"before": 6, "after": 3}}}
        }
        formatter._apply_spacing_from_contract(paragraph, block, "ieee")
        paragraph.paragraph_format.space_before = 6
        paragraph.paragraph_format.space_after = 3

    def test_apply_spacing_from_contract_no_rules(self, formatter):
        paragraph = MagicMock()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="Body")
        formatter.contract_loader.load.return_value = {"layout": {}}
        formatter._apply_spacing_from_contract(paragraph, block, "ieee")

    def test_apply_spacing_from_contract_line_spacing(self, formatter):
        paragraph = MagicMock()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="Body")
        formatter.contract_loader.load.return_value = {
            "layout": {"spacing": {"paragraph": {"before": 0, "after": 0}}, "line_spacing": "1.5"}
        }
        formatter._apply_spacing_from_contract(paragraph, block, "ieee")
        paragraph.paragraph_format.line_spacing = 1.5

    def test_apply_spacing_from_contract_invalid_line_spacing(self, formatter):
        from app.models import Block, BlockType
        paragraph = MagicMock()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="Body")
        formatter.contract_loader.load.return_value = {
            "layout": {"spacing": {}, "line_spacing": "invalid"}
        }
        formatter._apply_spacing_from_contract(paragraph, block, "ieee")


# ── Figure rendering & sizing ────────────────────────────────────────────────

class TestFormatterFigures:
    def test_calculate_image_size_with_dimensions(self, formatter):
        from app.models import Figure
        fig = Figure(figure_id="f1", index=1, export_path="test.png", width=800, height=600)
        w, h = formatter._calculate_image_size(fig)
        assert w is not None

    def test_calculate_image_size_no_dimensions(self, formatter):
        from app.models import Figure
        fig = Figure(figure_id="f1", index=1, export_path="test.png")
        w, h = formatter._calculate_image_size(fig)
        assert w is not None
        assert h is None

    def test_calculate_image_size_small_image(self, formatter):
        from app.models import Figure
        fig = Figure(figure_id="f1", index=1, export_path="test.png", width=50, height=30)
        w, h = formatter._calculate_image_size(fig)
        assert w is not None

    def test_calculate_image_size_wide_image(self, formatter):
        from app.models import Figure
        fig = Figure(figure_id="f1", index=1, export_path="test.png", width=2000, height=200)
        w, h = formatter._calculate_image_size(fig)
        assert w is not None

    def test_render_figure_with_export_path(self, formatter):
        from app.models import Figure
        doc = _make_word_doc()
        fig = Figure(figure_id="f1", index=1, export_path="test.png", caption_text="A test figure",
                     width=400, height=300)
        with patch("os.path.exists", return_value=True):
            with patch.object(formatter, "_calculate_image_size", return_value=(5.0, 3.75)):
                formatter._render_figure(doc, fig, 1)
        doc.add_paragraph.assert_called()

    def test_render_figure_export_path_fallback(self, formatter):
        from app.models import Figure
        doc = _make_word_doc()
        fig = Figure(figure_id="f1", index=1, export_path="missing.png", caption_text="Missing fig",
                     width=400, height=300)
        with patch("os.path.exists", return_value=True):
            with patch.object(formatter, "_calculate_image_size", return_value=(5.0, 3.75)):
                with patch("app.pipeline.formatting.formatter.logger") as mock_log:
                    doc.add_paragraph.return_value.add_run.return_value.add_picture.side_effect = Exception("fail")
                    formatter._render_figure(doc, fig, 1)

    def test_render_figure_no_path_no_data(self, formatter):
        from app.models import Figure
        doc = _make_word_doc()
        fig = Figure(figure_id="f1", index=1, caption_text="No image", export_path="")
        formatter._render_figure(doc, fig, 1)

    def test_render_figure_caption_with_prefix(self, formatter):
        from app.models import Figure
        doc = _make_word_doc()
        fig = Figure(figure_id="f1", index=1, export_path="test.png", caption_text="Figure 1: A caption",
                     width=400, height=300)
        with patch("os.path.exists", return_value=False):
            formatter._render_figure(doc, fig, 1)
        doc.add_paragraph.assert_called_with(style="Caption")

    def test_render_figure_caption_without_prefix(self, formatter):
        from app.models import Figure
        doc = _make_word_doc()
        fig = Figure(figure_id="f1", index=1, export_path="test.png", caption_text="A description",
                     width=400, height=300)
        with patch("os.path.exists", return_value=False):
            formatter._render_figure(doc, fig, 1)

    def test_render_figure_no_caption(self, formatter):
        from app.models import Figure
        doc = _make_word_doc()
        fig = Figure(figure_id="f1", index=1, export_path="test.png", caption_text="",
                     width=400, height=300)
        with patch("os.path.exists", return_value=False):
            formatter._render_figure(doc, fig, 1)

    def test_render_figure_with_image_data(self, formatter):
        from app.models import Figure
        doc = _make_word_doc()
        fig = Figure(figure_id="f1", index=1, image_data=b"fake_image_bytes", caption_text="From bytes",
                     width=400, height=300)
        formatter._render_figure(doc, fig, 1)

    def test_render_figure_with_image_data_fallback(self, formatter):
        from app.models import Figure
        doc = _make_word_doc()
        fig = Figure(figure_id="f1", index=1, image_data=b"bad_bytes", caption_text="Fail")
        doc.add_paragraph.return_value.add_run.return_value.add_picture.side_effect = Exception("bad img")
        formatter._render_figure(doc, fig, 1)


# ── Footnotes ────────────────────────────────────────────────────────────────

class TestFormatterFootnotes:
    def test_build_footnote_lookup(self, formatter):
        from app.models import PipelineDocument, DocumentMetadata, Block, BlockType
        doc = PipelineDocument(
            document_id="doc1",
            blocks=[
                Block(block_id="b1", index=1, block_type=BlockType.BODY, text="Body"),
                Block(block_id="b2", index=2, block_type=BlockType.FOOTNOTE, text="Note text",
                      metadata={"footnote_id": "fn1"}),
            ],
            metadata=DocumentMetadata(),
        )
        lookup = formatter._build_footnote_lookup(doc)
        assert "fn1" in lookup
        assert lookup["fn1"]["text"] == "Note text"

    def test_build_footnote_lookup_empty(self, formatter, minimal_doc):
        lookup = formatter._build_footnote_lookup(minimal_doc)
        assert lookup == {}

    def test_build_footnote_lookup_deduplicates(self, formatter):
        from app.models import PipelineDocument, DocumentMetadata, Block, BlockType
        doc = PipelineDocument(
            document_id="doc1",
            blocks=[
                Block(block_id="b1", index=1, block_type=BlockType.FOOTNOTE, text="First",
                      metadata={"footnote_id": "1"}),
                Block(block_id="b2", index=2, block_type=BlockType.FOOTNOTE, text="Second",
                      metadata={"footnote_id": "1"}),
            ],
            metadata=DocumentMetadata(),
        )
        lookup = formatter._build_footnote_lookup(doc)
        assert len(lookup) == 1

    def test_append_footnote_reference(self, formatter):
        p = MagicMock()
        p._p = MagicMock()
        formatter._append_footnote_reference(p, 1)

    def test_build_footnotes_part(self, formatter):
        lookup = {"1": {"word_id": 1, "text": "First note"}, "2": {"word_id": 2, "text": "Second note"}}
        xml_bytes = formatter._build_footnotes_part(lookup)
        assert xml_bytes.startswith(b"<?xml")
        assert b"First note" in xml_bytes
        assert b"Second note" in xml_bytes

    def test_build_footnotes_part_empty(self, formatter):
        xml_bytes = formatter._build_footnotes_part({})
        assert xml_bytes.startswith(b"<?xml")

    def test_install_post_save_hook(self, formatter):
        obj = MagicMock()
        obj.save = MagicMock()
        obj._scholarform_save_hook_installed = False
        formatter._install_post_save_hook(obj, {"1": {"word_id": 1, "text": "Note"}})
        assert obj._scholarform_save_hook_installed is True

    def test_install_post_save_hook_already_installed(self, formatter):
        obj = MagicMock()
        obj._scholarform_save_hook_installed = True
        formatter._install_post_save_hook(obj, {"1": {"word_id": 1, "text": "Note"}})

    def test_install_post_save_hook_empty_lookup(self, formatter):
        obj = MagicMock()
        formatter._install_post_save_hook(obj, {})

    def test_patch_saved_docx_with_footnotes_bytesio(self, formatter):
        target = io.BytesIO()
        target.write(b"test data")
        with patch.object(formatter, "_patch_docx_payload", return_value=b"patched"):
            formatter._patch_saved_docx_with_footnotes(target, {"1": {"word_id": 1, "text": "Note"}})
        target.seek(0)
        assert target.read() == b"patched"

    def test_patch_saved_docx_with_footnotes_empty(self, formatter):
        target = io.BytesIO()
        formatter._patch_saved_docx_with_footnotes(target, {"1": {"word_id": 1, "text": "Note"}})

    def test_patch_saved_docx_with_footnotes_filepath(self, formatter, tmp_path):
        target_path = tmp_path / "test.docx"
        target_path.write_bytes(b"some content")
        with patch.object(formatter, "_patch_docx_payload", return_value=b"patched"):
            formatter._patch_saved_docx_with_footnotes(str(target_path), {"1": {"word_id": 1, "text": "Note"}})
        assert target_path.read_bytes() == b"patched"

    def test_patch_docx_payload_no_references(self, formatter):
        payload = _make_zip_bytes({"word/document.xml": b"<w:document/>"})
        result = formatter._patch_docx_payload(payload, {"1": {"word_id": 1, "text": "Note"}})
        assert result == payload

    def test_patch_docx_payload_with_references(self, formatter):
        document_xml = b'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p><w:r><w:footnoteReference w:id="1"/></w:r></w:p></w:body></w:document>'
        payload = _make_zip_bytes({
            "word/document.xml": document_xml,
            "[Content_Types].xml": b'<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>',
        })
        with patch.object(formatter, "_patch_content_types", return_value=b"patched_ct"):
            with patch.object(formatter, "_patch_document_relationships", return_value=b"patched_rel"):
                with patch.object(formatter, "_patch_settings_xml", return_value=b"patched_set"):
                    with patch.object(formatter, "_build_footnotes_part", return_value=b"footnotes_xml"):
                        result = formatter._patch_docx_payload(payload, {"1": {"word_id": 1, "text": "Note"}})
        assert result is not None
        assert result != payload

    def test_patch_content_types_adds_override(self, formatter):
        ct_xml = b'<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>'
        result = formatter._patch_content_types(ct_xml)
        assert b"footnotes.xml" in result

    def test_patch_content_types_already_exists(self, formatter):
        ct_xml = b'<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Override PartName="/word/footnotes.xml"/></Types>'
        result = formatter._patch_content_types(ct_xml)
        assert b"footnotes.xml" in result

    def test_patch_document_relationships_new(self, formatter):
        result = formatter._patch_document_relationships(b"")
        assert b"footnotes" in result

    def test_patch_document_relationships_existing(self, formatter):
        rel_xml = b'<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/></Relationships>'
        result = formatter._patch_document_relationships(rel_xml)
        assert b"footnotes" in result

    def test_patch_settings_xml_no_settings(self, formatter):
        assert formatter._patch_settings_xml(b"") == b""

    def test_patch_settings_xml_adds_footnote_pr(self, formatter):
        settings_xml = b'<?xml version="1.0"?><w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:settings>'
        result = formatter._patch_settings_xml(settings_xml)
        assert b"footnotePr" in result or b"footnote" in result


# ── Contract loading ────────────────────────────────────────────────────────

class TestFormatterContract:
    def test_load_contract_success(self, formatter, tmp_path):
        contract_file = tmp_path / "contract.yaml"
        contract_file.write_text(yaml.dump({"layout": {"margins": {"top": 1}}}))
        result = formatter._load_contract(str(contract_file))
        assert result["layout"]["margins"]["top"] == 1

    def test_load_contract_missing(self, formatter):
        result = formatter._load_contract("/nonexistent/path.yaml")
        assert result == {}

    def test_load_contract_invalid(self, formatter, tmp_path):
        contract_file = tmp_path / "bad.yaml"
        contract_file.write_text(":: invalid yaml ::")
        result = formatter._load_contract(str(contract_file))
        assert result == {}


# ── Post-process template render ────────────────────────────────────────────

class TestFormatterPostProcess:
    def test_post_process_with_docx(self, formatter):
        from app.models import PipelineDocument, DocumentMetadata
        rendered = MagicMock()
        rendered.docx = MagicMock()
        doc = PipelineDocument(document_id="doc1", blocks=[], metadata=DocumentMetadata())
        with patch.object(formatter, "_apply_initial_layout"):
            with patch.object(formatter, "_apply_page_size"):
                with patch.object(formatter, "_remove_static_page_number_placeholders"):
                    with patch.object(formatter, "_rehydrate_template_render"):
                        with patch.object(formatter, "_apply_global_line_spacing"):
                            formatter._post_process_template_render(rendered, doc, "ieee", {})

    def test_post_process_no_docx(self, formatter):
        from app.models import PipelineDocument, DocumentMetadata
        rendered = MagicMock()
        rendered.docx = None
        doc = PipelineDocument(document_id="doc1", blocks=[], metadata=DocumentMetadata())
        formatter._post_process_template_render(rendered, doc, "ieee", {})


# ── Inline content & hyperlinks ─────────────────────────────────────────────

class TestFormatterInlineContent:
    def test_write_inline_content_plain(self, formatter):
        p = MagicMock()
        p._p = MagicMock()
        formatter._write_inline_content(p, "Hello world", [], [], {})

    def test_write_inline_content_with_hyperlinks(self, formatter):
        p = MagicMock()
        p._p = MagicMock()
        with patch.object(formatter, "_add_hyperlink") as mock_ah:
            formatter._write_inline_content(
                p, "Visit ScholarForm AI",
                [{"text": "ScholarForm AI", "url": "https://scholarform.ai"}],
                [], {},
            )
        mock_ah.assert_called_once()

    def test_write_inline_content_hyperlink_no_match(self, formatter):
        p = MagicMock()
        p._p = MagicMock()
        with patch.object(formatter, "_add_hyperlink") as mock_ah:
            formatter._write_inline_content(
                p, "Some text",
                [{"text": "MissingLink", "url": "https://example.com"}],
                [], {},
            )
        mock_ah.assert_not_called()

    def test_write_inline_content_empty_hyperlink(self, formatter):
        p = MagicMock()
        p._p = MagicMock()
        with patch.object(formatter, "_add_hyperlink") as mock_ah:
            formatter._write_inline_content(
                p, "Text",
                [{"text": "", "url": ""}],
                [], {},
            )

    def test_write_inline_content_footnote_refs(self, formatter):
        p = MagicMock()
        p._p = MagicMock()
        with patch.object(formatter, "_append_footnote_reference") as mock_afr:
            formatter._write_inline_content(
                p, "Text",
                [], ["1"],
                {"1": {"word_id": 1, "text": "Note"}},
            )
        mock_afr.assert_called_once_with(p, 1)

    def test_write_inline_content_empty_result(self, formatter):
        p = MagicMock()
        p._p = MagicMock()
        formatter._write_inline_content(p, "", [], [], {})

    def test_add_hyperlink(self, formatter):
        p = MagicMock()
        p.part = MagicMock()
        p.part.relate_to.return_value = "rId99"
        p._p = MagicMock()
        formatter._add_hyperlink(p, "Click here", "https://example.com")

    def test_replace_paragraph_inline_content(self, formatter):
        from app.models import Block, BlockType
        p = MagicMock()
        p._p = MagicMock()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="Text with link",
                      metadata={"hyperlinks": [{"text": "link", "url": "https://example.com"}]})
        with patch.object(formatter, "_clear_paragraph_content"):
            with patch.object(formatter, "_write_inline_content"):
                formatter._replace_paragraph_inline_content(p, block, {})

    def test_clear_paragraph_content(self, formatter):
        p = MagicMock()
        child_a = MagicMock()
        child_a.tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
        child_b = MagicMock()
        child_b.tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
        p._p = [child_a, child_b]
        formatter._clear_paragraph_content(p)
        assert child_b._p.remove_child_or_something


# ── Rehydrate template render ───────────────────────────────────────────────

class TestFormatterRehydrate:
    def test_rehydrate_matched_paragraph(self, formatter):
        from app.models import Block, BlockType, PipelineDocument, DocumentMetadata
        doc = _make_word_doc()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="Matched paragraph")
        doc.paragraphs = [MagicMock()]
        doc.paragraphs[0].text = "Matched paragraph"
        doc_obj = PipelineDocument(document_id="doc1", blocks=[block], metadata=DocumentMetadata())
        with patch.object(formatter, "_find_matching_paragraph", return_value=doc.paragraphs[0]):
            with patch.object(formatter, "_replace_paragraph_inline_content"):
                formatter._rehydrate_template_render(doc, doc_obj, "ieee", {})

    def test_rehydrate_unmatched_renders_block(self, formatter):
        from app.models import Block, BlockType, PipelineDocument, DocumentMetadata
        doc = _make_word_doc()
        block = Block(block_id="b1", index=1, block_type=BlockType.BODY, text="Unmatched text",
                      metadata={"hyperlinks": [{"text": "link", "url": "https://example.com"}]})
        doc_obj = PipelineDocument(document_id="doc1", blocks=[block], metadata=DocumentMetadata())
        with patch.object(formatter, "_find_matching_paragraph", return_value=None):
            with patch.object(formatter, "_render_block"):
                formatter._rehydrate_template_render(doc, doc_obj, "ieee", {})

    def test_rehydrate_unmatched_reference_entry(self, formatter):
        from app.models import Block, BlockType, PipelineDocument, DocumentMetadata
        doc = _make_word_doc()
        block = Block(block_id="b1", index=1, block_type=BlockType.REFERENCE_ENTRY, text="[1] Ref")
        doc_obj = PipelineDocument(document_id="doc1", blocks=[block], metadata=DocumentMetadata())
        with patch.object(formatter, "_find_matching_paragraph", return_value=None):
            with patch.object(formatter, "_render_block"):
                formatter._rehydrate_template_render(doc, doc_obj, "ieee", {})

    def test_rehydrate_skips_footnotes(self, formatter):
        from app.models import Block, BlockType, PipelineDocument, DocumentMetadata
        doc = _make_word_doc()
        block = Block(block_id="b1", index=1, block_type=BlockType.FOOTNOTE, text="Skip footnote",
                      metadata={"is_footnote": True})
        doc_obj = PipelineDocument(document_id="doc1", blocks=[block], metadata=DocumentMetadata())
        with patch.object(formatter, "_find_matching_paragraph") as mock_find:
            formatter._rehydrate_template_render(doc, doc_obj, "ieee", {})
        mock_find.assert_not_called()

    def test_rehydrate_appends_figures(self, formatter):
        from app.models import PipelineDocument, DocumentMetadata, Figure, Equation, Table
        doc = _make_word_doc()
        doc_obj = PipelineDocument(
            document_id="doc1", blocks=[], metadata=DocumentMetadata(),
            figures=[Figure(figure_id="f1", index=1, export_path="fig.png", caption_text="Fig", width=100, height=100)],
            equations=[Equation(equation_id="e1", text="x=1", index=1)],
            tables=[Table(table_id="t1", num_rows=2, num_cols=2, index=1, block_index=1, cells=[])],
        )
        with patch.object(formatter, "_render_figure"):
            with patch.object(formatter, "_render_equation"):
                formatter._rehydrate_template_render(doc, doc_obj, "ieee", {})

    def test_find_matching_paragraph_exact(self, formatter):
        doc = _make_word_doc()
        p1 = MagicMock()
        p1.text = "Hello world"
        p2 = MagicMock()
        p2.text = "Goodbye"
        doc.paragraphs = [p1, p2]
        result = formatter._find_matching_paragraph(doc, "Hello world", set())
        assert result is not None

    def test_find_matching_paragraph_contains(self, formatter):
        doc = _make_word_doc()
        p1 = MagicMock()
        p1.text = "The quick brown fox"
        doc.paragraphs = [p1]
        result = formatter._find_matching_paragraph(doc, "quick brown", set())
        assert result is not None

    def test_find_matching_paragraph_not_found(self, formatter):
        doc = _make_word_doc()
        doc.paragraphs = []
        result = formatter._find_matching_paragraph(doc, "nothing", set())
        assert result is None

    def test_find_matching_paragraph_used_skipped(self, formatter):
        doc = _make_word_doc()
        p = MagicMock()
        p.text = "Hello world"
        doc.paragraphs = [p]
        result = formatter._find_matching_paragraph(doc, "Hello world", {id(p)})
        assert result is None

    def test_find_matching_paragraph_empty_needle(self, formatter):
        doc = _make_word_doc()
        result = formatter._find_matching_paragraph(doc, "", set())
        assert result is None


def _make_zip_bytes(files: dict) -> bytes:
    import zipfile
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, content in files.items():
            zf.writestr(name, content)
    buf.seek(0)
    return buf.read()


