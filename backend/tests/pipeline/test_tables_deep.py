# SPDX-License-Identifier: MIT
# Copyright (c) 2026 ScholarForm AI

"""
Deep test suite for tables pipeline: caption matcher, extractor, renderer.
Targets uncovered branches and edge cases to push coverage >85% / >90%.
"""

from __future__ import annotations
from unittest.mock import patch, MagicMock, call
import pytest
from app.pipeline.tables.caption_matcher import TableCaptionMatcher
from app.pipeline.tables.extractor import TableExtractor
from app.pipeline.tables.renderer import TableRenderer
from app.models import PipelineDocument, Block, BlockType, Table, TableCell


# ===================================================================
# CAPTION MATCHER — Deep Coverage Tests
# ===================================================================

class TestTableCaptionMatcherDeep:
    """Covers reference boundaries, fallback search, tiebreaker, exceptions."""

    @pytest.fixture
    def matcher(self):
        return TableCaptionMatcher()

    # -- _find_references_start_index ---------------------------------

    def test_find_refs_via_block_type(self, matcher):
        blocks = [
            Block(block_id="r1", index=5, text="References", block_type=BlockType.REFERENCES_HEADING),
            Block(block_id="b1", index=6, text="[1] Some ref.", block_type=BlockType.REFERENCE_ENTRY),
        ]
        assert matcher._find_references_start_index(blocks) == 5

    def test_find_refs_via_keyword_heading(self, matcher):
        blocks = [
            Block(block_id="r1", index=3, text="References", block_type=BlockType.HEADING_1),
        ]
        assert matcher._find_references_start_index(blocks) == 3

    def test_find_refs_via_bibliography_keyword(self, matcher):
        blocks = [
            Block(block_id="b1", index=7, text="Bibliography", block_type=BlockType.HEADING_2),
        ]
        assert matcher._find_references_start_index(blocks) == 7

    def test_find_refs_via_works_cited_keyword(self, matcher):
        blocks = [
            Block(block_id="b1", index=9, text="Works Cited", block_type=BlockType.HEADING_3),
        ]
        assert matcher._find_references_start_index(blocks) == 9

    def test_keyword_non_heading_does_not_match(self, matcher):
        blocks = [
            Block(block_id="b1", index=4, text="references", block_type=BlockType.BODY),
        ]
        assert matcher._find_references_start_index(blocks) is None

    def test_no_references_returns_none(self, matcher):
        blocks = [
            Block(block_id="b1", index=0, text="Introduction", block_type=BlockType.HEADING_1),
        ]
        assert matcher._find_references_start_index(blocks) is None

    # -- References heading boundary (lines 60–68, 113–114) -----------

    def test_ref_boundary_skips_candidates_after_refs(self, matcher):
        """Blocks at/after references heading are excluded from caption matching."""
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="c1", index=0, text="Table 1. Before refs.", block_type=BlockType.BODY),
            Block(block_id="body", index=1, text="Body.", block_type=BlockType.BODY),
            Block(block_id="ref", index=2, text="References", block_type=BlockType.REFERENCES_HEADING),
            Block(block_id="c2", index=3, text="Table 1. After refs.", block_type=BlockType.BODY),
            Block(block_id="b1", index=4, text="Body after.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=4),
        ])
        result = matcher.process(doc)
        # ref_start_idx = 2.  Window for table at pos 4: lower=max(0,4-2)=2, upper=min(4,4+1)=4.
        # c1 (idx=0) at pos 0 is outside the window.  c2 (idx=3) at pos 2 is inside
        # the window but idx=3 >= ref_start_idx=2 → skipped.
        assert result.tables[0].caption_text is None

    def test_ref_boundary_no_candidates_before_refs(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="r1", index=0, text="References", block_type=BlockType.HEADING_1),
            Block(block_id="b1", index=1, text="Table 1. Missed.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
        ])
        result = matcher.process(doc)
        # ref_start_idx = 0; candidate b1 has index 1 >= 0 → skipped → no caption.
        assert result.tables[0].caption_text is None

    # -- Table index not in list_index_map (lines 81–91) --------------

    def test_table_idx_not_in_list_index_map_fallback(self, matcher):
        """When table.block_index does not match any block index, linear search fallback."""
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="b1", index=0, text="Table 1. Caption.", block_type=BlockType.BODY),
            Block(block_id="b2", index=2, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
        ])
        result = matcher.process(doc)
        # block_index=1 not in {0, 2} → fallback. table_list_pos found at `b2` (index >= 1 gives pos 1).
        # search window [0, 2], caption at pos 0 is in bounds → matched.
        assert result.tables[0].caption_text == "Table 1. Caption."

    def test_table_idx_fallback_to_end_of_list(self, matcher):
        """When table_idx > all block indices, table_list_pos = len(blocks) - 1."""
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="b1", index=10, text="Table 1. Caption.", block_type=BlockType.BODY),
            Block(block_id="b2", index=20, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=99),
        ])
        result = matcher.process(doc)
        # linear search: no index >= 99 → table_list_pos = len(blocks)-1 = 1
        # window: [max(0,1-2)=0, min(1,1+1)=1]
        # pos 0: b1 idx=10, candidate.index(10) <= prev_table_idx(-1)? No. >= next(inf)? No. matches caption → best.
        assert result.tables[0].caption_text == "Table 1. Caption."

    # -- Previous / next table boundary (lines 96–98, 115–116) --------

    def test_prev_table_boundary_skips_caption(self, matcher):
        """Caption before prev_table_idx is excluded."""
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="c1", index=0, text="Table 1. Caption one.", block_type=BlockType.BODY),
            Block(block_id="x1", index=1, text="Table data.", block_type=BlockType.BODY),
            Block(block_id="c2", index=2, text="Table 2. Caption two.", block_type=BlockType.BODY),
            Block(block_id="x2", index=3, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
            Table(table_id="t2", num_rows=1, num_cols=1, index=1, block_index=3),
        ])
        result = matcher.process(doc)
        # t1 gets c1.  t2: prev=1, next=inf. c2 idx=2 > 1 → ok. matches "Table 2" → correct.
        assert result.tables[1].caption_text == "Table 2. Caption two."

    def test_next_table_boundary_skips_caption(self, matcher):
        """Caption at/after next_table_idx is excluded from current table."""
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="c1", index=0, text="Table 1. First cap.", block_type=BlockType.BODY),
            Block(block_id="x1", index=1, text="Data.", block_type=BlockType.BODY),
            Block(block_id="c2", index=2, text="Table 2. Second cap.", block_type=BlockType.BODY),
            Block(block_id="x2", index=3, text="Data.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
            Table(table_id="t2", num_rows=1, num_cols=1, index=1, block_index=3),
        ])
        result = matcher.process(doc)
        assert result.tables[0].caption_text == "Table 1. First cap."
        assert result.tables[1].caption_text == "Table 2. Second cap."

    # -- Equal distance tiebreaker (lines 122–125) --------------------

    def test_equal_distance_prefers_caption_above(self, matcher):
        """When two candidates have equal distance, prefer the one above the table."""
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="above", index=0, text="Table 1. Above caption.", block_type=BlockType.BODY),
            Block(block_id="body1", index=1, text="Body.", block_type=BlockType.BODY),
            Block(block_id="below", index=2, text="Table 1. Below caption.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
        ])
        result = matcher.process(doc)
        # Both at distance 1, but "above" has pos 0 < table_list_pos 1 → preferred.
        assert result.tables[0].caption_text == "Table 1. Above caption."
        assert result.tables[0].caption_block_id == "above"

    # -- Already assigned block IDs (lines 109–110) -------------------

    def test_already_assigned_block_id_skipped(self, matcher):
        """A block already assigned as caption to one table cannot be reused."""
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="shared", index=0, text="Table 1. Shared cap.", block_type=BlockType.BODY),
            Block(block_id="body1", index=1, text="Body one.", block_type=BlockType.BODY),
            Block(block_id="body2", index=2, text="Body two.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
            Table(table_id="t2", num_rows=1, num_cols=1, index=1, block_index=2),
        ])
        result = matcher.process(doc)
        # t1 claims "shared".  t2 tries: "shared" block_id in assigned → skipped.
        assert result.tables[0].caption_text == "Table 1. Shared cap."
        assert result.tables[1].caption_text is None

    # -- Heading / REFERENCES_HEADING cannot be captions (111–112) ----

    def test_heading_skipped_as_caption(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="h1", index=0, text="Table 1. Introduction", block_type=BlockType.HEADING_1),
            Block(block_id="b1", index=1, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
        ])
        result = matcher.process(doc)
        assert result.tables[0].caption_text is None

    def test_references_heading_skipped_as_caption(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="rh", index=0, text="Table 1. References", block_type=BlockType.REFERENCES_HEADING),
            Block(block_id="b1", index=1, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
        ])
        result = matcher.process(doc)
        assert result.tables[0].caption_text is None

    # -- Exception handling per-table (lines 138–139) -----------------

    def test_per_table_exception_caught(self, matcher):
        """Exception inside per-table loop is caught; processing continues."""
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="b1", index=0, text="Table 1. Test", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="bad", num_rows=1, num_cols=1, index=0, block_index=0),
        ])
        with patch.object(matcher, "caption_regex") as mock_regex:
            mock_regex.match.side_effect = Exception("boom")
            result = matcher.process(doc)
        # Exception caught, document still returned
        assert result is doc

    # -- Top-level exception (lines 140–147) --------------------------

    def test_top_level_exception_adds_error_stage(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="b1", index=0, text="body", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=0),
        ])
        # Force exception at the top level (e.g. in log call after loop)
        with patch.object(matcher, "_find_references_start_index",
                          side_effect=Exception("catastrophe")):
            result = matcher.process(doc)
        stages = {s.stage_name: s.status for s in result.processing_history}
        assert stages.get("table_caption_matching") == "error"

    # -- Missing caption sets metadata (lines 135–137) ----------------

    def test_missing_caption_sets_status(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="b1", index=0, text="No table caption here.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=0),
        ])
        result = matcher.process(doc)
        assert result.tables[0].metadata.get("caption_status") == "Missing"

    def test_existing_caption_text_preserved_when_no_match(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="b1", index=0, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=0,
                  caption_text="Already set."),
        ])
        result = matcher.process(doc)
        # caption_text is already set → line 136: not table.caption_text → skip
        assert "caption_status" not in result.tables[0].metadata

    # -- Caption regex: various formats --------------------------------

    def test_roman_numeral_caption(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="b1", index=0, text="Table I: Results", block_type=BlockType.BODY),
            Block(block_id="b2", index=1, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
        ])
        result = matcher.process(doc)
        assert result.tables[0].caption_text == "Table I: Results"

    def test_letter_caption(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="b1", index=0, text="Table A: Appendix", block_type=BlockType.BODY),
            Block(block_id="b2", index=1, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
        ])
        result = matcher.process(doc)
        assert result.tables[0].caption_text == "Table A: Appendix"

    def test_dotted_number_caption(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="b1", index=0, text="Table 1.1: Sub results", block_type=BlockType.BODY),
            Block(block_id="b2", index=1, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
        ])
        result = matcher.process(doc)
        assert result.tables[0].caption_text == "Table 1.1: Sub results"

    # -- Search window boundaries -------------------------------------

    def test_search_window_above_respected(self, matcher):
        matcher.search_window_above = 1  # narrow window
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="c1", index=0, text="Table 1. Far above.", block_type=BlockType.BODY),
            Block(block_id="x1", index=1, text="Mid body.", block_type=BlockType.BODY),
            Block(block_id="x2", index=2, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=2),
        ])
        result = matcher.process(doc)
        # window: [max(0,2-1)=1, min(2,2+1)=2] → c1 at pos 0 not included
        assert result.tables[0].caption_text is None

    def test_search_window_below_respected(self, matcher):
        matcher.search_window_below = 0  # no search below
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="x1", index=0, text="Body.", block_type=BlockType.BODY),
            Block(block_id="c1", index=1, text="Table 1. Below.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=0),
        ])
        result = matcher.process(doc)
        # window: [0, 0] → c1 at pos 1 not included
        assert result.tables[0].caption_text is None

    def test_caption_below_table_in_window(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="x1", index=0, text="Body.", block_type=BlockType.BODY),
            Block(block_id="c1", index=1, text="Table 1. Below cap.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=0),
        ])
        result = matcher.process(doc)
        # window: [0, 1] → c1 at pos 1 included
        assert result.tables[0].caption_text == "Table 1. Below cap."

    # -- No tables / no blocks early return ---------------------------

    def test_no_tables_early_return(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="b1", index=0, text="Some text.", block_type=BlockType.BODY),
        ])
        result = matcher.process(doc)
        assert result is doc

    def test_no_blocks_early_return(self, matcher):
        doc = PipelineDocument(document_id="t", tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=0),
        ])
        result = matcher.process(doc)
        assert result is doc

    # -- Processing stage metadata ------------------------------------

    def test_success_stage_message_contains_match_count(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="c1", index=0, text="Table 1. Match.", block_type=BlockType.BODY),
            Block(block_id="x1", index=1, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
        ])
        result = matcher.process(doc)
        stage = next(s for s in result.processing_history if s.stage_name == "table_caption_matching")
        assert "1" in stage.message

    def test_caption_block_type_changed(self, matcher):
        doc = PipelineDocument(document_id="t", blocks=[
            Block(block_id="c1", index=0, text="Table 1. Types.", block_type=BlockType.BODY),
            Block(block_id="x1", index=1, text="Body.", block_type=BlockType.BODY),
        ], tables=[
            Table(table_id="t1", num_rows=1, num_cols=1, index=0, block_index=1),
        ])
        result = matcher.process(doc)
        caption_block = next(b for b in result.blocks if b.block_id == "c1")
        assert caption_block.block_type == BlockType.TABLE_CAPTION
        assert caption_block.metadata.get("linked_table_id") == "t1"
        assert caption_block.metadata.get("classification_method") == "deterministic_table_caption_rule"

    def test_custom_search_window_constructor(self):
        matcher = TableCaptionMatcher(search_window_above=5, search_window_below=3)
        assert matcher.search_window_above == 5
        assert matcher.search_window_below == 3

    # -- Tiebreaker / edge cases ---------------------------------------

    def test_equal_distance_tiebreaker_prefers_above(self, matcher):
        blocks = [
            Block(block_id="c1", index=0, text="Table 1. Above.", block_type=BlockType.BODY),
            Block(block_id="x1", index=1, text="Body.", block_type=BlockType.BODY),
            Block(block_id="c2", index=2, text="Table 1. Below.", block_type=BlockType.BODY),
        ]
        tables = [Table(table_id="t1", num_rows=2, num_cols=2, index=0, block_index=1)]
        doc = PipelineDocument(document_id="t", blocks=blocks, tables=tables)
        result = matcher.process(doc)
        assert result.tables[0].caption_text == "Table 1. Above."

    def test_missing_caption_sets_metadata(self, matcher):
        blocks = [Block(block_id="b1", index=0, text="Body.", block_type=BlockType.BODY)]
        tables = [Table(table_id="t1", num_rows=2, num_cols=2, index=0, block_index=0)]
        doc = PipelineDocument(document_id="t", blocks=blocks, tables=tables)
        result = matcher.process(doc)
        assert result.tables[0].metadata.get("caption_status") == "Missing"

    def test_convenience_function(self):
        from app.pipeline.tables.caption_matcher import match_table_captions
        doc = PipelineDocument(document_id="t", blocks=[], tables=[])
        result = match_table_captions(doc)
        assert result is doc


# ===================================================================
# EXTRACTOR — Deep Coverage Tests
# ===================================================================

class TestTableExtractorDeep:
    """Covers _normalize_cell_text, deep XML edge cases, nested tables."""

    @pytest.fixture
    def extractor(self):
        return TableExtractor()

    # -- _normalize_cell_text -----------------------------------------

    def test_normalize_cell_text_none(self, extractor):
        assert extractor._normalize_cell_text(None) == ""

    def test_normalize_cell_text_whitespace_only(self, extractor):
        assert extractor._normalize_cell_text("   \t  ") == ""

    def test_normalize_cell_text_preserves_internal(self, extractor):
        assert extractor._normalize_cell_text("  hello\nworld  ") == "hello\nworld"

    # -- Deep XML edge cases ------------------------------------------

    def test_deep_xml_iteration_exception(self, extractor):
        """_tc.iter() raises → _extract_deep_xml_text returns '' gracefully."""
        cell = MagicMock()
        cell._tc = MagicMock()
        cell._tc.iter.side_effect = RuntimeError("iter crash")
        result = extractor._extract_deep_xml_text(cell)
        assert result == ""

    def test_deep_xml_no_text_nodes(self, extractor):
        """iter() yields nodes but none match the ending '}t' tag."""
        cell = MagicMock()
        cell._tc = MagicMock()
        fake_node = MagicMock()
        fake_node.tag = "w:r"
        fake_node.text = "ignored"
        cell._tc.iter.return_value = [fake_node]
        result = extractor._extract_deep_xml_text(cell)
        assert result == ""

    def test_deep_xml_multiple_text_parts(self, extractor):
        cell = MagicMock()
        cell._tc = MagicMock()
        n1 = MagicMock(tag="}t", text="Hello ")
        n2 = MagicMock(tag="}t", text="World")
        cell._tc.iter.return_value = [n1, n2]
        result = extractor._extract_deep_xml_text(cell)
        assert result == "Hello World"

    # -- Outer extraction safeguards ----------------------------------

    def test_extract_deep_xml_exception_in_outer_loop(self, extractor):
        """When cell.text is empty and _extract_deep_xml_text raises,
        the outer try-except in extract() catches and logs."""
        cell = MagicMock()
        cell.text = ""
        cell._tc = MagicMock()
        cell._tc.iter.side_effect = RuntimeError("boom")
        run = MagicMock()
        run.bold = False
        cell.paragraphs = [MagicMock(runs=[run])]
        cell._element = MagicMock()
        cell._element.findall.return_value = []
        cell._parent = MagicMock()
        row = MagicMock()
        row.cells = [cell]
        tbl = MagicMock()
        tbl.rows = [row]
        result = extractor.extract(tbl, "tbl_deep_fail", 0, 0)
        assert result.data[0][0] == ""

    def test_extract_with_no_cell_text_and_deep_also_empty(self, extractor):
        cell = MagicMock()
        cell.text = ""
        cell._tc = MagicMock()
        cell._tc.iter.return_value = []  # no text nodes
        run = MagicMock()
        run.bold = False
        cell.paragraphs = [MagicMock(runs=[run])]
        cell._element = MagicMock()
        cell._element.findall.return_value = []
        cell._parent = MagicMock()
        row = MagicMock()
        row.cells = [cell]
        tbl = MagicMock()
        tbl.rows = [row]
        result = extractor.extract(tbl, "tbl_deep_empty", 0, 0)
        assert result.data[0][0] == ""

    # -- Uneven cells in original row (padding) -----------------------

    def test_padding_preserves_empty_cells(self, extractor):
        """Empty cells remain empty; no data corruption."""
        from unittest.mock import MagicMock
        cell = MagicMock(text="A", bold=False)
        cell._tc = MagicMock()
        cell._tc.iter.return_value = []
        cell._element = MagicMock()
        cell._element.findall.return_value = []
        cell._parent = MagicMock()
        empty_cell = MagicMock(text="", bold=False)
        empty_cell._tc = MagicMock()
        empty_cell._tc.iter.return_value = []
        empty_cell._element = MagicMock()
        empty_cell._element.findall.return_value = []
        empty_cell._parent = MagicMock()
        row = MagicMock()
        row.cells = [cell, empty_cell]
        tbl = MagicMock()
        tbl.rows = [row]
        result = extractor.extract(tbl, "tbl_pad", 0, 0)
        assert result.data[0] == ["A", ""]
        assert result.num_cols == 2

    # -- Header detection keywords ------------------------------------

    def test_header_keywords_unit(self, extractor):
        """'Unit' contains 'unit' keyword → header detected."""
        cell = MagicMock(text="Unit", bold=False)
        cell._tc = MagicMock()
        cell._tc.iter.return_value = []
        cell._element = MagicMock()
        cell._element.findall.return_value = []
        cell._parent = MagicMock()
        row = MagicMock()
        row.cells = [cell]
        tbl = MagicMock()
        tbl.rows = [row]
        result = extractor.extract(tbl, "tbl_unit", 0, 0)
        assert result.has_header is True

    def test_header_keywords_qty(self, extractor):
        cell = MagicMock(text="Qty", bold=False)
        cell._tc = MagicMock()
        cell._tc.iter.return_value = []
        cell._element = MagicMock()
        cell._element.findall.return_value = []
        cell._parent = MagicMock()
        row = MagicMock()
        row.cells = [cell]
        tbl = MagicMock()
        tbl.rows = [row]
        result = extractor.extract(tbl, "tbl_qty", 0, 0)
        assert result.has_header is True

    def test_header_keywords_amount(self, extractor):
        cell = MagicMock(text="Amount", bold=False)
        cell._tc = MagicMock()
        cell._tc.iter.return_value = []
        cell._element = MagicMock()
        cell._element.findall.return_value = []
        cell._parent = MagicMock()
        row = MagicMock()
        row.cells = [cell]
        tbl = MagicMock()
        tbl.rows = [row]
        result = extractor.extract(tbl, "tbl_amt", 0, 0)
        assert result.has_header is True

    # -- Nested tables -------------------------------------------------

    def test_multiple_nested_tables_in_one_cell(self, extractor):
        """A cell with 2 nested tables both get extracted."""
        cell = MagicMock(text="outer", bold=False)
        cell._tc = MagicMock()
        cell._tc.iter.return_value = []
        cell._parent = MagicMock()
        cell._parent._parent = MagicMock()
        cell._parent._parent._parent = MagicMock()

        inner1_tbl = MagicMock()
        inner1_row = MagicMock()
        inner1_cell = MagicMock(text="n1", bold=False)
        inner1_cell._tc = MagicMock()
        inner1_cell._tc.iter.return_value = []
        inner1_cell._element = MagicMock()
        inner1_cell._element.findall.return_value = []
        inner1_cell._parent = MagicMock()
        inner1_row.cells = [inner1_cell]
        inner1_tbl.rows = [inner1_row]

        inner2_tbl = MagicMock()
        inner2_row = MagicMock()
        inner2_cell = MagicMock(text="n2", bold=False)
        inner2_cell._tc = MagicMock()
        inner2_cell._tc.iter.return_value = []
        inner2_cell._element = MagicMock()
        inner2_cell._element.findall.return_value = []
        inner2_cell._parent = MagicMock()
        inner2_row.cells = [inner2_cell]
        inner2_tbl.rows = [inner2_row]

        xml1 = MagicMock()
        xml1._inner_mock = inner1_tbl
        xml2 = MagicMock()
        xml2._inner_mock = inner2_tbl

        def findall_side_effect(ns):
            return [xml1, xml2] if ns.endswith("}tbl") else []

        cell._element = MagicMock()
        cell._element.findall.side_effect = findall_side_effect

        from docx.oxml.ns import qn

        def patched_table_constructor(tbl_xml, parent):
            return tbl_xml._inner_mock

        row = MagicMock()
        row.cells = [cell]
        tbl = MagicMock()
        tbl.rows = [row]

        with patch("docx.table.Table", side_effect=patched_table_constructor):
            result = extractor.extract(tbl, "tbl_multi_nest", 0, 0)
        nested = result.cells[0].metadata.get("nested_tables", [])
        assert len(nested) == 2
        assert nested[0].data[0][0] == "n1"
        assert nested[1].data[0][0] == "n2"

    def test_nested_table_unexpected_structure_does_not_crash(self, extractor):
        """Even if docx structure is malformed, extractor does not crash."""
        cell = MagicMock(text="outer", bold=False)
        cell._tc = MagicMock()
        cell._tc.iter.side_effect = AttributeError("missing _tc structure")
        cell._parent = MagicMock()
        cell._parent._parent = MagicMock()
        cell._parent._parent._parent = MagicMock()

        cell._element = MagicMock()
        cell._element.findall.return_value = []

        row = MagicMock()
        row.cells = [cell]
        tbl = MagicMock()
        tbl.rows = [row]

        result = extractor.extract(tbl, "tbl_malformed", 0, 0)
        assert result.data[0][0] == "outer"


# ===================================================================
# RENDERER — Deep Coverage Tests
# ===================================================================

class TestTableRendererDeep:
    """Covers caption formatting, style fallback, nested rendering, edge cases."""

    @pytest.fixture
    def renderer(self):
        return TableRenderer()

    @pytest.fixture
    def simple_table(self):
        return Table(
            table_id="tbl_001",
            num_rows=2, num_cols=2,
            index=0, block_index=0,
            cells=[
                TableCell(row=0, col=0, text="A1", bold=True),
                TableCell(row=0, col=1, text="B1", bold=True),
                TableCell(row=1, col=0, text="A2"),
                TableCell(row=1, col=1, text="B2"),
            ],
            data=[["A1", "B1"], ["A2", "B2"]],
            rows=[["A1", "B1"], ["A2", "B2"]],
            has_header=True,
            has_header_row=True,
            header_rows=1,
        )

    # -- No-op cases ---------------------------------------------------

    def test_render_none_table(self, renderer):
        doc = MagicMock()
        renderer.render(doc, None)
        doc.add_table.assert_not_called()
        doc.add_paragraph.assert_not_called()

    def test_render_zero_rows(self, renderer):
        doc = MagicMock()
        tbl = Table(table_id="e", num_rows=0, num_cols=0, index=0, block_index=0)
        renderer.render(doc, tbl)
        doc.add_table.assert_not_called()

    def test_render_zero_cols(self, renderer):
        doc = MagicMock()
        tbl = Table(table_id="e", num_rows=1, num_cols=0, index=0, block_index=0,
                     data=[[]], rows=[[]])
        renderer.render(doc, tbl)
        doc.add_table.assert_not_called()

    # -- Caption formatting -------------------------------------------

    def test_caption_with_exact_table_prefix(self, renderer, simple_table):
        """When caption already starts with 'Table X: ', it is split correctly."""
        doc = MagicMock()
        doc.styles = {"Table Grid": MagicMock()}
        doc.add_table.return_value = MagicMock()
        doc.add_paragraph.return_value = MagicMock()
        simple_table.caption_text = "Table 1: Experimental results"
        renderer.render(doc, simple_table)
        para = doc.add_paragraph.return_value
        calls = para.add_run.call_args_list
        assert any(c[0][0] == "Table 1: " for c in calls)
        assert any("Experimental results" in str(c[0]) for c in calls)

    def test_caption_with_different_number_in_text(self, renderer, simple_table):
        """Caption says 'Table 2: …' but actual number differs."""
        doc = MagicMock()
        doc.styles = {"Table Grid": MagicMock()}
        doc.add_table.return_value = MagicMock()
        doc.add_paragraph.return_value = MagicMock()
        simple_table.caption_text = "Table 2: Results"
        renderer.render(doc, simple_table)
        para = doc.add_paragraph.return_value
        calls = para.add_run.call_args_list
        # table.index = 0, number not passed → table_num = 0+1 = 1
        # caption starts with "Table 2:" which is != "Table 1:" → falls to else branch
        assert any(c[0][0] == "Table 1: " for c in calls)

    def test_caption_without_table_prefix(self, renderer, simple_table):
        """Caption that doesn't start with 'Table X:' at all."""
        doc = MagicMock()
        doc.styles = {"Table Grid": MagicMock()}
        doc.add_table.return_value = MagicMock()
        doc.add_paragraph.return_value = MagicMock()
        simple_table.caption_text = "Performance metrics"
        renderer.render(doc, simple_table)
        para = doc.add_paragraph.return_value
        calls = para.add_run.call_args_list
        assert any(c[0][0] == "Table 1: " for c in calls)

    def test_caption_center_alignment(self, renderer, simple_table):
        doc = MagicMock()
        doc.styles = {"Table Grid": MagicMock()}
        doc.add_table.return_value = MagicMock()
        para = MagicMock()
        doc.add_paragraph.return_value = para
        simple_table.caption_text = "Some caption"
        renderer.render(doc, simple_table)
        assert para.alignment is not None

    # -- Style application edge cases ---------------------------------

    def test_style_not_present(self, renderer, simple_table):
        doc = MagicMock()
        doc.styles = {}  # 'Table Grid' not present
        doc.add_table.return_value = MagicMock()
        doc.add_paragraph.return_value = MagicMock()
        renderer.render(doc, simple_table)
        word_table = doc.add_table.return_value
        assert not hasattr(word_table, "style") or word_table.style != "Table Grid"

    def test_style_application_exception(self, renderer, simple_table):
        doc = MagicMock()
        doc.styles = {"Table Grid": MagicMock()}
        word_table = MagicMock()
        doc.add_table.return_value = word_table
        doc.add_paragraph.return_value = MagicMock()
        # Setting style raises
        def raise_on_style_set(name, val):
            raise Exception("style error")
        type(word_table).style = property(fget=lambda s: "mock", fset=raise_on_style_set)
        # Should not crash
        renderer.render(doc, simple_table)
        # Still should have created the table
        doc.add_table.assert_called_once()

    # -- Cell population edge cases -----------------------------------

    def test_cell_out_of_bounds_handled(self, renderer, simple_table):
        """Cell row/col beyond word_table dimensions is skipped."""
        doc = MagicMock()
        doc.styles = {"Table Grid": MagicMock()}
        word_tbl = MagicMock()
        row0 = MagicMock()
        row0.cells = [MagicMock(), MagicMock()]
        word_tbl.rows = [row0]  # only 1 row
        doc.add_table.return_value = word_tbl
        doc.add_paragraph.return_value = MagicMock()
        simple_table.cells = [
            TableCell(row=0, col=0, text="ok"),
            TableCell(row=5, col=5, text="out"),
        ]
        renderer.render(doc, simple_table)
        # Should not raise; out-of-bounds cell should be skipped

    def test_cell_with_empty_text(self, renderer, simple_table):
        doc = MagicMock()
        doc.styles = {"Table Grid": MagicMock()}
        word_tbl = MagicMock()
        row0 = MagicMock()
        c00 = MagicMock()
        row0.cells = [c00]
        word_tbl.rows = [row0]
        doc.add_table.return_value = word_tbl
        doc.add_paragraph.return_value = MagicMock()
        simple_table.cells = [TableCell(row=0, col=0, text="")]
        renderer.render(doc, simple_table)
        assert c00.text == ""

    # -- Nested table rendering edge cases ---------------------------

    def test_nested_table_rendering_failure_logged(self, renderer, simple_table):
        """When a nested table render fails, fallback text is added."""
        doc = MagicMock()
        doc.styles = {"Table Grid": MagicMock()}
        word_tbl = MagicMock()
        row0 = MagicMock()
        c00 = MagicMock()
        row0.cells = [c00]
        word_tbl.rows = [row0]
        doc.add_table.return_value = word_tbl
        doc.add_paragraph.return_value = MagicMock()

        failing_nested = Table(table_id="bad_nest", num_rows=1, num_cols=1,
                               index=1, block_index=1,
                               data=[["x"]], cells=[TableCell(row=0, col=0, text="x")])
        simple_table.cells = [TableCell(row=0, col=0, text="parent",
                                        metadata={"nested_tables": [failing_nested]})]

        # Make render raise on nested call by passing None as doc
        original_render = renderer.render

        def render_side_effect(d, tbl, number=None):
            if d is c00:
                raise Exception("nested fail")
            return original_render(d, tbl, number=number)

        with patch.object(renderer, "render", side_effect=render_side_effect):
            # Should not raise; nested failure triggers fallback paragraph
            renderer.render(doc, simple_table)
        # Exception was caught and fallback add_paragraph on cell was called
        c00.add_paragraph.assert_called_once()

    # -- Full rendering smoke test -------------------------------------

    def test_smoke_render_simple_table_does_not_crash(self, renderer, simple_table):
        doc = MagicMock()
        doc.styles = {"Table Grid": MagicMock()}
        word_tbl = MagicMock()
        row0 = MagicMock()
        row0.cells = [MagicMock(), MagicMock()]
        row1 = MagicMock()
        row1.cells = [MagicMock(), MagicMock()]
        word_tbl.rows = [row0, row1]
        doc.add_table.return_value = word_tbl
        doc.add_paragraph.return_value = MagicMock()
        renderer.render(doc, simple_table)
        doc.add_table.assert_called_once_with(rows=2, cols=2)

    def test_outer_exception_caught_and_raised(self, renderer, simple_table):
        """Trigger the outer try/except at renderer.py:81-83."""
        doc = MagicMock()
        doc.styles = {"Table Grid": MagicMock()}
        doc.add_paragraph.side_effect = RuntimeError("caption failure")
        simple_table.caption_text = "Table 1: Test"
        with pytest.raises(RuntimeError, match="caption failure"):
            renderer.render(doc, simple_table)
