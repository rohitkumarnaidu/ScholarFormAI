# SPDX-License-Identifier: MIT
# Copyright (c) 2026 ScholarForm AI

from __future__ import annotations
from unittest.mock import patch, MagicMock, PropertyMock
import pytest
pytestmark = [pytest.mark.pipeline]


# ══════════════════════════════════════════════════════════════════════════════
# APA 7 Fallback Formatter (apafallback.py)
# ══════════════════════════════════════════════════════════════════════════════

class TestAPA7Formatter:
    """Coverage gap: apafallback.py was 0%"""

    def test_init_default(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f.hanging_indent is True

    def test_init_no_hanging(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        assert f.hanging_indent is False

    def test_format_intext_citation_parenthetical(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f.format_intext_citation(["Smith, J."], year=2020)
        assert result == "(Smith, 2020)"

    def test_format_intext_citation_narrative(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f.format_intext_citation(["Smith, J."], year=2020, narrative=True)
        assert result == "Smith (2020)"

    def test_format_intext_citation_with_page(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f.format_intext_citation(["Smith, J."], year=2020, page="42")
        assert result == "(Smith, 2020, p. 42)"

    def test_format_intext_citation_narrative_with_page(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f.format_intext_citation(["Smith, J."], year=2020, page="42", narrative=True)
        assert result == "Smith (2020, p. 42)"

    def test_format_intext_citation_no_year(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f.format_intext_citation(["Smith, J."])
        assert result == "(Smith, n.d.)"

    def test_format_intext_authors_empty(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_intext_authors([]) == "Unknown"

    def test_format_intext_authors_one(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_intext_authors(["Smith, J."]) == "Smith"

    def test_format_intext_authors_two(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_intext_authors(["Smith, J.", "Doe, A."])
        assert result == "Smith and Doe"

    def test_format_intext_authors_three_plus(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_intext_authors(["Smith, J.", "Doe, A.", "Lee, K."])
        assert result == "Smith et al."

    def test_extract_surname_empty(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._extract_surname("") == "Unknown"

    def test_extract_surname_comma_format(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._extract_surname("Smith, J.") == "Smith"

    def test_extract_surname_prefix(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._extract_surname("van der Waals, J.") == "van der Waals"

    def test_extract_surname_single(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._extract_surname("Smith") == "Smith"

    def test_extract_surname_reversed(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._extract_surname("John Smith") == "Smith"

    def test_extract_surname_last_as_given_with_prefix(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._extract_surname("Ludwig van Beethoven")
        assert result == "van Beethoven"

    def test_format_reference_entry_journal(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="A Study",
            journal="J. Science", volume="10", issue="2", pages="100-110",
            doi="10.1234/abc"
        )
        assert "Smith, J." in result
        assert "(2020)." in result
        assert "A Study." in result
        assert "J. Science" in result
        assert "https://doi.org/10.1234/abc" in result

    def test_format_reference_entry_journal_no_doi(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="A Study",
            journal="J. Science", volume="10", issue="2", pages="100-110"
        )
        assert "doi.org" not in result

    def test_format_reference_entry_book(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="My Book",
            publisher="Academic Press", reference_type="book", doi="10.1234/book"
        )
        assert "My Book" in result
        assert "Academic Press" in result
        assert "https://doi.org/10.1234/book" in result

    def test_format_reference_entry_book_with_edition(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="My Book",
            publisher="Academic Press", reference_type="book", edition="3rd"
        )
        assert "3rd ed." in result

    def test_format_reference_entry_book_chapter(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="My Chapter",
            book_title="Big Book", pages="50-70", publisher="Press",
            reference_type="book_chapter"
        )
        assert "My Chapter" in result
        assert "Big Book" in result
        assert "pp. 50-70" in result

    def test_format_reference_entry_book_chapter_no_pages(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="My Chapter",
            book_title="Big Book", reference_type="book_chapter"
        )
        assert "Big Book" in result

    def test_format_reference_entry_conference_paper(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="Paper",
            conference="ICML 2020", reference_type="conference_paper",
            doi="10.1234/conf"
        )
        assert "Paper" in result
        assert "ICML 2020" in result

    def test_format_reference_entry_thesis(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="My Thesis",
            publisher="MIT", reference_type="thesis"
        )
        assert "My Thesis" in result
        assert "Doctoral dissertation" in result
        assert "MIT" in result

    def test_format_reference_entry_thesis_no_publisher(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="My Thesis",
            reference_type="thesis"
        )
        assert "Doctoral dissertation" in result
        assert "Unknown Institution" not in result

    def test_format_reference_entry_web_page(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="My Page",
            publisher="Website Co.", url="https://example.com",
            reference_type="web_page"
        )
        assert "My Page" in result
        assert "Website Co." in result
        assert "https://example.com" in result

    def test_format_reference_entry_web_page_no_url(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="My Page",
            reference_type="web_page"
        )
        assert "Website" in result

    def test_format_reference_entry_default(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="Other",
            journal="Some Journal", reference_type="other"
        )
        assert "Some Journal" in result

    def test_format_reference_entry_default_no_venue(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="Other",
            reference_type="other"
        )
        assert "Other" in result
        assert "(2020)." in result

    def test_format_authors_empty(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_authors([]) == "Unknown"

    def test_format_authors_one(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_authors(["Smith, J."]) == "Smith, J."

    def test_format_authors_two(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_authors(["Smith, J.", "Doe, A."])
        assert result == "Smith, J., & Doe, A."

    def test_format_authors_three(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_authors(["Smith, J.", "Doe, A.", "Lee, K."])
        assert "&" in result

    def test_format_authors_twenty_one_plus(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        authors = [f"Author{i}, A." for i in range(21)]
        result = f._format_authors(authors)
        assert "..." in result
        assert "Author20, A." in result

    def test_format_single_author_comma_given(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_single_author("Smith, John") == "Smith, J."

    def test_format_single_author_comma_initial(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_single_author("Smith, J.") == "Smith, J."

    def test_format_single_author_single_word(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_single_author("Unknown") == "Unknown"

    def test_format_single_author_reversed(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_single_author("John Smith") == "Smith, J."

    def test_format_single_author_empty(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_single_author("") == "Unknown"

    def test_format_single_author_comma_no_given(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_single_author("Smith,") == "Smith"

    def test_format_title_empty(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_title("") == ""

    def test_format_title_adds_period(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_title("My Title")
        assert result == "My Title."

    def test_format_title_preserves_period(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_title("My Title.") == "My Title."

    def test_format_journal_article_full(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_journal_article("J. Sci.", "10", "2", "100-110")
        assert "*J. Sci.*" in result
        assert "*10*" in result
        assert "(2)" in result
        assert "100-110" in result

    def test_format_journal_article_no_pages(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_journal_article("J. Sci.", "10", "2", None)
        assert "100-110" not in result

    def test_format_journal_article_no_issue(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_journal_article("J. Sci.", "10", None, "100-110")
        assert "(2)" not in result

    def test_format_journal_article_no_journal(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_journal_article(None, "10", "2", "100-110")
        assert result == "*10*(2) 100-110."

    def test_format_journal_article_empty(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_journal_article(None, None, None, None) == ""

    def test_format_doi_empty(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        assert f._format_doi(None) == ""
        assert f._format_doi("") == ""

    def test_format_doi_https(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_doi("https://doi.org/10.1234/abc")
        assert result == "https://doi.org/10.1234/abc"

    def test_format_doi_plain(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        result = f._format_doi("10.1234/abc")
        assert result == "https://doi.org/10.1234/abc"

    def test_sort_references(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        refs = [
            {"authors": ["Zeta, B."], "title": "Z"},
            {"authors": ["Alpha, A."], "title": "A"},
        ]
        sorted_refs = f.sort_references(refs)
        assert sorted_refs[0]["authors"][0] == "Alpha, A."

    def test_sort_references_no_authors(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter()
        refs = [{"title": "Z"}, {"title": "A"}]
        sorted_refs = f.sort_references(refs)
        assert len(sorted_refs) == 2


# ══════════════════════════════════════════════════════════════════════════════
# Vancouver Fallback Formatter (vancouver_fallback.py)
# ══════════════════════════════════════════════════════════════════════════════

class TestVancouverFormatter:
    """Coverage gap: vancouver_fallback.py was 0%"""

    def test_init(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f._ref_counter == 0

    def test_format_intext_citation_empty(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f.format_intext_citation([]) == ""

    def test_format_intext_citation_single(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f.format_intext_citation([1]) == "[1]"

    def test_format_intext_citation_multiple(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f.format_intext_citation([1, 3, 5]) == "[1,3,5]"

    def test_format_intext_citation_consecutive(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f.format_intext_citation([1, 2, 3]) == "[1-3]"

    def test_format_intext_citation_mixed(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f.format_intext_citation([1, 2, 5, 6, 7]) == "[1-2,5-7]"

    def test_format_intext_citation_unsorted(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f.format_intext_citation([3, 1, 2]) == "[1-3]"

    def test_format_reference_entry_journal_full(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J", "Doe A"], title="A Study",
            journal="J Sci", year=2020, volume="10", issue="2",
            pages="100-110", doi="10.1234/abc"
        )
        assert "Smith J, Doe A" in result
        assert "A Study." in result
        assert "J Sci." in result
        assert "2020" in result
        assert "10(2)" in result
        assert ":100-110." in result
        assert "doi:" in result

    def test_format_reference_entry_journal_minimal(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="A Study"
        )
        assert "A Study." in result

    def test_format_reference_entry_book(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="My Book",
            publisher="Academic Press", year=2020,
            reference_type="book"
        )
        assert "My Book." in result
        assert "Academic Press." in result
        assert "2020" in result

    def test_format_reference_entry_book_with_edition(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="My Book",
            publisher="Academic Press", year=2020,
            edition="3rd", reference_type="book"
        )
        assert "3rd ed." in result

    def test_format_reference_entry_book_chapter(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="My Chapter",
            book_title="Big Book", year=2020,
            publisher="Press", pages="50-70",
            reference_type="book_chapter"
        )
        assert "My Chapter." in result
        assert "Big Book" in result
        assert "2020" in result
        assert "p. 50-70" in result

    def test_format_reference_entry_book_chapter_no_book_title(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="My Chapter",
            pages="50-70", reference_type="book_chapter"
        )
        assert "p. 50-70" in result

    def test_format_reference_entry_thesis(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="My Thesis",
            year=2020, publisher="MIT",
            reference_type="thesis"
        )
        assert "My Thesis." in result
        assert "[Dissertation]" in result
        assert "MIT." in result

    def test_format_reference_entry_thesis_no_publisher(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="My Thesis",
            year=2020, reference_type="thesis"
        )
        assert "Unknown Institution" in result

    def test_format_reference_entry_conference_paper(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="My Paper",
            conference="ICML 2020", year=2020,
            reference_type="conference_paper"
        )
        assert "My Paper." in result
        assert "ICML 2020" in result

    def test_format_reference_entry_conference_paper_no_conf(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="My Paper",
            year=2020, reference_type="conference_paper"
        )
        assert "My Paper." in result

    def test_format_reference_entry_default_type(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="Other", year=2020,
            reference_type="technical_report"
        )
        assert "Other." in result

    def test_format_reference_entry_doi_http(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="A Study",
            doi="https://doi.org/10.1234/abc",
            reference_type="journal_article"
        )
        assert "doi: abc." in result

    def test_format_authors_empty(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f._format_authors([]) == "Unknown"

    def test_format_authors_up_to_six(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        authors = ["A", "B", "C", "D", "E", "F"]
        result = f._format_authors(authors)
        assert "et al." not in result
        assert result == "A, B, C, D, E, F"

    def test_format_authors_seven_plus(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        authors = ["A", "B", "C", "D", "E", "F", "G", "H"]
        result = f._format_authors(authors)
        assert "et al." in result

    def test_format_doi_empty(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f._format_doi(None) == ""
        assert f._format_doi("") == ""

    def test_format_doi_plain(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f._format_doi("10.1234/abc") == "doi: 10.1234/abc."

    def test_format_doi_http(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f._format_doi("https://doi.org/10.1234/abc") == "doi: abc."

    def test_format_doi_http_no_slash(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f._format_doi("https://example.com/")
        assert "doi:" in result

    def test_sort_references_with_index(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        refs = [{"index": 3}, {"index": 1}, {"index": 2}]
        sorted_refs = f.sort_references(refs)
        assert [r["index"] for r in sorted_refs] == [1, 2, 3]

    def test_sort_references_without_index(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        refs = [{"title": "B"}, {"title": "A"}]
        sorted_refs = f.sort_references(refs)
        assert sorted_refs == refs

    def test_sort_references_empty(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        assert f.sort_references([]) == []


# ══════════════════════════════════════════════════════════════════════════════
# Caption Matcher (caption_matcher.py)
# ══════════════════════════════════════════════════════════════════════════════

class TestTableCaptionMatcher:
    """Coverage gap: caption_matcher.py was 7.95%"""

    def test_init(self):
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher
        m = TableCaptionMatcher()
        assert m.search_window_above == 2
        assert m.search_window_below == 1
        assert m.caption_regex is not None

    def test_init_custom_window(self):
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher
        m = TableCaptionMatcher(search_window_above=3, search_window_below=2)
        assert m.search_window_above == 3
        assert m.search_window_below == 2

    def test_process_no_tables(self):
        from app.models import PipelineDocument
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher
        doc = MagicMock(spec=PipelineDocument)
        doc.tables = []
        doc.blocks = [MagicMock()]
        m = TableCaptionMatcher()
        result = m.process(doc)
        assert result is doc

    def test_process_no_blocks(self):
        from app.models import PipelineDocument
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher
        doc = MagicMock(spec=PipelineDocument)
        doc.tables = [MagicMock()]
        doc.blocks = []
        m = TableCaptionMatcher()
        result = m.process(doc)
        assert result is doc

    def test_process_match_caption(self):
        from app.models import Block, BlockType, PipelineDocument, Table
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher

        doc = MagicMock(spec=PipelineDocument)
        blocks = [
            MagicMock(spec=Block, block_id="b1", block_type=BlockType.BODY,
                      text="Table 1: Results", index=0, metadata={},
                      is_heading=lambda: False),
            MagicMock(spec=Block, block_id="b2", block_type=BlockType.BODY,
                      text="data row", index=1, metadata={},
                      is_heading=lambda: False),
        ]

        table = MagicMock(spec=Table)
        table.block_index = 1
        table.table_id = "tbl_1"
        table.caption_text = None
        table.caption_block_id = None
        table.metadata = {}

        doc.blocks = blocks
        doc.tables = [table]

        m = TableCaptionMatcher()
        result = m.process(doc)
        assert result is doc
        assert table.caption_text == "Table 1: Results"
        assert table.caption_block_id == "b1"

    def test_process_match_caption_below(self):
        from app.models import Block, BlockType, PipelineDocument, Table
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher

        doc = MagicMock(spec=PipelineDocument)
        blocks = [
            MagicMock(spec=Block, block_id="b1", block_type=BlockType.BODY,
                      text="data row", index=0, metadata={},
                      is_heading=lambda: False),
            MagicMock(spec=Block, block_id="b2", block_type=BlockType.BODY,
                      text="Table 1: Results", index=1, metadata={},
                      is_heading=lambda: False),
        ]

        table = MagicMock(spec=Table)
        table.block_index = 0
        table.table_id = "tbl_1"
        table.caption_text = None
        table.caption_block_id = None
        table.metadata = {}

        doc.blocks = blocks
        doc.tables = [table]

        m = TableCaptionMatcher()
        result = m.process(doc)
        assert result is doc
        assert table.caption_text == "Table 1: Results"
        assert table.caption_block_id == "b2"

    def test_process_skips_heading_blocks(self):
        from app.models import Block, BlockType, PipelineDocument, Table
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher

        doc = MagicMock(spec=PipelineDocument)
        blocks = [
            MagicMock(spec=Block, block_id="b1", block_type=BlockType.HEADING_1,
                      text="Table 1: Results", index=0, metadata={},
                      is_heading=lambda: True),
            MagicMock(spec=Block, block_id="b2", block_type=BlockType.BODY,
                      text="data", index=1, metadata={},
                      is_heading=lambda: False),
        ]

        table = MagicMock(spec=Table)
        table.block_index = 1
        table.table_id = "tbl_1"
        table.caption_text = None
        table.caption_block_id = None
        table.metadata = {}

        doc.blocks = blocks
        doc.tables = [table]

        m = TableCaptionMatcher()
        result = m.process(doc)
        assert table.caption_text is None

    def test_process_sets_missing_status(self):
        from app.models import Block, BlockType, PipelineDocument, Table
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher

        doc = MagicMock(spec=PipelineDocument)
        block = MagicMock(spec=Block, block_id="b1", block_type=BlockType.BODY,
                          text="some text", index=0, metadata={},
                          is_heading=lambda: False)

        table = MagicMock(spec=Table)
        table.block_index = 0
        table.table_id = "tbl_1"
        table.caption_text = None
        table.caption_block_id = None
        table.metadata = {}

        doc.blocks = [block]
        doc.tables = [table]

        m = TableCaptionMatcher()
        m.process(doc)
        assert table.metadata.get("caption_status") == "Missing"

    def test_process_preserves_existing_caption(self):
        from app.models import Block, BlockType, PipelineDocument, Table
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher

        doc = MagicMock(spec=PipelineDocument)
        block = MagicMock(spec=Block, block_id="b1", block_type=BlockType.BODY,
                          text="some text", index=0, metadata={},
                          is_heading=lambda: False)

        table = MagicMock(spec=Table)
        table.block_index = 0
        table.table_id = "tbl_1"
        table.caption_text = "Existing"
        table.caption_block_id = None
        table.metadata = {}

        doc.blocks = [block]
        doc.tables = [table]

        m = TableCaptionMatcher()
        m.process(doc)
        assert table.metadata.get("caption_status") != "Missing"

    def test_caption_regex_matches(self):
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher
        m = TableCaptionMatcher()
        assert m.caption_regex.match("Table 1: Results")
        assert m.caption_regex.match("TABLE 2. Data")
        assert m.caption_regex.match("Table 3 – Summary")
        assert m.caption_regex.match("Table 1.1: Sub")
        assert m.caption_regex.match("Table I: Roman")
        assert m.caption_regex.match("Table A: Letter")
        assert not m.caption_regex.match("Not a caption")
        assert not m.caption_regex.match("Figure 1: Caption")

    def test_find_references_start_index_by_heading(self):
        from app.models import Block, BlockType
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher
        block = MagicMock(spec=Block, block_type=BlockType.REFERENCES_HEADING,
                          text="References", index=10, is_heading=lambda: True)
        m = TableCaptionMatcher()
        assert m._find_references_start_index([block]) == 10

    def test_find_references_start_index_by_keyword(self):
        from app.models import Block, BlockType
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher
        block = MagicMock(spec=Block, block_type=BlockType.HEADING_1,
                          text="References", index=10, is_heading=lambda: True)
        m = TableCaptionMatcher()
        assert m._find_references_start_index([block]) == 10

    def test_find_references_start_index_not_found(self):
        from app.models import Block, BlockType
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher
        block = MagicMock(spec=Block, block_type=BlockType.BODY,
                          text="Some text", index=5, is_heading=lambda: False)
        m = TableCaptionMatcher()
        assert m._find_references_start_index([block]) is None

    def test_match_table_captions_convenience(self):
        from app.pipeline.tables.caption_matcher import match_table_captions
        doc = MagicMock()
        doc.tables = []
        result = match_table_captions(doc)
        assert result is doc

    def test_process_exception_handling(self):
        from app.pipeline.tables.caption_matcher import TableCaptionMatcher
        doc = MagicMock()
        doc.blocks = [MagicMock()]
        doc.tables = [MagicMock()]
        doc.tables[0].block_index = 0
        doc.add_processing_stage = MagicMock()
        m = TableCaptionMatcher()
        result = m.process(doc)
        assert result is doc


# ══════════════════════════════════════════════════════════════════════════════
# Table Extractor (extractor.py)
# ══════════════════════════════════════════════════════════════════════════════

class TestTableExtractor:
    """Coverage gap: extractor.py was 8.46%"""

    @pytest.fixture
    def mock_docx_table(self):
        """Build a mock docx Table with 2 rows x 2 cols."""
        import docx
        cell_00 = MagicMock()
        cell_00.text = "Header1"
        cell_00.paragraphs = []

        cell_01 = MagicMock()
        cell_01.text = "Header2"
        cell_01.paragraphs = []

        cell_10 = MagicMock()
        cell_10.text = "Data1"
        cell_10.paragraphs = []

        cell_11 = MagicMock()
        cell_11.text = "Data2"
        cell_11.paragraphs = []

        row0 = MagicMock()
        row0.cells = [cell_00, cell_01]

        row1 = MagicMock()
        row1.cells = [cell_10, cell_11]

        table = MagicMock()
        table.rows = [row0, row1]
        table.__class__ = docx.table.Table
        return table

    def test_extract(self, mock_docx_table):
        from app.pipeline.tables.extractor import TableExtractor
        extractor = TableExtractor()
        result = extractor.extract(mock_docx_table, "tbl_1", 0, 0)
        assert result.table_id == "tbl_1"
        assert result.num_rows == 2
        assert result.num_cols == 2
        assert result.index == 0
        assert result.block_index == 0
        assert result.data == [["Header1", "Header2"], ["Data1", "Data2"]]

    def test_extract_with_bold_header(self, mock_docx_table):
        from app.pipeline.tables.extractor import TableExtractor
        row0 = mock_docx_table.rows[0]
        for cell in row0.cells:
            run = MagicMock()
            run.bold = True
            para = MagicMock()
            para.runs = [run]
            cell.paragraphs = [para]

        extractor = TableExtractor()
        result = extractor.extract(mock_docx_table, "tbl_1", 0, 0)
        assert result.has_header is True

    def test_extract_header_keywords(self, mock_docx_table):
        from app.pipeline.tables.extractor import TableExtractor
        cell_00 = mock_docx_table.rows[0].cells[0]
        cell_00.text = "Name"
        cell_01 = mock_docx_table.rows[0].cells[1]
        cell_01.text = "Date"

        extractor = TableExtractor()
        result = extractor.extract(mock_docx_table, "tbl_1", 0, 0)
        assert result.has_header is True

    def test_extract_uneven_rows(self, mock_docx_table):
        from app.pipeline.tables.extractor import TableExtractor
        cell_01 = mock_docx_table.rows[0].cells[1]
        cell_01.text = ""

        extractor = TableExtractor()
        result = extractor.extract(mock_docx_table, "tbl_1", 0, 0)
        assert result.num_cols == 2
        assert len(result.data[1]) == 2

    def test_normalize_cell_text(self):
        from app.pipeline.tables.extractor import TableExtractor
        e = TableExtractor()
        assert e._normalize_cell_text("  hello  ") == "hello"
        assert e._normalize_cell_text("") == ""
        assert e._normalize_cell_text(None) == ""

    def test_contains_header_keywords(self):
        from app.pipeline.tables.extractor import TableExtractor
        e = TableExtractor()
        assert e._contains_header_keywords(["Name", "Date", "Value"]) is True
        assert e._contains_header_keywords(["x", "y", "z"]) is False
        assert e._contains_header_keywords(["Qty"]) is True

    def test_deep_xml_text_normal(self):
        from app.pipeline.tables.extractor import TableExtractor
        e = TableExtractor()
        cell = MagicMock()
        cell._tc = MagicMock()
        cell._tc.iter.return_value = []
        assert e._extract_deep_xml_text(cell) == ""

    def test_deep_xml_text_exception(self):
        from app.pipeline.tables.extractor import TableExtractor
        e = TableExtractor()
        cell = MagicMock()
        cell._tc.iter.side_effect = Exception("boom")
        assert e._extract_deep_xml_text(cell) == ""

    def test_is_cell_bold_true(self):
        from app.pipeline.tables.extractor import TableExtractor
        e = TableExtractor()
        run = MagicMock()
        run.bold = True
        para = MagicMock()
        para.runs = [run]
        cell = MagicMock()
        cell.paragraphs = [para]
        assert e._is_cell_bold(cell) is True

    def test_is_cell_bold_false(self):
        from app.pipeline.tables.extractor import TableExtractor
        e = TableExtractor()
        run = MagicMock()
        run.bold = False
        para = MagicMock()
        para.runs = [run]
        cell = MagicMock()
        cell.paragraphs = [para]
        assert e._is_cell_bold(cell) is False


# ══════════════════════════════════════════════════════════════════════════════
# Table Renderer (renderer.py)
# ══════════════════════════════════════════════════════════════════════════════

class TestTableRenderer:
    """Coverage gap: renderer.py was 13.04%"""

    def test_render_no_table(self):
        from app.pipeline.tables.renderer import TableRenderer
        doc = MagicMock()
        renderer = TableRenderer()
        renderer.render(doc, None)
        doc.add_table.assert_not_called()

    def test_render_no_rows(self):
        from app.models import Table
        from app.pipeline.tables.renderer import TableRenderer
        doc = MagicMock()
        table = MagicMock(spec=Table)
        table.rows = []
        renderer = TableRenderer()
        renderer.render(doc, table)
        doc.add_table.assert_not_called()

    def test_render_with_caption(self):
        from app.models import Table, TableCell
        from app.pipeline.tables.renderer import TableRenderer

        doc = MagicMock()
        doc.styles = {"Caption": MagicMock(), "Table Grid": MagicMock()}
        table = MagicMock(spec=Table)
        table.rows = [["a", "b"]]
        table.cells = [
            MagicMock(spec=TableCell, row=0, col=0, text="a", metadata={}),
            MagicMock(spec=TableCell, row=0, col=1, text="b", metadata={}),
        ]
        table.caption_text = "Table 1: Results"
        table.index = 0
        table.num_cols = 2

        renderer = TableRenderer()
        renderer.render(doc, table, number=1)
        doc.add_paragraph.assert_called_once()
        doc.add_table.assert_called_once()

    def test_render_without_caption(self):
        from app.models import Table, TableCell
        from app.pipeline.tables.renderer import TableRenderer

        doc = MagicMock()
        doc.styles = {}
        table = MagicMock(spec=Table)
        table.rows = [["a", "b"]]
        table.cells = [
            MagicMock(spec=TableCell, row=0, col=0, text="a", metadata={}),
            MagicMock(spec=TableCell, row=0, col=1, text="b", metadata={}),
        ]
        table.caption_text = None
        table.index = 0
        table.num_cols = 2

        renderer = TableRenderer()
        renderer.render(doc, table)
        doc.add_paragraph.assert_not_called()
        doc.add_table.assert_called_once()

    def test_render_nested_tables(self):
        from app.models import Table, TableCell
        from app.pipeline.tables.renderer import TableRenderer

        doc = MagicMock()
        doc.styles = {}
        nested_table = MagicMock(spec=Table)
        nested_table.rows = [["x"]]
        nested_table.cells = []
        nested_table.caption_text = None

        cell = MagicMock(spec=TableCell, row=0, col=0, text="a", metadata={})
        cell.metadata = {"nested_tables": [nested_table]}

        table = MagicMock(spec=Table)
        table.rows = [["a"]]
        table.cells = [cell]
        table.caption_text = None
        table.index = 0
        table.num_cols = 1

        renderer = TableRenderer()
        renderer.render(doc, table)
        doc.add_table.assert_called()

    def test_render_caption_not_prefixed(self):
        from app.models import Table, TableCell
        from app.pipeline.tables.renderer import TableRenderer

        doc = MagicMock()
        doc.styles = {"Caption": MagicMock(), "Table Grid": MagicMock()}
        table = MagicMock(spec=Table)
        table.rows = [["a"]]
        table.cells = [
            MagicMock(spec=TableCell, row=0, col=0, text="a", metadata={}),
        ]
        table.caption_text = "Experimental Results"
        table.index = 0
        table.num_cols = 1

        renderer = TableRenderer()
        renderer.render(doc, table)
        doc.add_paragraph.assert_called_once()

    def test_render_zero_cols(self):
        from app.models import Table
        from app.pipeline.tables.renderer import TableRenderer

        doc = MagicMock()
        table = MagicMock(spec=Table)
        table.rows = [[]]
        table.caption_text = None
        table.num_cols = 0

        renderer = TableRenderer()
        renderer.render(doc, table)
        doc.add_table.assert_not_called()

    def test_render_style_exception(self):
        from app.models import Table, TableCell
        from app.pipeline.tables.renderer import TableRenderer

        doc = MagicMock()
        doc.styles.__contains__.side_effect = KeyError("missing")
        doc.styles.__getitem__.side_effect = KeyError("missing")
        table = MagicMock(spec=Table)
        table.rows = [["a"]]
        table.cells = [
            MagicMock(spec=TableCell, row=0, col=0, text="a", metadata={}),
        ]
        table.caption_text = None
        table.index = 0
        table.num_cols = 1

        renderer = TableRenderer()
        renderer.render(doc, table)
        doc.add_table.assert_called_once()


# ══════════════════════════════════════════════════════════════════════════════
# Document Validator v3 (validator_v3.py)
# ══════════════════════════════════════════════════════════════════════════════

class TestDocumentValidatorV3:
    """Coverage gap: validator_v3.py was 15.67%"""

    @pytest.fixture
    def validator(self):
        from app.pipeline.validation.validator_v3 import DocumentValidator
        with patch("app.pipeline.validation.validator_v3.ContractLoader"), \
             patch("app.pipeline.validation.validator_v3.SectionOrderValidator"), \
             patch("app.pipeline.validation.validator_v3.CrossReferenceEngine"), \
             patch("app.pipeline.validation.validator_v3.CrossRefClient"), \
             patch("app.pipeline.validation.validator_v3.ReviewManager"):
            yield DocumentValidator()

    def test_init(self, validator):
        assert validator.contract_loader is not None
        assert validator.order_validator is not None
        assert validator.integrity_engine is not None
        assert validator.crossref_client is not None

    def test_as_bool_none(self, validator):
        assert validator._as_bool(None) is False
        assert validator._as_bool(None, True) is True

    def test_as_bool_bool(self, validator):
        assert validator._as_bool(True) is True
        assert validator._as_bool(False) is False

    def test_as_bool_int(self, validator):
        assert validator._as_bool(1) is True
        assert validator._as_bool(0) is False

    def test_as_bool_string_true(self, validator):
        assert validator._as_bool("true") is True
        assert validator._as_bool("1") is True
        assert validator._as_bool("yes") is True
        assert validator._as_bool("on") is True

    def test_as_bool_string_false(self, validator):
        assert validator._as_bool("false") is False
        assert validator._as_bool("0") is False

    def test_as_bool_string_other(self, validator):
        assert validator._as_bool("maybe") is False

    def test_process_calls_validate(self, validator):
        doc = MagicMock()
        with patch.object(validator, "validate", return_value=MagicMock()) as mock_val:
            result = validator.process(doc)
            mock_val.assert_called_once_with(doc)
        assert result is doc

    def test_process_safe_execution_error(self, validator):
        doc = MagicMock()
        with patch.object(validator, "validate", side_effect=Exception("boom")):
            result = validator.process(doc)
        assert result is doc

    def test_validate_full_happy_path(self, validator):
        from app.models import PipelineDocument
        doc = MagicMock(spec=PipelineDocument)
        doc.references = []
        doc.figures = []
        doc.tables = []
        doc.formatting_options = {}
        doc.get_section_names.return_value = []
        doc.get_stats.return_value = {"blocks": 0}

        with patch.object(validator, "_check_sections", return_value=([], [])), \
             patch.object(validator, "_check_figures", return_value=([], [])), \
             patch.object(validator, "_check_references", return_value=([], [])), \
             patch.object(validator, "_check_tables", return_value=([], [])), \
             patch.object(validator, "_check_reference_integrity", return_value=([], [])):
            validator.integrity_engine.validate_integrity.return_value = []
            result = validator.validate(doc)
        assert result.is_valid is True

    def test_validate_with_fast_mode(self, validator):
        from app.models import PipelineDocument
        doc = MagicMock(spec=PipelineDocument)
        doc.references = []
        doc.figures = []
        doc.tables = []
        doc.formatting_options = {"fast_mode": "true"}
        doc.get_section_names.return_value = []
        doc.get_stats.return_value = {"blocks": 0}

        with patch.object(validator, "_check_sections", return_value=([], [])), \
             patch.object(validator, "_check_figures", return_value=([], [])), \
             patch.object(validator, "_check_references", return_value=([], [])), \
             patch.object(validator, "_check_tables", return_value=([], [])), \
             patch.object(validator, "_check_reference_integrity") as mock_doi:
            validator.integrity_engine.validate_integrity.return_value = []
            result = validator.validate(doc)
        mock_doi.assert_not_called()

    def test_validate_with_errors(self, validator):
        from app.models import PipelineDocument
        doc = MagicMock(spec=PipelineDocument)
        doc.references = []
        doc.figures = []
        doc.tables = []
        doc.formatting_options = {}
        doc.get_section_names.return_value = []
        doc.get_stats.return_value = {"blocks": 0}

        with patch.object(validator, "_check_sections", return_value=(["Missing required section"], [])), \
             patch.object(validator, "_check_figures", return_value=([], [])), \
             patch.object(validator, "_check_references", return_value=([], [])), \
             patch.object(validator, "_check_tables", return_value=([], [])), \
             patch.object(validator, "_check_reference_integrity", return_value=([], [])):
            validator.integrity_engine.validate_integrity.return_value = ["Dangling ref"]
            result = validator.validate(doc)
        assert result.is_valid is False

    def test_check_sections(self, validator):
        from app.models import PipelineDocument
        doc = MagicMock(spec=PipelineDocument)
        doc.template = MagicMock()
        doc.template.template_name = "ieee"
        with patch.object(validator.order_validator, "validate_order",
                          return_value=["Missing required section", "Reorder warning"]):
            errors, warnings = validator._check_sections(doc)
        assert "Missing required section" in errors
        assert "Reorder warning" in warnings

    def test_check_sections_fallback_publisher(self, validator):
        doc = MagicMock()
        doc.template = None
        with patch.object(validator.order_validator, "validate_order",
                          return_value=[]):
            errors, warnings = validator._check_sections(doc)
        assert errors == []

    def test_check_sections_exception(self, validator):
        from app.models import PipelineDocument
        doc = MagicMock(spec=PipelineDocument)
        doc.template = MagicMock()
        doc.template.template_name = "ieee"
        validator.order_validator.validate_order.side_effect = Exception("order fail")
        errors, warnings = validator._check_sections(doc)
        assert "Section order check skipped due to internal error" in warnings

    def test_check_figures(self, validator):
        from app.models import Figure
        fig1 = MagicMock(spec=Figure)
        fig1.has_caption.return_value = True
        fig1.figure_id = "fig1"
        fig2 = MagicMock(spec=Figure)
        fig2.has_caption.return_value = False
        fig2.figure_id = "fig2"

        doc = MagicMock()
        doc.figures = [fig1, fig2]
        errors, warnings = validator._check_figures(doc)
        assert len(errors) == 0
        assert any("fig2" in w for w in warnings)

    def test_check_references_empty_section_found(self, validator):
        doc = MagicMock()
        doc.references = []
        doc.get_section_names.return_value = ["References"]
        errors, warnings = validator._check_references(doc)
        assert any("References section found" in w for w in warnings)

    def test_check_references_empty_no_section(self, validator):
        doc = MagicMock()
        doc.references = []
        doc.get_section_names.return_value = []
        errors, warnings = validator._check_references(doc)
        assert errors == []
        assert warnings == []

    def test_check_references_with_issues(self, validator):
        from app.models import Reference
        ref1 = MagicMock(spec=Reference)
        ref1.citation_key = "ref1"
        ref1.year = None
        ref1.authors = []
        ref1.title = None
        doc = MagicMock()
        doc.references = [ref1]
        doc.get_section_names.return_value = ["References"]
        errors, warnings = validator._check_references(doc)
        assert any("ref1" in e for e in errors)
        assert any("ref1" in w for w in warnings)

    def test_check_references_partial(self, validator):
        from app.models import Reference
        ref = MagicMock(spec=Reference)
        ref.citation_key = "ref1"
        ref.year = 2020
        ref.authors = ["Smith, J."]
        ref.title = "A Study"
        doc = MagicMock()
        doc.references = [ref]
        doc.get_section_names.return_value = ["References"]
        errors, warnings = validator._check_references(doc)
        assert errors == []
        assert warnings == []

    def test_check_tables(self, validator):
        from app.models import Table
        tbl1 = MagicMock(spec=Table)
        tbl1.caption_text = "Table 1: Results"
        tbl2 = MagicMock(spec=Table)
        tbl2.caption_text = None
        doc = MagicMock()
        doc.tables = [tbl1, tbl2]
        errors, warnings = validator._check_tables(doc)
        assert errors == []
        assert any("Table 2" in w for w in warnings)

    def test_check_tables_no_attr(self, validator):
        doc = MagicMock()
        del doc.tables
        errors, warnings = validator._check_tables(doc)
        assert errors == []
        assert warnings == []

    def test_check_reference_integrity_no_refs(self, validator):
        doc = MagicMock()
        doc.references = []
        errors, warnings = validator._check_reference_integrity(doc)
        assert errors == []
        assert warnings == []

    def test_check_reference_integrity_valid_doi(self, validator):
        from app.models import Reference
        ref = MagicMock(spec=Reference)
        ref.citation_key = "ref1"
        ref.doi = "10.1234/abc"
        ref.has_doi.return_value = True
        ref.metadata = {}
        ref.title = "A Study"
        ref.year = 2020
        ref.authors = ["Smith, J."]

        validator.crossref_client.validate_doi.return_value = True
        validator.crossref_client.get_metadata.return_value = {"title": "Real"}
        validator.crossref_client.calculate_confidence.return_value = 0.95

        doc = MagicMock()
        doc.references = [ref]
        errors, warnings = validator._check_reference_integrity(doc)
        assert errors == []
        assert warnings == []
        assert ref.metadata["validation"]["doi_valid"] is True

    def test_check_reference_integrity_low_confidence(self, validator):
        from app.models import Reference
        ref = MagicMock(spec=Reference)
        ref.citation_key = "ref1"
        ref.doi = "10.1234/abc"
        ref.has_doi.return_value = True
        ref.metadata = {}
        ref.title = "A Study"
        ref.year = 2020
        ref.authors = ["Smith, J."]

        validator.crossref_client.validate_doi.return_value = True
        validator.crossref_client.get_metadata.return_value = {"title": "Real"}
        validator.crossref_client.calculate_confidence.return_value = 0.3

        doc = MagicMock()
        doc.references = [ref]
        errors, warnings = validator._check_reference_integrity(doc)
        assert any("low confidence" in w for w in warnings)

    def test_check_reference_integrity_invalid_doi(self, validator):
        from app.models import Reference
        ref = MagicMock(spec=Reference)
        ref.citation_key = "ref1"
        ref.doi = "10.1234/abc"
        ref.has_doi.return_value = True
        ref.metadata = {}

        validator.crossref_client.validate_doi.return_value = False

        doc = MagicMock()
        doc.references = [ref]
        errors, warnings = validator._check_reference_integrity(doc)
        assert any("invalid DOI" in w for w in warnings)
        assert ref.metadata["validation"]["confidence"] == 0.0

    def test_check_reference_integrity_metadata_fetch_fails(self, validator):
        from app.models import Reference
        ref = MagicMock(spec=Reference)
        ref.citation_key = "ref1"
        ref.doi = "10.1234/abc"
        ref.has_doi.return_value = True
        ref.metadata = {}

        validator.crossref_client.validate_doi.return_value = True
        validator.crossref_client.get_metadata.side_effect = Exception("API error")

        doc = MagicMock()
        doc.references = [ref]
        errors, warnings = validator._check_reference_integrity(doc)
        assert any("Failed to fetch" in w for w in warnings)

    def test_check_reference_integrity_validate_fails(self, validator):
        from app.models import Reference
        ref = MagicMock(spec=Reference)
        ref.citation_key = "ref1"
        ref.doi = "10.1234/abc"
        ref.has_doi.return_value = True
        ref.metadata = {}

        validator.crossref_client.validate_doi.side_effect = Exception("boom")

        doc = MagicMock()
        doc.references = [ref]
        errors, warnings = validator._check_reference_integrity(doc)
        assert any("CrossRef validation failed" in w for w in warnings)

    def test_check_reference_integrity_safe_wrapper(self, validator):
        doc = MagicMock()
        doc.references = [MagicMock()]
        validator.crossref_client.validate_doi.side_effect = Exception("API down")
        with patch("app.pipeline.validation.validator_v3.safe_function",
                   lambda **kw: lambda f: f):
            errors, warnings = validator._check_reference_integrity(doc)
        assert any("CrossRef validation failed" in w for w in warnings)

    def test_integrity_violations_dangling(self, validator):
        from app.models import PipelineDocument
        doc = MagicMock(spec=PipelineDocument)
        doc.references = []
        doc.figures = []
        doc.tables = []
        doc.formatting_options = {}
        doc.get_section_names.return_value = []
        doc.get_stats.return_value = {"blocks": 0}

        validator.integrity_engine.validate_integrity.return_value = [
            "Dangling reference to Figure 2"
        ]

        with patch.object(validator, "_check_sections", return_value=([], [])), \
             patch.object(validator, "_check_figures", return_value=([], [])), \
             patch.object(validator, "_check_references", return_value=([], [])), \
             patch.object(validator, "_check_tables", return_value=([], [])), \
             patch.object(validator, "_check_reference_integrity", return_value=([], [])):
            result = validator.validate(doc)
        assert any("Dangling" in e for e in result.errors)

    def test_validate_document_convenience(self):
        from app.models import PipelineDocument
        doc = MagicMock(spec=PipelineDocument)
        doc.references = []
        doc.figures = []
        doc.tables = []
        doc.formatting_options = {}
        doc.get_section_names.return_value = []
        doc.get_stats.return_value = {"blocks": 0}

        with patch("app.pipeline.validation.validator_v3.DocumentValidator") as mock_cls:
            instance = mock_cls.return_value
            instance.validate.return_value = MagicMock(is_valid=True)
            from app.pipeline.validation.validator_v3 import validate_document
            result = validate_document(doc)
            assert result.is_valid is True


# ══════════════════════════════════════════════════════════════════════════════
# AI Explainer (ai_explainer.py)
# ══════════════════════════════════════════════════════════════════════════════

class TestAIExplainer:
    """Coverage gap: ai_explainer.py was 14.29%"""

    def test_init(self):
        from app.pipeline.validation.ai_explainer import AIExplainer
        e = AIExplainer()
        assert "missing_sections" in e.explanation_map
        assert "citation_format" in e.explanation_map

    def test_explain_results_empty(self):
        from app.pipeline.validation.ai_explainer import AIExplainer
        e = AIExplainer()
        result = e.explain_results({"errors": []})
        assert result == []

    def test_explain_results_missing_sections(self):
        from app.pipeline.validation.ai_explainer import AIExplainer
        e = AIExplainer()
        result = e.explain_results({"errors": ["missing section: Introduction"]}, "IEEE")
        assert any("missing" in r.lower() for r in result)
        assert any("IEEE" in r for r in result)

    def test_explain_results_reference_error(self):
        from app.pipeline.validation.ai_explainer import AIExplainer
        e = AIExplainer()
        result = e.explain_results({"errors": ["reference missing DOI"]}, "APA")
        assert any("reference" in r.lower() for r in result)

    def test_explain_results_general_error(self):
        from app.pipeline.validation.ai_explainer import AIExplainer
        e = AIExplainer()
        result = e.explain_results({"errors": ["some other error"]}, "IEEE")
        assert len(result) == 1

    def test_explain_results_dict_errors(self):
        from app.pipeline.validation.ai_explainer import AIExplainer
        e = AIExplainer()
        result = e.explain_results({
            "errors": [{"category": "citation_format", "message": "wrong style"}]
        }, "IEEE")
        assert any("citations" in r.lower() for r in result)

    def test_explain_results_dict_unknown_category(self):
        from app.pipeline.validation.ai_explainer import AIExplainer
        e = AIExplainer()
        result = e.explain_results({
            "errors": [{"category": "unknown_cat", "message": "weird issue"}]
        }, "IEEE")
        assert len(result) == 1

    def test_explain_results_mixed_errors(self):
        from app.pipeline.validation.ai_explainer import AIExplainer
        e = AIExplainer()
        result = e.explain_results({
            "errors": [
                "missing section: Methods",
                {"category": "figure_captions", "message": "no labels"}
            ]
        }, "Nature")
        assert len(result) == 2


# ══════════════════════════════════════════════════════════════════════════════
# Review Manager (review_manager.py)
# ══════════════════════════════════════════════════════════════════════════════

class TestReviewManager:
    """Coverage gap: review_manager.py was 7.95%"""

    def test_init_defaults(self):
        from app.pipeline.validation.review_manager import ReviewManager
        rm = ReviewManager()
        assert rm.review_threshold == 0.70
        assert rm.critical_threshold == 0.45

    def test_init_custom(self):
        from app.pipeline.validation.review_manager import ReviewManager
        rm = ReviewManager(review_threshold=0.8, critical_threshold=0.3)
        assert rm.review_threshold == 0.8
        assert rm.critical_threshold == 0.3

    def test_init_critical_gte_review(self):
        from app.pipeline.validation.review_manager import ReviewManager
        with pytest.raises(ValueError, match="must be less than"):
            ReviewManager(review_threshold=0.5, critical_threshold=0.5)

    def test_init_invalid_threshold(self):
        from app.pipeline.validation.review_manager import ReviewManager
        with pytest.raises(ValueError, match="between 0.0 and 1.0"):
            ReviewManager(review_threshold=1.5, critical_threshold=0.5)

    def test_evaluate_all_ok(self):
        from app.pipeline.validation.review_manager import ReviewManager
        from app.models import ReviewStatus

        doc = MagicMock()
        doc.blocks = []
        doc.metadata = MagicMock()
        doc.metadata.ai_hints = {}
        doc.review = None

        rm = ReviewManager()
        result = rm.evaluate(doc)
        assert result.review.status == ReviewStatus.OK

    def test_evaluate_critical_threshold(self):
        from app.pipeline.validation.review_manager import ReviewManager
        from app.models import ReviewStatus

        doc = MagicMock()
        block = MagicMock()
        block.metadata = {"classification_confidence": 0.3}
        block.semantic_intent = "body"
        block.block_id = "abcdefgh1234"
        doc.blocks = [block]
        doc.metadata = MagicMock()
        doc.metadata.ai_hints = {}

        rm = ReviewManager()
        result = rm.evaluate(doc)
        assert result.review.status == ReviewStatus.CRITICAL

    def test_evaluate_review_threshold(self):
        from app.pipeline.validation.review_manager import ReviewManager
        from app.models import ReviewStatus

        doc = MagicMock()
        block = MagicMock()
        block.metadata = {"classification_confidence": 0.6}
        block.semantic_intent = "body"
        block.block_id = "abcdefgh1234"
        doc.blocks = [block]
        doc.metadata = MagicMock()
        doc.metadata.ai_hints = {}

        rm = ReviewManager()
        result = rm.evaluate(doc)
        assert result.review.status == ReviewStatus.REVIEW

    def test_evaluate_missing_confidence(self):
        from app.pipeline.validation.review_manager import ReviewManager
        from app.models import ReviewStatus

        doc = MagicMock()
        block = MagicMock()
        block.metadata = {}
        block.semantic_intent = "body"
        block.block_id = "abcdefgh1234"
        doc.blocks = [block]
        doc.metadata = MagicMock()
        doc.metadata.ai_hints = {}

        rm = ReviewManager()
        result = rm.evaluate(doc)
        assert result.review.status == ReviewStatus.OK

    def test_evaluate_uses_nlp_confidence_fallback(self):
        from app.pipeline.validation.review_manager import ReviewManager
        from app.models import ReviewStatus

        doc = MagicMock()
        block = MagicMock()
        block.metadata = {"nlp_confidence": 0.6}
        block.classification_confidence = None
        block.semantic_intent = "body"
        block.block_id = "abcdefgh1234"
        doc.blocks = [block]
        doc.metadata = MagicMock()
        doc.metadata.ai_hints = {}

        rm = ReviewManager()
        result = rm.evaluate(doc)
        assert result.review.status == ReviewStatus.REVIEW

    def test_evaluate_uses_classification_confidence_attribute(self):
        from app.pipeline.validation.review_manager import ReviewManager
        from app.models import ReviewStatus

        doc = MagicMock()
        block = MagicMock()
        block.metadata = {}
        block.classification_confidence = 0.3
        block.semantic_intent = "body"
        block.block_id = "abcdefgh1234"
        doc.blocks = [block]
        doc.metadata = MagicMock()
        doc.metadata.ai_hints = {}

        rm = ReviewManager()
        result = rm.evaluate(doc)
        assert result.review.status == ReviewStatus.CRITICAL

    def test_evaluate_with_ai_hints(self):
        from app.pipeline.validation.review_manager import ReviewManager
        from app.models import ReviewStatus

        doc = MagicMock()
        block = MagicMock()
        block.metadata = {"classification_confidence": 0.9}
        block.classification_confidence = 0.9
        block.semantic_intent = "body"
        block.block_id = "abcdefgh1234"
        doc.blocks = [block]
        doc.metadata = MagicMock()
        doc.metadata.ai_hints = {"semantic_advice": {"confidence": 0.5}}

        rm = ReviewManager()
        result = rm.evaluate(doc)
        assert result.review.status == ReviewStatus.REVIEW

    def test_evaluate_invalid_confidence_value(self):
        from app.pipeline.validation.review_manager import ReviewManager
        from app.models import ReviewStatus

        doc = MagicMock()
        block = MagicMock()
        block.metadata = {"classification_confidence": "invalid"}
        block.semantic_intent = "body"
        block.block_id = "abcdefgh1234"
        doc.blocks = [block]
        doc.metadata = MagicMock()
        doc.metadata.ai_hints = {}

        rm = ReviewManager()
        result = rm.evaluate(doc)
        assert result.review is not None

    def test_evaluate_semantic_intent_from_metadata(self):
        from app.pipeline.validation.review_manager import ReviewManager
        from app.models import ReviewStatus

        doc = MagicMock()
        block = MagicMock()
        block.metadata = {"classification_confidence": 0.3, "semantic_intent": "abstract"}
        block.semantic_intent = None
        block.block_id = "abcdefgh1234"
        doc.blocks = [block]
        doc.metadata = MagicMock()
        doc.metadata.ai_hints = {}

        rm = ReviewManager()
        result = rm.evaluate(doc)
        assert result.review.status == ReviewStatus.CRITICAL
        assert any("abstract" in f for f in result.review.flags)

    def test_evaluate_flags_limited_to_top_5(self):
        from app.pipeline.validation.review_manager import ReviewManager
        from app.models import ReviewStatus

        doc = MagicMock()
        blocks = []
        for i in range(10):
            b = MagicMock()
            b.metadata = {"classification_confidence": 0.3}
            b.semantic_intent = f"section_{i}"
            b.block_id = f"block{i:08d}"
            blocks.append(b)
        doc.blocks = blocks
        doc.metadata = MagicMock()
        doc.metadata.ai_hints = {}

        rm = ReviewManager()
        result = rm.evaluate(doc)
        assert len(result.review.flags) <= 5


# ══════════════════════════════════════════════════════════════════════════════
# Extra edge-case tests to push coverage past 90%
# ══════════════════════════════════════════════════════════════════════════════

class TestCoverageEdgeCases:

    # --- APA additional edge cases ---

    def test_apa_book_chapter_no_publisher_no_pages(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="My Chapter",
            book_title="Big Book", reference_type="book_chapter"
        )
        assert "Big Book" in result
        assert "pp." not in result

    def test_apa_thesis_with_publisher_and_doi(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="My Thesis",
            publisher="MIT", doi="10.1234/thesis", reference_type="thesis"
        )
        assert "MIT" in result
        assert "https://doi.org/10.1234/thesis" in result

    def test_apa_default_empty(self):
        from app.pipeline.references.csl.apafallback import APA7Formatter
        f = APA7Formatter(hanging_indent=False)
        result = f.format_reference_entry(
            ["Smith, J."], year=2020, title="Misc",
            reference_type="patent"
        )
        assert "Misc" in result

    # --- Vancouver additional edge cases ---

    def test_vancouver_journal_no_pages_no_doi(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f.format_reference_entry(
            ["Smith J"], title="Title",
            journal="J Sci", year=2020, volume="10", issue="2",
            reference_type="journal_article"
        )
        assert "Title" in result
        assert result.endswith(".")

    def test_vancouver_doi_invalid_format(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f._format_doi("httpinvalid")
        assert "doi:" in result

    def test_vancouver_doi_http_no_slash(self):
        from app.pipeline.references.csl.vancouver_fallback import VancouverFormatter
        f = VancouverFormatter()
        result = f._format_doi("http://")
        assert "doi:" in result

    # --- Validator additional edge cases ---

    def test_validator_integrity_warning(self):
        from app.pipeline.validation.validator_v3 import DocumentValidator
        from app.models import PipelineDocument
        with patch("app.pipeline.validation.validator_v3.ContractLoader"), \
             patch("app.pipeline.validation.validator_v3.SectionOrderValidator"), \
             patch("app.pipeline.validation.validator_v3.CrossReferenceEngine"), \
             patch("app.pipeline.validation.validator_v3.CrossRefClient"), \
             patch("app.pipeline.validation.validator_v3.ReviewManager"):
            v = DocumentValidator()
            doc = MagicMock(spec=PipelineDocument)
            doc.references = []
            doc.figures = []
            doc.tables = []
            doc.formatting_options = {}
            doc.get_section_names.return_value = []
            doc.get_stats.return_value = {"blocks": 0}
            doc.blocks = []

            v.integrity_engine.validate_integrity.return_value = [
                "Cross-reference mismatch in Section 3"
            ]

            with patch.object(v, "_check_sections", return_value=([], [])), \
                 patch.object(v, "_check_figures", return_value=([], [])), \
                 patch.object(v, "_check_references", return_value=([], [])), \
                 patch.object(v, "_check_tables", return_value=([], [])), \
                 patch.object(v, "_check_reference_integrity", return_value=([], [])):
                result = v.validate(doc)
        assert any("Cross-reference mismatch" in w for w in result.warnings)

    def test_validator_template_exception(self):
        from app.pipeline.validation.validator_v3 import DocumentValidator
        with patch("app.pipeline.validation.validator_v3.ContractLoader"), \
             patch("app.pipeline.validation.validator_v3.SectionOrderValidator"), \
             patch("app.pipeline.validation.validator_v3.CrossReferenceEngine"), \
             patch("app.pipeline.validation.validator_v3.CrossRefClient"), \
             patch("app.pipeline.validation.validator_v3.ReviewManager"):
            v = DocumentValidator()
            doc = MagicMock()
            doc.template = MagicMock()
            doc.template.template_name = "ieee"
            v.order_validator.validate_order.side_effect = Exception("order fail")
            errors, warnings = v._check_sections(doc)
        assert any("internal error" in w for w in warnings)

    # --- Extractor edge cases ---

    def test_extractor_deep_xml_fallback(self):
        from app.pipeline.tables.extractor import TableExtractor
        e = TableExtractor()
        cell = MagicMock()
        cell._tc = MagicMock()
        node1 = MagicMock()
        node1.tag = "}t"
        node1.text = "deep"
        node2 = MagicMock()
        node2.tag = "}t"
        node2.text = " text"
        cell._tc.iter.return_value = [node1, node2]
        result = e._extract_deep_xml_text(cell)
        assert result == "deep text"
