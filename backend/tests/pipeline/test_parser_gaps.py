# SPDX-License-Identifier: MIT
# Copyright (c) 2026 ScholarForm AI

from __future__ import annotations

from datetime import UTC, datetime
from typing import Any
from unittest.mock import MagicMock, patch

from app.models import Block, BlockType, Equation, Figure, Table
from app.pipeline.parsing.parser import DocxParser

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_run(
    text: str = "",
    bold: bool | None = False,
    italic: bool | None = False,
    underline: bool | None = False,
    font_name: str | None = None,
    font_size_pt: float | None = None,
    element_findall_result: list | None = None,
) -> MagicMock:
    """Build a MagicMock that looks like a python-docx Run."""
    run = MagicMock()
    run.text = text
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = font_name
    run.font.size.pt = font_size_pt if font_size_pt is not None else None
    if font_size_pt is None:
        run.font.size = None
    run._element = MagicMock()
    run._element.findall.return_value = element_findall_result if element_findall_result is not None else []
    run.part = MagicMock()
    return run


def _make_paragraph(
    text: str = "",
    style_name: str | None = None,
    runs: list[MagicMock] | None = None,
    alignment: Any = None,
    font_bold: bool | None = None,
    font_italic: bool | None = None,
    font_name: str | None = None,
    font_size_pt: float | None = None,
    element_findall_result: list | None = None,
) -> MagicMock:
    """Build a MagicMock that looks like a python-docx Paragraph."""
    para = MagicMock()
    para.text = text
    para.alignment = alignment
    para.runs = runs if runs is not None else []
    para._element = MagicMock()
    para._element.findall.return_value = element_findall_result if element_findall_result is not None else []
    para._element.find.return_value = None

    style_mock = MagicMock()
    style_mock.name = style_name
    style_mock.font.bold = font_bold
    style_mock.font.italic = font_italic
    style_mock.font.name = font_name
    style_mock.font.size.pt = font_size_pt if font_size_pt is not None else None
    if font_size_pt is None:
        style_mock.font.size = None
    para.style = style_mock
    return para


def _make_docx_mock(
    body_elements: list | None = None,
    sections: list | None = None,
    footnotes_part: Any = None,
    endnotes_part: Any = None,
) -> MagicMock:
    """Build a MagicMock that looks like a python-docx Document."""
    docx = MagicMock()
    docx.core_properties = MagicMock(
        title=None, author=None, subject=None, keywords=None, created=None,
    )
    docx.element.body = body_elements if body_elements is not None else []
    docx.sections = sections if sections is not None else []
    docx.part = MagicMock(footnotes_part=footnotes_part, endnotes_part=endnotes_part)
    return docx


# ---------------------------------------------------------------------------
# Parse() method – missing branches
# ---------------------------------------------------------------------------

class TestParseGaps:
    """Cover gaps in DocxParser.parse() and surrounding logic."""

    def test_parse_with_headers_footers(self, tmp_path):
        """Line 141: header_footer_blocks is non-empty -> extend."""
        f = tmp_path / "with_hf.docx"
        f.write_text("dummy")
        with patch("app.pipeline.parsing.parser.DocxDocument") as mock_docx_cls:
            docx = _make_docx_mock()
            mock_docx_cls.return_value = docx
            p = DocxParser()

            hf_block = Block(
                block_id="hf_001", text="Header", index=0,
                block_type=BlockType.UNKNOWN,
            )
            with patch.object(p, "_extract_headers_and_footers", return_value=[hf_block]):
                doc = p.parse(str(f), "doc1")
                assert len(doc.blocks) == 1
                assert "Header" in [b.text for b in doc.blocks]

    def test_parse_with_notes(self, tmp_path):
        """Line 148: note_blocks is non-empty -> extend."""
        f = tmp_path / "with_notes.docx"
        f.write_text("dummy")
        with patch("app.pipeline.parsing.parser.DocxDocument") as mock_docx_cls:
            docx = _make_docx_mock()
            mock_docx_cls.return_value = docx
            p = DocxParser()

            note = Block(
                block_id="fn_001", text="A footnote", index=0,
                block_type=BlockType.UNKNOWN,
            )
            note.metadata["is_footnote"] = True
            with patch.object(p, "_extract_footnotes_and_endnotes", return_value=[note]):
                doc = p.parse(str(f), "doc1")
                assert len(doc.blocks) == 1
                assert doc.blocks[0].metadata.get("is_footnote") is True

    def test_parse_with_notes_and_headers(self, tmp_path):
        """Line 159: status message includes counts when extras present."""
        f = tmp_path / "both.docx"
        f.write_text("dummy")
        with patch("app.pipeline.parsing.parser.DocxDocument") as mock_docx_cls:
            docx = _make_docx_mock()
            mock_docx_cls.return_value = docx
            p = DocxParser()

            hf = Block(block_id="hf_001", text="H", index=0, block_type=BlockType.UNKNOWN)
            note = Block(block_id="fn_001", text="F", index=0, block_type=BlockType.UNKNOWN)
            with (
                patch.object(p, "_extract_headers_and_footers", return_value=[hf]),
                patch.object(p, "_extract_footnotes_and_endnotes", return_value=[note]),
            ):
                doc = p.parse(str(f), "doc1")
                last_stage = doc.processing_history[-1]
                assert "notes" in last_stage.message
                assert "header/footers" in last_stage.message

    def test_parse_with_tables(self, tmp_path):
        """Lines 401-413: CT_Tbl branch in _extract_body_content."""
        from docx.oxml.table import CT_Tbl

        f = tmp_path / "with_tbl.docx"
        f.write_text("dummy")
        with patch("app.pipeline.parsing.parser.DocxDocument") as mock_docx_cls:
            tbl_element = MagicMock(spec=CT_Tbl)
            docx = _make_docx_mock(body_elements=[tbl_element])
            mock_docx_cls.return_value = docx
            p = DocxParser()

            fake_table = MagicMock(spec=Table)
            fake_table.table_id = "tbl_000"

            with patch("app.pipeline.parsing.parser.DocxTable") as mock_tbl_cls, \
                 patch("app.pipeline.parsing.parser.TableExtractor") as mock_ext_cls:
                mock_tbl = MagicMock()
                mock_tbl_cls.return_value = mock_tbl
                mock_extractor = MagicMock()
                mock_ext_cls.return_value = mock_extractor
                mock_extractor.extract.return_value = fake_table

                doc = p.parse(str(f), "doc1")
                assert len(doc.tables) == 1
                assert doc.tables[0].table_id == "tbl_000"

    def test_parse_with_inline_images(self, tmp_path):
        """Lines 372->378, 381->388, 383, 386: block + inline images."""
        f = tmp_path / "with_imgs.docx"
        f.write_text("dummy")
        with patch("app.pipeline.parsing.parser.DocxDocument") as mock_docx_cls:
            from docx.oxml.text.paragraph import CT_P

            p_elem = MagicMock(spec=CT_P)
            docx = _make_docx_mock(body_elements=[p_elem])
            mock_docx_cls.return_value = docx

            # Use _make_paragraph which returns a MagicMock, but we need real Block.metadata
            # So we patch _extract_paragraph to return a real Block with real metadata
            p = DocxParser()

            fake_block = Block(
                block_id="b001", text="With image", index=0,
                block_type=BlockType.UNKNOWN,
            )
            fake_figure = MagicMock(spec=Figure)
            fake_figure.metadata = {}

            with patch("app.pipeline.parsing.parser.DocxParagraph") as mock_para_cls, \
                 patch.object(p, "_extract_paragraph", return_value=fake_block), \
                 patch.object(p, "_extract_inline_images", return_value=[fake_figure]), \
                 patch.object(p, "_extract_equations", return_value=[]):
                mock_para_cls.return_value = MagicMock()
                doc = p.parse(str(f), "doc1")
                assert len(doc.figures) == 1
                assert fake_block.metadata.get("has_figure") is True

    def test_parse_with_equations(self, tmp_path):
        """Lines 393-398: equations attached to a block."""
        f = tmp_path / "with_eqns.docx"
        f.write_text("dummy")
        with patch("app.pipeline.parsing.parser.DocxDocument") as mock_docx_cls:
            from docx.oxml.text.paragraph import CT_P

            p_elem = MagicMock(spec=CT_P)
            docx = _make_docx_mock(body_elements=[p_elem])
            mock_docx_cls.return_value = docx

            p = DocxParser()
            fake_eqn = MagicMock(spec=Equation)
            fake_eqn.metadata = {}
            fake_block = Block(
                block_id="b001", text="E=mc2", index=0,
                block_type=BlockType.UNKNOWN,
            )

            with patch("app.pipeline.parsing.parser.DocxParagraph"), \
                 patch.object(p, "_extract_paragraph", return_value=fake_block), \
                 patch.object(p, "_extract_inline_images", return_value=[]), \
                 patch.object(p, "_extract_equations", return_value=[fake_eqn]):
                doc = p.parse(str(f), "doc1")
                assert len(doc.equations) == 1
                assert fake_block.metadata.get("has_equation") is True


# ---------------------------------------------------------------------------
# _extract_core_properties – line 339
# ---------------------------------------------------------------------------

class TestCorePropertiesGaps:
    """Cover the created-date branch."""

    def test_core_properties_with_created(self):
        """Line 339: core_props.created is truthy."""
        p = DocxParser()
        docx = MagicMock()
        docx.core_properties = MagicMock(
            title="T", author="A", subject="S", keywords="kw1; kw2",
            created=datetime(2024, 1, 15, tzinfo=UTC),
        )
        meta = p._extract_core_properties(docx)
        assert meta.publication_date is not None
        # keywords with semicolons should be handled too
        assert "kw1" in meta.keywords


# ---------------------------------------------------------------------------
# _extract_footnotes_and_endnotes – lines 184-270
# ---------------------------------------------------------------------------

class TestFootnotesEndnotesGaps:
    """Cover footnote/endnote extraction with real content and edge-cases."""

    def _make_footnotes_part(self, text: str = "Footnote text", fn_id: str = "1") -> MagicMock:
        part = MagicMock()
        fn_element = MagicMock()
        fn_element.get.return_value = fn_id
        p_element = MagicMock()
        r_element = MagicMock()
        t_element = MagicMock()
        t_element.text = text
        r_element.find.return_value = t_element
        p_element.findall.return_value = [r_element]
        fn_element.findall.return_value = [p_element]
        part.element.findall.return_value = [fn_element]
        return part

    def test_footnotes_with_content(self):
        """Lines 184->188, 190-230: footnote part exists with text."""
        p = DocxParser()
        docx = MagicMock()
        docx.part = MagicMock()
        docx.part.footnotes_part = self._make_footnotes_part("A real footnote", "1")
        docx.part.endnotes_part = None

        blocks = p._extract_footnotes_and_endnotes(docx)
        assert len(blocks) == 1
        assert blocks[0].text == "A real footnote"
        assert blocks[0].metadata.get("is_footnote") is True
        assert blocks[0].metadata.get("footnote_id") == "1"

    def test_footnotes_skips_empty_text(self):
        """Lines 207-209: footnote paragraph with no text -> skip."""
        p = DocxParser()
        docx = MagicMock()
        docx.part = MagicMock()
        docx.part.footnotes_part = self._make_footnotes_part("   ", "2")
        docx.part.endnotes_part = None

        blocks = p._extract_footnotes_and_endnotes(docx)
        assert len(blocks) == 0

    def test_footnotes_exception_caught(self):
        """Line 228-230: footnote extraction raises -> logged, not crashed."""
        p = DocxParser()
        docx = MagicMock()
        docx.part = MagicMock()
        # Setting footnotes_part so that part is not None but element raises
        fn_part = MagicMock()
        fn_part.element.findall.side_effect = RuntimeError("footnote error")
        docx.part.footnotes_part = fn_part
        docx.part.endnotes_part = None

        blocks = p._extract_footnotes_and_endnotes(docx)
        assert blocks == []

    def test_endnotes_with_content(self):
        """Lines 235->239, 240-270: endnote part exists with text."""
        p = DocxParser()
        docx = MagicMock()
        docx.part = MagicMock()
        docx.part.footnotes_part = None

        en_part = self._make_footnotes_part("An endnote", "5")
        docx.part.endnotes_part = en_part

        blocks = p._extract_footnotes_and_endnotes(docx)
        assert len(blocks) == 1
        assert blocks[0].text == "An endnote"
        assert blocks[0].metadata.get("is_endnote") is True
        assert blocks[0].metadata.get("endnote_id") == "5"

    def test_endnotes_skips_empty_text(self):
        """Lines 251-252: endnote paragraph with no text -> skip."""
        p = DocxParser()
        docx = MagicMock()
        docx.part = MagicMock()
        docx.part.footnotes_part = None
        en_part = self._make_footnotes_part("", "6")
        docx.part.endnotes_part = en_part
        blocks = p._extract_footnotes_and_endnotes(docx)
        assert len(blocks) == 0

    def test_endnotes_exception_caught(self):
        """Line 269-270: endnote exception caught."""
        p = DocxParser()
        docx = MagicMock()
        docx.part = MagicMock()
        docx.part.footnotes_part = None
        en_part = MagicMock()
        en_part.element.findall.side_effect = RuntimeError("endnote error")
        docx.part.endnotes_part = en_part
        blocks = p._extract_footnotes_and_endnotes(docx)
        assert blocks == []

    @patch("app.pipeline.parsing.parser.hasattr")
    def test_footnotes_no_part_attr(self, mock_hasattr):
        """Line 184: hasattr(docx, 'part') is False."""
        mock_hasattr.side_effect = lambda obj, name: name != "part"
        p = DocxParser()
        docx = MagicMock()
        blocks = p._extract_footnotes_and_endnotes(docx)
        assert blocks == []


# ---------------------------------------------------------------------------
# _extract_headers_and_footers – lines 282-301
# ---------------------------------------------------------------------------

class TestHeaderFooterGaps:
    """Cover header/footer extraction with content."""

    def test_header_with_content(self):
        """Lines 282->293, 284-290: header paragraph yields a block."""
        p = DocxParser()
        docx = MagicMock()
        section = MagicMock()
        header_para = _make_paragraph(text="Page Header", style_name="Header")
        section.header.paragraphs = [header_para]
        section.footer.paragraphs = []
        docx.sections = [section]

        blocks = p._extract_headers_and_footers(docx)
        assert len(blocks) == 1
        assert blocks[0].text == "Page Header"
        assert blocks[0].metadata.get("is_header") is True
        assert blocks[0].metadata.get("section_index") == 0

    def test_footer_with_content(self):
        """Lines 293->280, 295-301: footer paragraph yields a block."""
        p = DocxParser()
        docx = MagicMock()
        section = MagicMock()
        section.header.paragraphs = []
        footer_para = _make_paragraph(text="Page Footer", style_name="Footer")
        section.footer.paragraphs = [footer_para]
        docx.sections = [section]

        blocks = p._extract_headers_and_footers(docx)
        assert len(blocks) == 1
        assert blocks[0].text == "Page Footer"
        assert blocks[0].metadata.get("is_footer") is True

    def test_header_and_footer_with_multiple_sections(self):
        """Lines 293->280: multiple sections loop."""
        p = DocxParser()
        docx = MagicMock()
        s1 = MagicMock()
        s1.header.paragraphs = [_make_paragraph(text="H1")]
        s1.footer.paragraphs = [_make_paragraph(text="F1")]
        s2 = MagicMock()
        s2.header.paragraphs = [_make_paragraph(text="H2")]
        s2.footer.paragraphs = [_make_paragraph(text="F2")]
        docx.sections = [s1, s2]

        blocks = p._extract_headers_and_footers(docx)
        assert len(blocks) == 4

    def test_headers_exception_caught(self):
        """Line 302-303: exception in header/footer extraction."""
        p = DocxParser()
        docx = MagicMock()
        docx.sections = [None]  # iterating causes AttributeError
        blocks = p._extract_headers_and_footers(docx)
        assert blocks == []


# ---------------------------------------------------------------------------
# _extract_paragraph – hyperlinks / footnote_refs / list_info gaps
# ---------------------------------------------------------------------------

class TestParagraphMetadataGaps:
    """Cover hyperlinks, footnote_refs, and list_info truthy branches."""

    def test_paragraph_with_hyperlinks(self):
        """Line 463: hyperlinks list is non-empty."""
        p = DocxParser()
        para = _make_paragraph(text="Link here")
        link = {"text": "click", "url": "https://example.com"}
        with patch.object(p, "_extract_hyperlinks", return_value=[link]):
            block = p._extract_paragraph(para)
            assert "hyperlinks" in block.metadata
            assert block.metadata["hyperlinks"][0]["url"] == "https://example.com"

    def test_paragraph_with_footnote_refs(self):
        """Line 467: footnote_refs list is non-empty."""
        p = DocxParser()
        para = _make_paragraph(text="Ref here")
        with patch.object(p, "_extract_note_references", return_value=["1", "2"]):
            block = p._extract_paragraph(para)
            assert "footnote_refs" in block.metadata
            assert "1" in block.metadata["footnote_refs"]


# ---------------------------------------------------------------------------
# _extract_hyperlinks – edge cases (lines 484-500)
# ---------------------------------------------------------------------------

class TestHyperlinksGaps:
    """Cover missing r_id, empty text, KeyError, and outer exception."""

    def test_no_rid(self):
        """Line 484->481: hyperlink.get(qn('r:id')) returns None/falsy."""
        p = DocxParser()
        para = MagicMock()
        hyperlink = MagicMock()
        r_id_mock = MagicMock()
        r_id_mock.get.return_value = None if True else None
        hyperlink.get = MagicMock(return_value=None)
        para._element.findall.return_value = [hyperlink]
        links = p._extract_hyperlinks(para)
        assert links == []

    def test_text_node_is_none(self):
        """Line 492->491: run text_node.text is None -> skip."""
        p = DocxParser()
        para = MagicMock()
        hyperlink = MagicMock()
        hyperlink.get.return_value = "rId1"
        run_elem = MagicMock()
        text_node = MagicMock()
        text_node.text = None
        run_elem.findall.return_value = [text_node]
        hyperlink.findall.return_value = [run_elem]
        para._element.findall.return_value = [hyperlink]
        para.part.rels = {"rId1": MagicMock(target_ref="https://x.com")}
        links = p._extract_hyperlinks(para)
        # text is empty, so link is not added
        assert links == []

    def test_text_empty_after_strip(self):
        """Line 495->481: joined text is empty -> skip."""
        p = DocxParser()
        para = MagicMock()
        hyperlink = MagicMock()
        hyperlink.get.return_value = "rId1"
        run_elem = MagicMock()
        text_node = MagicMock()
        text_node.text = "   "
        run_elem.findall.return_value = [text_node]
        hyperlink.findall.return_value = [run_elem]
        para._element.findall.return_value = [hyperlink]
        para.part.rels = {"rId1": MagicMock(target_ref="https://x.com")}
        links = p._extract_hyperlinks(para)
        assert links == []

    def test_key_error_on_resolve(self):
        """Lines 497-500: KeyError/AttributeError inside inner try."""
        p = DocxParser()
        para = MagicMock()
        hyperlink = MagicMock()
        hyperlink.get.return_value = "rId_missing"
        para._element.findall.return_value = [hyperlink]
        para.part.rels = {}  # rId_missing not in rels -> KeyError
        # Should not propagate; inner except catches and continues
        links = p._extract_hyperlinks(para)
        assert links == []

    def test_outer_exception_caught(self):
        """Lines 499-500: outer try-except catches unexpected errors."""
        p = DocxParser()
        para = MagicMock()
        para._element.findall.side_effect = RuntimeError("outer boom")
        links = p._extract_hyperlinks(para)
        assert links == []

    def test_hyperlink_with_text_added(self):
        """Happy path where text is non-empty -> link is appended."""
        p = DocxParser()
        para = MagicMock()
        hyperlink = MagicMock()
        hyperlink.get.return_value = "rId1"
        run_elem = MagicMock()
        text_node = MagicMock()
        text_node.text = "Click Here"
        run_elem.findall.return_value = [text_node]
        hyperlink.findall.return_value = [run_elem]
        para._element.findall.return_value = [hyperlink]
        para.part.rels = {"rId1": MagicMock(target_ref="https://example.com")}
        links = p._extract_hyperlinks(para)
        assert len(links) == 1
        assert links[0]["text"] == "Click Here"
        assert links[0]["url"] == "https://example.com"


# ---------------------------------------------------------------------------
# _extract_note_references – lines 509-512
# ---------------------------------------------------------------------------

class TestNoteReferencesGaps:
    """Cover missing note_id and exception branches."""

    def test_no_note_id(self):
        """Line 509->507: note_ref.get(qn('w:id')) returns None."""
        p = DocxParser()
        para = MagicMock()
        note_ref = MagicMock()
        note_ref.get.return_value = None
        para._element.findall.return_value = [note_ref]
        refs = p._extract_note_references(para, "w:footnoteReference")
        assert refs == []

    def test_exception_caught(self):
        """Lines 511-512: findall raises -> warning logged, empty returned."""
        p = DocxParser()
        para = MagicMock()
        para._element.findall.side_effect = RuntimeError("boom")
        refs = p._extract_note_references(para, "w:footnoteReference")
        assert refs == []


# ---------------------------------------------------------------------------
# _get_list_info – lines 525-551
# ---------------------------------------------------------------------------

class TestListInfoGaps:
    """Cover missing ilvl, missing numId, style fallback, and exception."""

    def test_numpr_no_ilvl(self):
        """Line 528: numPr exists but ilvl missing -> default level 0."""
        from lxml import etree
        p = DocxParser()
        para = MagicMock()
        xml = '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:numPr><w:numId w:val="7"/></w:numPr></w:pPr>'
        pPr = etree.fromstring(xml)
        para._element = MagicMock()
        para._element.find.return_value = pPr
        result = p._get_list_info(para)
        assert result is not None
        assert result["is_list_item"] is True
        assert result["list_level"] == 0  # default
        assert result["list_id"] == "7"

    def test_numpr_with_numid(self):
        """Lines 531->534: numId is present -> list_id set."""
        from lxml import etree
        p = DocxParser()
        para = MagicMock()
        xml = '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:numPr><w:ilvl w:val="2"/><w:numId w:val="5"/></w:numPr></w:pPr>'
        pPr = etree.fromstring(xml)
        para._element = MagicMock()
        para._element.find.return_value = pPr
        result = p._get_list_info(para)
        assert result is not None
        assert result["list_level"] == 2
        assert result["list_id"] == "5"

    def test_style_fallback_no_number(self):
        """Lines 538->552: pStyle matches but no trailing number -> level 0."""
        from lxml import etree
        p = DocxParser()
        para = MagicMock()
        xml = '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pStyle w:val="ListBullet"/></w:pPr>'
        pPr = etree.fromstring(xml)
        para._element = MagicMock()
        para._element.find.return_value = pPr
        result = p._get_list_info(para)
        assert result is not None
        assert result["is_list_item"] is True
        assert result["list_level"] == 0

    def test_style_fallback_with_number(self):
        """Lines 540->552: style name has trailing number -> level = number - 1."""
        from lxml import etree
        p = DocxParser()
        para = MagicMock()
        xml = '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pStyle w:val="ListNumber3"/></w:pPr>'
        pPr = etree.fromstring(xml)
        para._element = MagicMock()
        para._element.find.return_value = pPr
        result = p._get_list_info(para)
        assert result is not None
        assert result["list_level"] == 2  # 3 - 1

    def test_style_fallback_only_number_matches(self):
        """Lines 540->552: style name only matches on 'number' keyword."""
        from lxml import etree
        p = DocxParser()
        para = MagicMock()
        xml = '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pStyle w:val="Numbered1"/></w:pPr>'
        pPr = etree.fromstring(xml)
        para._element = MagicMock()
        para._element.find.return_value = pPr
        result = p._get_list_info(para)
        assert result is not None
        assert result["is_list_item"] is True

    def test_exception_caught(self):
        """Lines 550-551: exception in _get_list_info -> logged, None returned."""
        p = DocxParser()
        para = MagicMock()
        para._element.find.side_effect = RuntimeError("list error")
        result = p._get_list_info(para)
        assert result is None


# ---------------------------------------------------------------------------
# _extract_paragraph_style – lines 578-603
# ---------------------------------------------------------------------------

class TestParagraphStyleGaps:
    """Cover empty-run skip, None checks, and style fallback branches."""

    def test_skip_empty_run(self):
        """Line 578->577: first run has empty text -> skipped."""
        p = DocxParser()
        para = _make_paragraph(text="real", style_name="Normal")
        empty_run = _make_run(text="   ", bold=None, italic=None, underline=None)
        content_run = _make_run(
            text="real", bold=False, italic=False, underline=False,
            font_name=None, font_size_pt=None,
        )
        para.runs = [empty_run, content_run]

        block = p._extract_paragraph(para)
        assert block is not None
        assert block.style.bold is False

    def test_bold_is_none(self):
        """Line 579->581: run.bold is None -> skip assignment."""
        p = DocxParser()
        para = _make_paragraph(text="t")
        run = _make_run(text="t", bold=None, italic=False, underline=False)
        para.runs = [run]
        block = p._extract_paragraph(para)
        assert block.style.bold is False  # default

    def test_italic_is_none(self):
        """Line 581->583: run.italic is None -> skip assignment."""
        p = DocxParser()
        para = _make_paragraph(text="t")
        run = _make_run(text="t", bold=False, italic=None, underline=False)
        para.runs = [run]
        block = p._extract_paragraph(para)
        assert block.style.italic is False

    def test_underline_set_name_none(self):
        """Lines 583->585: underline is not None, but font.name is None."""
        p = DocxParser()
        para = _make_paragraph(text="t")
        run = _make_run(text="t", bold=False, italic=False, underline=False, font_name=None)
        para.runs = [run]
        block = p._extract_paragraph(para)
        assert block.style.underline is True  # underline=False acts as "not None" -> True
        assert block.style.font_name is None

    def test_underline_is_none(self):
        """Line 583: underline is None -> skip."""
        p = DocxParser()
        para = _make_paragraph(text="t")
        run = _make_run(text="t", bold=False, italic=False, underline=None)
        para.runs = [run]
        block = p._extract_paragraph(para)
        assert block.style.underline is False

    def test_fallback_italic_true(self):
        """Line 597: no runs, style.font.italic is True."""
        p = DocxParser()
        para = _make_paragraph(
            text="styled", style_name="Body", runs=[],
            font_bold=False, font_italic=True, font_name="Arial", font_size_pt=12.0,
        )
        block = p._extract_paragraph(para)
        assert block.style.italic is True

    def test_fallback_attribute_error(self):
        """Lines 602-603: style has no font -> AttributeError caught."""
        p = DocxParser()
        para = MagicMock()
        para.text = "fallback"
        para.alignment = None
        para.runs = []
        para._element = MagicMock()
        para._element.findall.return_value = []
        para._element.find.return_value = None
        style_mock = MagicMock()
        style_mock.name = None
        style_mock.font = None  # None.bold raises AttributeError
        para.style = style_mock

        block = p._extract_paragraph(para)
        assert block is not None


# ---------------------------------------------------------------------------
# _extract_inline_images – lines 629-658
# ---------------------------------------------------------------------------

class TestInlineImagesGaps:
    """Cover the inline image extraction loop body."""

    def test_inline_images_finds_shapes(self):
        """Lines 629-658: run has inline shapes -> extraction proceeds."""
        p = DocxParser()
        para = MagicMock()
        run = MagicMock()
        run._element = MagicMock()
        # First findall returns inline shapes, second returns anchored shapes
        inline_shape = MagicMock()
        run._element.findall.side_effect = [
            [inline_shape],  # inline
            [],              # anchor
        ]
        run.part = MagicMock()
        para.runs = [run]

        fake_figure = MagicMock(spec=Figure)
        with patch.object(p, "_extract_image_from_inline", return_value=fake_figure):
            figures = p._extract_inline_images(para)
            assert len(figures) == 1

    def test_inline_images_no_shapes(self):
        """Both findall calls return empty lists -> no figures."""
        p = DocxParser()
        para = MagicMock()
        run = MagicMock()
        run._element = MagicMock()
        run._element.findall.side_effect = [[], []]
        run.part = MagicMock()
        para.runs = [run]

        figures = p._extract_inline_images(para)
        assert figures == []

    def test_inline_images_exception_logged(self):
        """Line 655-658: exception during extraction is caught."""
        p = DocxParser()
        para = MagicMock()
        run = MagicMock()
        run._element = MagicMock()
        run._element.findall.side_effect = RuntimeError("image error")
        para.runs = [run]

        figures = p._extract_inline_images(para)
        assert figures == []

    def test_inline_images_no_element(self):
        """Line 629: run has no _element."""
        p = DocxParser()
        para = MagicMock()
        run = MagicMock()
        del run._element  # simulate hasattr returning False
        para.runs = [run]

        figures = p._extract_inline_images(para)
        assert figures == []


# ---------------------------------------------------------------------------
# _extract_image_from_inline – lines 702-728
# ---------------------------------------------------------------------------

class TestImageFromInlineGaps:
    """Cover extent dimension edge-cases and exception handler."""

    def test_no_extent(self):
        """Line 702->712: extent is None -> width/height remain None."""
        p = DocxParser()
        inline = MagicMock()
        blip = MagicMock()
        blip.get.return_value = "rId1"
        inline.find.side_effect = lambda x, ns=None: blip if "blip" in x else None

        part = MagicMock()
        image_part = MagicMock()
        image_part.blob = b"data"
        image_part.content_type = "image/png"
        part.related_parts = {"rId1": image_part}

        result = p._extract_image_from_inline(inline, part)
        assert result is not None
        assert result.width is None
        assert result.height is None

    def test_cx_only(self):
        """Lines 706->708: cx present but cy absent."""
        p = DocxParser()
        inline = MagicMock()
        blip = MagicMock()
        blip.get.return_value = "rId1"
        extent = MagicMock()
        extent.get.side_effect = lambda k: {"cx": "100000", "cy": None}.get(k)
        inline.find.side_effect = lambda x, ns=None: blip if "blip" in x else extent

        part = MagicMock()
        image_part = MagicMock()
        image_part.blob = b"data"
        image_part.content_type = "image/png"
        part.related_parts = {"rId1": image_part}

        result = p._extract_image_from_inline(inline, part)
        assert result is not None
        assert result.width is not None
        assert result.height is None

        # Verify width conversion
        expected = 100000 / 9525
        assert abs(result.width - expected) < 0.01

    def test_cy_only(self):
        """Lines 708->712: cx absent but cy present."""
        p = DocxParser()
        inline = MagicMock()
        blip = MagicMock()
        blip.get.return_value = "rId1"
        extent = MagicMock()
        extent.get.side_effect = lambda k: {"cx": None, "cy": "80000"}.get(k)
        inline.find.side_effect = lambda x, ns=None: blip if "blip" in x else extent

        part = MagicMock()
        image_part = MagicMock()
        image_part.blob = b"data"
        image_part.content_type = "image/png"
        part.related_parts = {"rId1": image_part}

        result = p._extract_image_from_inline(inline, part)
        assert result is not None
        assert result.width is None
        assert result.height is not None

        expected = 80000 / 9525
        assert abs(result.height - expected) < 0.01

    def test_exception_caught(self):
        """Lines 726-728: exception during image extraction -> return None."""
        p = DocxParser()
        inline = MagicMock()
        inline.find.side_effect = RuntimeError("image extraction failed")
        part = MagicMock()
        result = p._extract_image_from_inline(inline, part)
        assert result is None


# ---------------------------------------------------------------------------
# _extract_table – lines 765-773
# ---------------------------------------------------------------------------

class TestTableGaps:
    """Cover the _extract_table method."""

    def test_extract_table_direct(self):
        """Lines 765-773: instantiate TableExtractor and delegate."""
        p = DocxParser()
        table_mock = MagicMock()

        with patch("app.pipeline.parsing.parser.TableExtractor") as mock_ext_cls:
            mock_extractor = MagicMock()
            mock_ext_cls.return_value = mock_extractor
            fake_table = MagicMock(spec=Table)
            fake_table.table_id = "tbl_007"
            mock_extractor.extract.return_value = fake_table

            result = p._extract_table(table_mock, block_index=42)
            assert result.table_id == "tbl_007"
            assert p.table_counter == 1  # counter incremented

            # Verify correct arguments passed to extract()
            call_args = mock_extractor.extract.call_args
            assert call_args[0][0] is table_mock  # docx_table
            assert call_args[0][3] == 42          # block_index


# ---------------------------------------------------------------------------
# _extract_equations – lines 789-804
# ---------------------------------------------------------------------------

class TestEquationExtractionGaps:
    """Cover oMathPara loop, inline skip, and math element extraction."""

    def test_omathpara_block_equations(self):
        """Lines 789-792: oMathPara loop extracts block equations."""

        p = DocxParser()
        para = MagicMock()
        om_para = MagicMock()
        om_math = MagicMock()
        om_para.findall.return_value = [om_math]
        para._element = MagicMock()

        # side_effect for three findall calls on para._element:
        # call 1: qn('m:oMathPara') -> [om_para]
        # call 2: qn('m:oMath') -> [] (no inline)
        # call 3: any-check inside inline loop -> [] (none match)
        para._element.findall.side_effect = [
            [om_para],   # oMathPara
            [],           # oMath (inline)
            [],           #  nested oMath inside oMathPara (inline skip check)
        ]

        fake_eqn = MagicMock(spec=Equation)
        fake_eqn.equation_id = "eqn_000"
        with patch.object(p, "_extract_math_element", return_value=fake_eqn):
            equations = p._extract_equations(para)
            assert len(equations) == 1
            p._extract_math_element.assert_called_once_with(om_math, is_block=True)

    def test_inline_equations_skip_processed(self):
        """Lines 799-804: inline oMath already processed -> continue."""

        p = DocxParser()
        para = MagicMock()
        om_math = MagicMock()
        para._element = MagicMock()

        # First you need to understand the check:
        # if any(om is b_om for b_om in para._element.findall(f".//{qn('m:oMathPara')}/{qn('m:oMath')}")):
        #     continue
        # This means: for each inline om, check if it appears as a child of any oMathPara.
        # We want this check to be True -> continue.
        # So para._element.findall for the nested path must return a list containing the same om.

        # side_effect:
        # call 1: oMathPara -> [] (no block equations)
        # call 2: oMath -> [om_math] (one inline equation)
        # call 3: nested path -> [om_math] (om_math is in the list, so "any" is True -> continue)
        para._element.findall.side_effect = [
            [],            # oMathPara (no block)
            [om_math],     # oMath (inline found)
            [om_math],     # nested: om IS b_om -> any -> True -> continue
        ]

        equations = p._extract_equations(para)
        assert equations == []  # because equation was skipped

    def test_inline_equations_not_processed(self):
        """Line 799: inline oMath not in oMathPara -> extracted."""

        p = DocxParser()
        para = MagicMock()
        om_math = MagicMock()
        para._element = MagicMock()

        para._element.findall.side_effect = [
            [],            # oMathPara -> no block
            [om_math],     # oMath -> inline
            [],            # nested -> not found -> any() is False -> proceed
        ]

        fake_eqn = MagicMock(spec=Equation)
        fake_eqn.equation_id = "eqn_001"
        with patch.object(p, "_extract_math_element", return_value=fake_eqn):
            equations = p._extract_equations(para)
            assert len(equations) == 1
            p._extract_math_element.assert_called_once_with(om_math, is_block=False)


# ---------------------------------------------------------------------------
# _extract_math_element – lines 810-834
# ---------------------------------------------------------------------------

class TestMathElementGaps:
    """Cover text extraction, fallback, and exception handling."""

    def test_extract_math_success(self):
        """Lines 810-831: successful equation extraction."""
        p = DocxParser()
        om_element = MagicMock()

        # findall with m:t returns non-empty
        om_element.findall.return_value = [MagicMock(text="x + y")]

        with patch("lxml.etree.tostring", return_value="<math>...</math>"):
            result = p._extract_math_element(om_element, is_block=True)

        assert result is not None
        assert result.text == "x + y"
        assert result.is_block is True
        assert result.omml == "<math>...</math>"

    def test_extract_math_fallback_text(self):
        """Lines 818-819: m:t empty -> fall back to w:t."""
        p = DocxParser()
        om_element = MagicMock()

        # First findall (m:t) returns empty; second (w:t) returns content
        om_element.findall.side_effect = [
            [],                                      # m:t -> empty
            [MagicMock(text="fallback eq")],          # w:t -> found
        ]

        with patch("lxml.etree.tostring", return_value="<math>...</math>"):
            result = p._extract_math_element(om_element, is_block=False)

        assert result is not None
        assert result.text == "fallback eq"

    def test_extract_math_no_text(self):
        """Both m:t and w:t empty -> result with empty text."""
        p = DocxParser()
        om_element = MagicMock()
        om_element.findall.return_value = []

        with patch("lxml.etree.tostring", return_value="<math/>"):
            result = p._extract_math_element(om_element, is_block=False)

        assert result is not None
        assert result.text == ""

    def test_extract_math_exception(self):
        """Lines 832-834: exception during extraction -> None returned."""
        p = DocxParser()
        om_element = MagicMock()
        om_element.findall.side_effect = RuntimeError("math error")

        result = p._extract_math_element(om_element, is_block=False)
        assert result is None


# ---------------------------------------------------------------------------
# _extract_body_content – miscellaneous edge cases
# ---------------------------------------------------------------------------

class TestBodyContentGaps:
    """Cover edge cases in _extract_body_content."""

    def test_empty_paragraph_skipped(self):
        """Line 372: _extract_paragraph returns None -> skip."""
        from docx.oxml.text.paragraph import CT_P

        p = DocxParser()
        docx = MagicMock()
        p_elem = MagicMock(spec=CT_P)
        docx.element.body = [p_elem]

        para = _make_paragraph(text="")
        # Simulate no block returned for an empty/whitespace paragraph
        p._extract_paragraph = MagicMock(return_value=None)
        # _extract_inline_images and _extract_equations also called but shouldn't crash
        p._extract_inline_images = MagicMock(return_value=[])
        p._extract_equations = MagicMock(return_value=[])

        with patch("app.pipeline.parsing.parser.DocxParagraph", return_value=para):
            blocks, figures, tables, equations = p._extract_body_content(docx)
            assert len(blocks) == 0
            # Even with no block, inline images/equations can still be accumulated
            # (lines 378, 388: figures are added regardless)
            assert len(figures) == 0
