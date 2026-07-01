# SPDX-License-Identifier: MIT
# Copyright (c) 2026 ScholarForm AI

"""
Shared fixtures and mock utilities for table/figure pipeline tests.
"""

from __future__ import annotations
from typing import List, Optional, Any, Dict
from unittest.mock import MagicMock, PropertyMock
from dataclasses import dataclass, field
import pytest


# ---------------------------------------------------------------------------
# python-docx mock helpers
# ---------------------------------------------------------------------------

@dataclass
class MockDocxCell:
    """Simulates a python-docx Cell."""
    text: str = ""
    paragraphs: List = field(default_factory=list)
    bold: bool = False
    _tc: Any = None
    _element: Any = None
    _parent: Any = None
    _tbl: Any = None

    def __post_init__(self):
        if self._tc is None:
            self._tc = MagicMock()
            self._tc.iter.return_value = []
        if self._element is None:
            self._element = MagicMock()
            self._element.findall.return_value = []
        if self._parent is None:
            self._parent = MagicMock()
        # Ensure paragraphs reflects bold status
        run = MagicMock()
        run.bold = self.bold
        para = MagicMock()
        para.runs = [run]
        self.paragraphs = [para]


@dataclass
class MockDocxRow:
    """Simulates a python-docx Row."""
    cells: List[MockDocxCell] = field(default_factory=list)


@dataclass
class MockDocxTable:
    """Simulates a python-docx Table."""
    rows: List[MockDocxRow] = field(default_factory=list)


def make_mock_table(
    data: List[List[str]],
    header_row: bool = False,
    nested_tables_at: Optional[Dict[str, List[List[List[str]]]]] = None,
) -> MockDocxTable:
    """
    Build a MockDocxTable from raw string data.

    Args:
        data: 2D list of cell texts.
        header_row: If True, first row cells get bold=True.
        nested_tables_at: Dict mapping "row,col" -> list of table data 3D lists.
                          e.g. {"0,1": [[["a","b"],["c","d"]]]}
    """
    if nested_tables_at is None:
        nested_tables_at = {}

    rows = []
    for r_idx, row_data in enumerate(data):
        cells = []
        for c_idx, text in enumerate(row_data):
            is_bold = header_row and r_idx == 0
            cell = MockDocxCell(text=text, bold=is_bold)

            # Inject nested table XML into _element.findall
            key = f"{r_idx},{c_idx}"
            if key in nested_tables_at:
                nested_tables_data = nested_tables_at[key]
                from unittest.mock import MagicMock
                tbl_elements = []
                for nt_data in nested_tables_data:
                    inner_tbl = make_mock_table(nt_data)
                    mock_tbl_xml = MagicMock()
                    # wire recursion hook so MockDocxTable can be wrapped
                    mock_tbl_xml.tag = "w:tbl"
                    # We store the inner mock on the xml so DocxTableWrapper can access it
                    from docx.oxml.ns import qn
                    # The extractor does: tbl_xml.findall(qn('w:tbl'))
                    # So we skip nested on the inner level
                    inner_mock = MagicMock()
                    inner_mock.rows = inner_tbl.rows
                    # We'll store it as a side effect
                    mock_tbl_xml._inner_mock = inner_mock
                    from docx.oxml import OxmlElement
                    tbl_elements.append(mock_tbl_xml)

                # Override findall to return our nested table XML elements
                cell._element.findall.side_effect = lambda ns, elems=tbl_elements: elems if ns.endswith('}tbl') else []

                # For the DocxTableWrapper constructor: need the parent chain
                cell._parent._parent._parent = MagicMock()

            cells.append(cell)
        rows.append(MockDocxRow(cells=cells))

    return MockDocxTable(rows=rows)


# ---------------------------------------------------------------------------
# PIL image helpers
# ---------------------------------------------------------------------------

@pytest.fixture
def temp_image(tmp_path):
    """Create a temporary PNG image for figure analysis tests."""
    from PIL import Image
    path = tmp_path / "test_figure.png"
    img = Image.new("RGB", (800, 600), color=(255, 255, 255))
    img.save(str(path))
    return str(path)


@pytest.fixture
def temp_image_low_res(tmp_path):
    """Create a low-resolution image."""
    from PIL import Image
    path = tmp_path / "low_res.png"
    img = Image.new("RGB", (100, 80), color=(0, 0, 0))
    img.save(str(path))
    return str(path)


@pytest.fixture
def temp_jpeg_image(tmp_path):
    """Create a temporary JPEG image."""
    from PIL import Image
    path = tmp_path / "test_figure.jpg"
    img = Image.new("RGB", (1024, 768), color=(128, 128, 128))
    img.save(str(path), "JPEG", quality=95)
    return str(path)


# ---------------------------------------------------------------------------
# Document / Block / Table / Figure fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def doc_empty():
    from app.models import PipelineDocument
    return PipelineDocument(document_id="test_doc", blocks=[])


@pytest.fixture
def simple_table():
    from app.models import Table, TableCell
    return Table(
        table_id="tbl_001",
        num_rows=2,
        num_cols=2,
        index=0,
        block_index=0,
        cells=[
            TableCell(row=0, col=0, text="A1", bold=True),
            TableCell(row=0, col=1, text="B1", bold=True),
            TableCell(row=1, col=0, text="A2"),
            TableCell(row=1, col=1, text="B2"),
        ],
        data=[["A1", "B1"], ["A2", "B2"]],
        rows=[["A1", "B1"], ["A2", "B2"]],
        has_header=True,
        has_header_row=True,
        header_rows=1,
    )


@pytest.fixture
def simple_figure():
    from app.models import Figure
    return Figure(figure_id="fig_001", index=0)


@pytest.fixture
def doc_with_blocks():
    from app.models import PipelineDocument, Block, BlockType
    return PipelineDocument(
        document_id="doc_blocks",
        blocks=[
            Block(block_id="b1", index=0, text="Introduction", block_type=BlockType.HEADING_1),
            Block(block_id="b2", index=1, text="Some body text here.", block_type=BlockType.BODY),
            Block(block_id="b3", index=2, text="Table 1. Performance results.", block_type=BlockType.BODY),
            Block(block_id="b4", index=3, text="Row data here.", block_type=BlockType.BODY),
            Block(block_id="b5", index=4, text="Figure 1. Architecture diagram.", block_type=BlockType.BODY),
        ],
    )
